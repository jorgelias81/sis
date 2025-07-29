# -*- coding: utf-8 -*-
import sqlite3
from datetime import datetime
import tkinter as tk
from tkinter import ttk, messagebox, simpledialog, filedialog
from tkcalendar import DateEntry
import hashlib
import logging
from logging.handlers import RotatingFileHandler
import os
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd
import webbrowser
import os
from PIL import Image, ImageTk
import threading
import re
import sys

logging.basicConfig(stream=sys.stdout, level=logging.DEBUG)

# Configuración inicial
class Config:
    DB_NAME = "stock_informatico.db"
    LOG_FILE = "stock_system.log"
    MAX_LOG_SIZE = 1024 * 1024  # 1MB
    BACKUP_COUNT = 5
    UI_THEME = "clam"
    PDF_OPTIONS = {
        'pagesize': letter,
        'margins': (20, 20, 20, 20)
    }
    # Nueva paleta de colores profesional
    PRIMARY_COLOR = "#2c3e50"      # Azul oscuro (barra superior)
    SECONDARY_COLOR = "#3498db"    # Azul claro (botones)
    BACKGROUND_COLOR = "#f8f9fa"   # Gris muy claro (fondo)
    TEXT_COLOR = "#2c3e50"         # Azul oscuro (texto)
    SUCCESS_COLOR = "#28a745"      # Verde (éxito)
    WARNING_COLOR = "#ffc107"      # Amarillo (advertencia)
    ERROR_COLOR = "#dc3545"        # Rojo (error)
    HIGHLIGHT_COLOR = "#e9ecef"    # Gris claro (resaltado)
    HEADER_COLOR = "#343a40"       # Gris oscuro (encabezados)
    BORDER_COLOR = "#dee2e6"       # Gris bordes

# Configuración de logging
def setup_logging():
    logger = logging.getLogger()
    logger.setLevel(logging.INFO)
    
    formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    
    file_handler = RotatingFileHandler(
        Config.LOG_FILE, 
        maxBytes=Config.MAX_LOG_SIZE, 
        backupCount=Config.BACKUP_COUNT)
    file_handler.setFormatter(formatter)
    logger.addHandler(file_handler)

setup_logging()



class Database:
    """Capa de acceso a datos con patrón Singleton"""
    _instance = None
    
    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
            cls._instance._initialize()
        return cls._instance
    
    def _initialize(self):
        self.conn = sqlite3.connect(Config.DB_NAME, check_same_thread=False)
        self.conn.execute("PRAGMA foreign_keys = ON")
        self.cursor = self.conn.cursor()
        self._create_tables()
        self._insert_default_data()
        logging.info("Database initialized")
    
    def _create_tables(self):
        try:
            tables = [
                '''CREATE TABLE IF NOT EXISTS tipos_equipo (
                    id INTEGER PRIMARY KEY AUTOINCREMENT, 
                    nombre TEXT NOT NULL UNIQUE
                )''',
                '''CREATE TABLE IF NOT EXISTS marcas (
                    id INTEGER PRIMARY KEY AUTOINCREMENT, 
                    nombre TEXT NOT NULL UNIQUE
                )''',
                '''CREATE TABLE IF NOT EXISTS modelos (
                    id INTEGER PRIMARY KEY AUTOINCREMENT, 
                    id_marca INTEGER NOT NULL, 
                    nombre TEXT NOT NULL, 
                    UNIQUE(id_marca, nombre),
                    FOREIGN KEY (id_marca) REFERENCES marcas(id) ON DELETE CASCADE
                )''',
                '''CREATE TABLE IF NOT EXISTS equipos (
                    id INTEGER PRIMARY KEY AUTOINCREMENT, 
                    pj TEXT NOT NULL, 
                    id_tipo_equipo INTEGER NOT NULL,
                    serie TEXT NOT NULL UNIQUE, 
                    id_marca INTEGER NOT NULL,
                    id_modelo INTEGER NOT NULL,
                    ubicacion TEXT, 
                    fecha_ingreso TEXT NOT NULL, 
                    fecha_salida TEXT, 
                    falla TEXT NOT NULL, 
                    estado TEXT DEFAULT 'En reparación', 
                    observaciones TEXT,
                    FOREIGN KEY (id_tipo_equipo) REFERENCES tipos_equipo(id),
                    FOREIGN KEY (id_marca) REFERENCES marcas(id),
                    FOREIGN KEY (id_modelo) REFERENCES modelos(id)
                )''',
                '''CREATE TABLE IF NOT EXISTS repuestos (
                    id INTEGER PRIMARY KEY AUTOINCREMENT, 
                    id_equipo INTEGER NOT NULL, 
                    nombre TEXT NOT NULL, 
                    cantidad INTEGER NOT NULL, 
                    costo REAL, 
                    FOREIGN KEY (id_equipo) REFERENCES equipos(id) ON DELETE CASCADE
                )''',
                '''CREATE TABLE IF NOT EXISTS usuarios (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    username TEXT UNIQUE NOT NULL,
                    password TEXT NOT NULL,
                    rol TEXT NOT NULL,
                    activo BOOLEAN DEFAULT 1
                )''',
                '''CREATE TABLE IF NOT EXISTS auditoria (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    usuario_id INTEGER,
                    accion TEXT NOT NULL,
                    tabla_afectada TEXT,
                    registro_id INTEGER,
                    fecha TEXT NOT NULL,
                    detalles TEXT,
                    FOREIGN KEY (usuario_id) REFERENCES usuarios(id)
                )''',

                # --- TABLAS PARA TONER ---
                '''CREATE TABLE IF NOT EXISTS marcas_toner (
                    id INTEGER PRIMARY KEY AUTOINCREMENT, 
                    nombre TEXT NOT NULL UNIQUE,
                    fecha_creacion TEXT DEFAULT CURRENT_TIMESTAMP,
                    usuario_id INTEGER,
                    FOREIGN KEY (usuario_id) REFERENCES usuarios(id)
                )''',

                '''CREATE TABLE IF NOT EXISTS modelos_toner (
                    id INTEGER PRIMARY KEY AUTOINCREMENT, 
                    id_marca INTEGER NOT NULL, 
                    nombre TEXT NOT NULL, 
                    fecha_creacion TEXT DEFAULT CURRENT_TIMESTAMP,
                    usuario_id INTEGER,
                    UNIQUE(id_marca, nombre),
                    FOREIGN KEY (id_marca) REFERENCES marcas_toner(id) ON DELETE CASCADE,
                    FOREIGN KEY (usuario_id) REFERENCES usuarios(id)
                )''',

                '''CREATE TABLE IF NOT EXISTS stock_toner (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    id_modelo INTEGER NOT NULL UNIQUE,
                    cantidad INTEGER NOT NULL DEFAULT 0 CHECK(cantidad >= 0),
                    FOREIGN KEY (id_modelo) REFERENCES modelos_toner(id) ON DELETE CASCADE
                )''',

                '''CREATE TABLE IF NOT EXISTS tipos_movimiento_toner (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    nombre TEXT NOT NULL UNIQUE,  -- Ej: "Ingreso", "Retiro", "Envio Recarga", "Recepcion Recarga"
                    afecta_stock INTEGER NOT NULL CHECK(afecta_stock IN (-1, 0, 1))  -- -1 = resta, +1 = suma, 0 = no afecta
                )''',

                '''CREATE TABLE IF NOT EXISTS empresas_recarga (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    nombre TEXT NOT NULL UNIQUE,
                    contacto TEXT,
                    telefono TEXT
                )''',

                '''CREATE TABLE IF NOT EXISTS movimientos_toner (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    id_modelo INTEGER NOT NULL,
                    id_tipo_movimiento INTEGER NOT NULL,
                    cantidad INTEGER NOT NULL CHECK(cantidad > 0),
                    id_responsable INTEGER,  -- Puede ser NULL si es una recarga
                    id_sector INTEGER,       -- Puede ser NULL si es una recarga
                    id_empresa_recarga INTEGER,  -- Solo para movimientos de recarga
                    observaciones TEXT,
                    fecha TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                    usuario_id INTEGER NOT NULL,
                    FOREIGN KEY (id_modelo) REFERENCES modelos_toner(id),
                    FOREIGN KEY (id_tipo_movimiento) REFERENCES tipos_movimiento_toner(id),
                    FOREIGN KEY (id_responsable) REFERENCES usuarios(id),
                    FOREIGN KEY (id_sector) REFERENCES sectores(id),
                    FOREIGN KEY (id_empresa_recarga) REFERENCES empresas_recarga(id),
                    FOREIGN KEY (usuario_id) REFERENCES usuarios(id)
                )''',

                # Tabla adicional para seguimiento de recargas
                '''CREATE TABLE IF NOT EXISTS recargas_toner (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    id_movimiento_envio INTEGER NOT NULL,  -- Movimiento de tipo "Envio Recarga"
                    id_movimiento_recepcion INTEGER,       -- Movimiento de tipo "Recepcion Recarga" (NULL hasta que se reciba)
                    estado TEXT NOT NULL DEFAULT 'Enviado', -- "Enviado", "Recibido", "Cancelado"
                    fecha_envio TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                    fecha_recepcion TEXT,
                    usuario_recepcion INTEGER,
                    FOREIGN KEY (id_movimiento_envio) REFERENCES movimientos_toner(id),
                    FOREIGN KEY (id_movimiento_recepcion) REFERENCES movimientos_toner(id),
                    FOREIGN KEY (usuario_recepcion) REFERENCES usuarios(id)
                 )'''
                 
            ]
            
            for table in tables:
                self.cursor.execute(table)
            self.conn.commit()
            logging.info("Tables created successfully")
        except sqlite3.Error as e:
            logging.error(f"Error creating tables: {e}")
            raise
    
    def _insert_default_data(self):
        try:
            default_tipos = ["CPU", "Impresora", "Monitor", "Switch", "Router", "Otro"]
            default_marcas = ["HP", "Dell", "Lenovo", "Epson", "Brother", "Otro"]
        
            # Insertar tipos de equipo
            for tipo in default_tipos:
                self.cursor.execute(
                    "INSERT OR IGNORE INTO tipos_equipo (nombre) VALUES (?)", (tipo,))
        
            # Insertar marcas
            for marca in default_marcas:
                self.cursor.execute(
                    "INSERT OR IGNORE INTO marcas (nombre) VALUES (?)", (marca,))
        
            # Insertar usuario admin por defecto
            hashed_pw = hashlib.sha256("admin123".encode('utf-8')).hexdigest()
            self.cursor.execute(
                """INSERT OR IGNORE INTO usuarios 
                (username, password, rol, activo) 
                VALUES (?, ?, ?, ?)""",
                ("admin", hashed_pw, "admin", 1))
        
            self.conn.commit()
            logging.info("Default data inserted")
        except sqlite3.Error as e:
            logging.error(f"Error inserting default data: {e}")
            self.conn.rollback()
            raise

        # --- DATOS INICIALES PARA TONER ---
        try:
            # Insertar marcas de toner por defecto
            default_marcas_toner = ["HP", "Brother", "Epson", "Xerox"]
            for marca in default_marcas_toner:
                self.cursor.execute(
                    "INSERT OR IGNORE INTO marcas_toner (nombre) VALUES (?)", 
                    (marca,))
    
            # Insertar modelos de toner para HP como ejemplo
            self.cursor.execute("SELECT id FROM marcas_toner WHERE nombre='HP'")
            hp_id = self.cursor.fetchone()[0]
            modelos_hp = ["CF400X", "CF401X", "CF500X"]
            for modelo in modelos_hp:
                self.cursor.execute(
                    "INSERT OR IGNORE INTO modelos_toner (id_marca, nombre) VALUES (?, ?)",
                    (hp_id, modelo))
        
            self.conn.commit()
            logging.info("Default toner data inserted")
        except sqlite3.Error as e:
            logging.error(f"Error inserting default toner data: {e}")
            self.conn.rollback()
            raise


            
    
    def autenticar_usuario(self, username, password):
        """Autentica un usuario con sus credenciales"""
        try:
            hashed_pw = hashlib.sha256(password.encode('utf-8')).hexdigest()
            
            self.cursor.execute(
                """SELECT id, rol, username 
                FROM usuarios 
                WHERE username=? AND password=? AND activo=1""",
                (username, hashed_pw))
            
            resultado = self.cursor.fetchone()
            
            if resultado:
                logging.info(f"Autenticación exitosa para usuario: {username}")
                return resultado
            else:
                logging.warning(f"Intento fallido de autenticación para usuario: {username}")
                return None
                
        except sqlite3.Error as e:
            logging.error(f"Error en autenticación: {e}")
            return None
    
    def obtener_tipos_equipo(self):
        """Obtiene todos los tipos de equipo disponibles"""
        self.cursor.execute("SELECT id, nombre FROM tipos_equipo ORDER BY nombre")
        return self.cursor.fetchall()
    
    def obtener_marcas(self):
        """Obtiene todas las marcas disponibles"""
        self.cursor.execute("SELECT id, nombre FROM marcas ORDER BY nombre")
        return self.cursor.fetchall()
    
    def agregar_marca(self, nombre, usuario_id):
        """Agrega una nueva marca al sistema"""
        try:
            self.cursor.execute(
                "INSERT INTO marcas (nombre) VALUES (?)", 
                (nombre,))
            
            marca_id = self.cursor.lastrowid
            
            self.registrar_auditoria(
                usuario_id, 'ALTA_MARCA', 'marcas', marca_id,
                f"Nueva marca agregada: {nombre}")
            
            self.conn.commit()
            return marca_id
        except sqlite3.IntegrityError:
            raise ValueError("La marca ya existe")
        except sqlite3.Error as e:
            self.conn.rollback()
            logging.error(f"Error agregando marca: {e}")
            raise
    
    def obtener_modelos(self, id_marca):
        """Obtiene los modelos de una marca específica"""
        self.cursor.execute(
            "SELECT id, nombre FROM modelos WHERE id_marca=? ORDER BY nombre", 
            (id_marca,))
        return self.cursor.fetchall()
    
    def agregar_modelo(self, id_marca, nombre, usuario_id):
        """Agrega un nuevo modelo a una marca"""
        try:
            self.cursor.execute(
                "INSERT INTO modelos (id_marca, nombre) VALUES (?, ?)", 
                (id_marca, nombre))
            
            modelo_id = self.cursor.lastrowid
            
            self.registrar_auditoria(
                usuario_id, 'ALTA_MODELO', 'modelos', modelo_id,
                f"Nuevo modelo agregado: {nombre} para marca ID: {id_marca}")
            
            self.conn.commit()
            return modelo_id
        except sqlite3.IntegrityError:
            raise ValueError("El modelo ya existe para esta marca")
        except sqlite3.Error as e:
            self.conn.rollback()
            logging.error(f"Error agregando modelo: {e}")
            raise
    
    def agregar_equipo(self, datos, usuario_id):
        """Agrega un nuevo equipo al sistema"""
        try:
            query = '''INSERT INTO equipos 
                (pj, id_tipo_equipo, serie, id_marca, id_modelo, ubicacion, 
                 fecha_ingreso, fecha_salida, falla, estado, observaciones) 
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)'''
            self.cursor.execute(query, datos)
            equipo_id = self.cursor.lastrowid
            
            self.registrar_auditoria(
                usuario_id, 'ALTA_EQUIPO', 'equipos', equipo_id,
                f"Equipo ingresado: {datos[2]} (PJ: {datos[0]})")
            
            self.conn.commit()
            return equipo_id
        except sqlite3.IntegrityError as e:
            if "serie" in str(e):
                raise ValueError("El número de serie ya existe")
            raise
        except sqlite3.Error as e:
            self.conn.rollback()
            logging.error(f"Error agregando equipo: {e}")
            raise


    def obtener_resumen_repuestos(self, año=None, mes=None):
        """Obtiene un resumen de repuestos usados agrupados por nombre y cantidad total"""
        query = """
            SELECT 
                r.nombre, 
                SUM(r.cantidad) as cantidad_total,
                strftime('%m', e.fecha_ingreso) as mes,
                strftime('%Y', e.fecha_ingreso) as año
            FROM repuestos r
            JOIN equipos e ON r.id_equipo = e.id
        """
    
        params = []
        conditions = []
    
        if año:
            conditions.append("strftime('%Y', e.fecha_ingreso) = ?")
            params.append(str(año))
        if mes:
            conditions.append("strftime('%m', e.fecha_ingreso) = ?")
            params.append(f"{mes:02d}")
    
        if conditions:
            query += " WHERE " + " AND ".join(conditions)
    
        query += " GROUP BY r.nombre, mes, año ORDER BY año, mes, cantidad_total DESC"
    
        self.cursor.execute(query, params)
        return self.cursor.fetchall()  
    

    def actualizar_equipo(self, equipo_id, datos, usuario_id):
        """Actualiza un equipo existente en el sistema"""
        try:
            query = '''UPDATE equipos SET
                pj = ?, id_tipo_equipo = ?, serie = ?, id_marca = ?, id_modelo = ?,
                ubicacion = ?, fecha_ingreso = ?, falla = ?, observaciones = ?
                WHERE id = ?'''
            
            self.cursor.execute(query, (*datos, equipo_id))
            
            self.registrar_auditoria(
                usuario_id, 'MODIFICAR_EQUIPO', 'equipos', equipo_id,
                f"Equipo modificado: {datos[2]} (PJ: {datos[0]})")
            
            self.conn.commit()
            return True
        except sqlite3.IntegrityError as e:
            if "serie" in str(e):
                raise ValueError("El número de serie ya existe en otro equipo")
            raise
        except sqlite3.Error as e:
            self.conn.rollback()
            logging.error(f"Error actualizando equipo: {e}")
            raise
    
    def obtener_equipos(self, filtro=None):
        """Obtiene todos los equipos con opción de filtrado"""
        query = """
            SELECT e.id, e.pj, te.nombre, ma.nombre, mo.nombre, e.ubicacion, 
                   e.fecha_ingreso, e.fecha_salida, e.estado
            FROM equipos e
            JOIN tipos_equipo te ON e.id_tipo_equipo = te.id
            JOIN marcas ma ON e.id_marca = ma.id
            JOIN modelos mo ON e.id_modelo = mo.id
        """
        if filtro:
            query += f" WHERE {filtro}"
        query += " ORDER BY e.id DESC"
        
        self.cursor.execute(query)
        return self.cursor.fetchall()
    
    def obtener_equipo_por_id(self, equipo_id):
        """Obtiene un equipo específico por su ID"""
        query = """
            SELECT e.*, te.nombre as tipo_equipo, ma.nombre as marca, mo.nombre as modelo
            FROM equipos e
            JOIN tipos_equipo te ON e.id_tipo_equipo = te.id
            JOIN marcas ma ON e.id_marca = ma.id
            JOIN modelos mo ON e.id_modelo = mo.id
            WHERE e.id = ?
        """
        self.cursor.execute(query, (equipo_id,))
        result = self.cursor.fetchone()
        return dict(zip([column[0] for column in self.cursor.description], result)) if result else None
    
    def obtener_repuestos_por_equipo(self, equipo_id):
        """Obtiene los repuestos utilizados en un equipo"""
        self.cursor.execute(
            "SELECT id, nombre, cantidad, costo FROM repuestos WHERE id_equipo=?",
            (equipo_id,))
        return self.cursor.fetchall()
    
    def obtener_repuestos_por_periodo(self, fecha_inicio, fecha_fin):
        """Obtiene repuestos utilizados en un período específico"""
        query = """
            SELECT r.nombre, SUM(r.cantidad) as cantidad_total, 
                   SUM(r.costo * r.cantidad) as costo_total, 
                   e.serie, te.nombre as tipo_equipo
            FROM repuestos r
            JOIN equipos e ON r.id_equipo = e.id
            JOIN tipos_equipo te ON e.id_tipo_equipo = te.id
            WHERE e.fecha_ingreso BETWEEN ? AND ?
            GROUP BY r.nombre, e.serie, te.nombre
            ORDER BY cantidad_total DESC
        """
        self.cursor.execute(query, (fecha_inicio, fecha_fin))
        return self.cursor.fetchall()
    
    def agregar_repuesto(self, equipo_id, nombre, cantidad, costo, usuario_id):
        """Agrega un repuesto a un equipo"""
        try:
            # Verificar si el equipo ya está reparado
            equipo = self.obtener_equipo_por_id(equipo_id)
            if equipo and equipo['estado'] == "Reparado":
                raise ValueError("No se pueden agregar repuestos a equipos ya reparados")
                
            self.cursor.execute(
                "INSERT INTO repuestos (id_equipo, nombre, cantidad, costo) VALUES (?, ?, ?, ?)",
                (equipo_id, nombre, cantidad, costo))
            
            repuesto_id = self.cursor.lastrowid
            
            self.registrar_auditoria(
                usuario_id, 'AGREGAR_REPUESTO', 'repuestos', repuesto_id,
                f"Repuesto agregado: {nombre} (Cant: {cantidad}, Costo: {costo})")
            
            self.conn.commit()
            return repuesto_id
        except sqlite3.Error as e:
            self.conn.rollback()
            logging.error(f"Error agregando repuesto: {e}")
            raise
    
    def eliminar_repuesto(self, repuesto_id, usuario_id):
        """Elimina un repuesto del sistema"""
        try:
            # Obtener datos del repuesto antes de eliminarlo para auditoría
            self.cursor.execute(
                "SELECT nombre, cantidad, costo FROM repuestos WHERE id = ?",
                (repuesto_id,))
            repuesto = self.cursor.fetchone()
            
            if not repuesto:
                raise ValueError("Repuesto no encontrado")
            
            # Verificar si el equipo ya está reparado
            self.cursor.execute(
                "SELECT id_equipo FROM repuestos WHERE id = ?",
                (repuesto_id,))
            equipo_id = self.cursor.fetchone()[0]
            
            equipo = self.obtener_equipo_por_id(equipo_id)
            if equipo and equipo['estado'] == "Reparado":
                raise ValueError("No se pueden modificar repuestos de equipos ya reparados")
            
            self.cursor.execute(
                "DELETE FROM repuestos WHERE id = ?",
                (repuesto_id,))
            
            self.registrar_auditoria(
                usuario_id, 'ELIMINAR_REPUESTO', 'repuestos', repuesto_id,
                f"Repuesto eliminado: {repuesto[0]} (Cant: {repuesto[1]}, Costo: {repuesto[2]})")
            
            self.conn.commit()
        except sqlite3.Error as e:
            self.conn.rollback()
            logging.error(f"Error eliminando repuesto: {e}")
            raise
    
    def actualizar_estado_equipo(self, equipo_id, estado, observaciones, usuario_id):
        """Actualiza el estado de un equipo"""
        try:
            # Verificar si ya está en el estado solicitado
            equipo = self.obtener_equipo_por_id(equipo_id)
            if equipo and equipo['estado'] == estado:
                return  # No hacer nada si el estado es el mismo
            
            query = """UPDATE equipos 
                      SET estado = ?, 
                          observaciones = ?,
                          fecha_salida = ?
                      WHERE id = ?"""
            
            fecha_salida = datetime.now().strftime("%Y-%m-%d") if estado == "Reparado" else None
            
            self.cursor.execute(query, (
                estado,
                observaciones or None,
                fecha_salida,
                equipo_id
            ))
            
            self.registrar_auditoria(
                usuario_id, 'ACTUALIZAR_ESTADO', 'equipos', equipo_id,
                f"Estado actualizado a: {estado}")
            
            self.conn.commit()
        except sqlite3.Error as e:
            self.conn.rollback()
            logging.error(f"Error actualizando estado del equipo: {e}")
            raise
    
    def registrar_auditoria(self, usuario_id, accion, tabla=None, registro_id=None, detalles=None):
        """Registra una acción en el log de auditoría"""
        try:
            self.cursor.execute(
                """INSERT INTO auditoria 
                (usuario_id, accion, tabla_afectada, registro_id, fecha, detalles) 
                VALUES (?, ?, ?, ?, ?, ?)""",
                (usuario_id, accion, tabla, registro_id, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), detalles))
            self.conn.commit()
        except sqlite3.Error as e:
            logging.error(f"Error registrando auditoría: {e}")
            self.conn.rollback()
    
    def __del__(self):
        if hasattr(self, 'conn'):
            self.conn.close()



    def crear_tabla_toner():
        conexion = sqlite3.connect("BaseDatos.db")
        cursor = conexion.cursor()
    
        cursor.execute("""
        CREATE TABLE IF NOT EXISTS Toner (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            modelo TEXT NOT NULL UNIQUE,
            stock INTEGER DEFAULT 0,
            compatible_con TEXT,
            color TEXT,
            fecha_actualizacion TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """)
    
        conexion.commit()
        conexion.close()
        
    # Métodos para marcas de toner
    def obtener_marcas_toner(self):
        self.cursor.execute("SELECT id, nombre FROM marcas_toner ORDER BY nombre")
        return self.cursor.fetchall()

    def agregar_marca_toner(self, nombre, usuario_id):
        try:
            self.cursor.execute("INSERT INTO marcas_toner (nombre) VALUES (?)", (nombre,))
            marca_id = self.cursor.lastrowid
            self.registrar_auditoria(usuario_id, 'ALTA_MARCA_TONER', 'marcas_toner', marca_id, f"Nueva marca de toner: {nombre}")
            self.conn.commit()
            return marca_id
        except sqlite3.IntegrityError:
            raise ValueError("La marca ya existe")

    # Métodos para modelos de toner
    def obtener_modelos_toner(self, id_marca):
        self.cursor.execute(
            "SELECT id, nombre FROM modelos_toner WHERE id_marca=? ORDER BY nombre", 
            (id_marca,))
        return self.cursor.fetchall()

    def agregar_modelo_toner(self, id_marca, nombre, usuario_id):
        try:
            self.cursor.execute("INSERT INTO modelos_toner (id_marca, nombre) VALUES (?, ?)", (id_marca, nombre))
            modelo_id = self.cursor.lastrowid
            self.registrar_auditoria(usuario_id, 'ALTA_MODELO_TONER', 'modelos_toner', modelo_id, f"Nuevo modelo de toner: {nombre}")
            self.conn.commit()
            return modelo_id
        except sqlite3.IntegrityError:
            raise ValueError("El modelo ya existe para esta marca")

    # Métodos para stock de toner
    def obtener_stock_toner(self):
        query = """
            SELECT m.nombre as marca, mo.nombre as modelo, s.cantidad, s.id
            FROM stock_toner s
            JOIN modelos_toner mo ON s.id_modelo = mo.id
            JOIN marcas_toner m ON mo.id_marca = m.id
            ORDER BY m.nombre, mo.nombre
        """
        self.cursor.execute(query)
        return self.cursor.fetchall()

    def actualizar_stock_toner(self, id_modelo, cantidad, usuario_id):
        try:
            # Verificar si ya existe registro para este modelo
            self.cursor.execute("SELECT id FROM stock_toner WHERE id_modelo=?", (id_modelo,))
            exists = self.cursor.fetchone()
        
            if exists:
                self.cursor.execute("UPDATE stock_toner SET cantidad=? WHERE id_modelo=?", (cantidad, id_modelo))
            else:
                self.cursor.execute("INSERT INTO stock_toner (id_modelo, cantidad) VALUES (?, ?)", (id_modelo, cantidad))
        
            self.registrar_auditoria(usuario_id, 'ACTUALIZAR_STOCK_TONER', 'stock_toner', id_modelo, f"Stock actualizado a {cantidad}")
            self.conn.commit()
        except sqlite3.Error as e:
            self.conn.rollback()
            raise

    # Métodos para movimientos de toner
    def registrar_movimiento_toner(self, id_modelo, tipo, cantidad, responsable, sector, empresa_recarga, observaciones, usuario_id):
        try:
            query = """
                INSERT INTO movimientos_toner 
                (id_modelo, tipo, cantidad, responsable, sector, empresa_recarga, observaciones, fecha, usuario_id)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """
            fecha = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            self.cursor.execute(query, (id_modelo, tipo, cantidad, responsable, sector, empresa_recarga, observaciones, fecha, usuario_id))
        
            # Actualizar stock según el tipo de movimiento
            if tipo == 'retiro':
                self.cursor.execute("UPDATE stock_toner SET cantidad = cantidad - ? WHERE id_modelo=?", (cantidad, id_modelo))
            elif tipo == 'ingreso':
                self.cursor.execute("UPDATE stock_toner SET cantidad = cantidad + ? WHERE id_modelo=?", (cantidad, id_modelo))
        
            movimiento_id = self.cursor.lastrowid
            self.registrar_auditoria(usuario_id, 'MOVIMIENTO_TONER', 'movimientos_toner', movimiento_id, 
                                   f"Movimiento de toner: {tipo} - Cantidad: {cantidad}")
            self.conn.commit()
            return movimiento_id
        except sqlite3.Error as e:
            self.conn.rollback()
            raise

    # Métodos para recargas de toner
    def registrar_recarga_toner(self, id_modelo, cantidad, empresa, observaciones, usuario_id):
        try:
            query = """
                INSERT INTO recargas_toner 
                (id_modelo, cantidad, empresa, fecha_envio, estado, observaciones, usuario_envio_id)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """
            fecha = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            self.cursor.execute(query, (id_modelo, cantidad, empresa, fecha, 'Enviado', observaciones, usuario_id))
        
            # Actualizar stock (restar los enviados a recarga)
            self.cursor.execute("UPDATE stock_toner SET cantidad = cantidad - ? WHERE id_modelo=?", (cantidad, id_modelo))
        
            recarga_id = self.cursor.lastrowid
            self.registrar_auditoria(usuario_id, 'ENVIO_RECARGA_TONER', 'recargas_toner', recarga_id, 
                                   f"Envío a recarga: {cantidad} unidades")
            self.conn.commit()
            return recarga_id
        except sqlite3.Error as e:
            self.conn.rollback()
            raise

    def recibir_recarga_toner(self, recarga_id, observaciones, usuario_id):
        try:
            # Obtener datos de la recarga
            self.cursor.execute("SELECT id_modelo, cantidad FROM recargas_toner WHERE id=?", (recarga_id,))
            recarga = self.cursor.fetchone()
            if not recarga:
                raise ValueError("Recarga no encontrada")
        
            id_modelo, cantidad = recarga
        
            # Actualizar recarga
            fecha = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            self.cursor.execute("""
                UPDATE recargas_toner 
                SET estado='Recibido', fecha_recibo=?, observaciones=?, usuario_recibo_id=?
                WHERE id=?
             """, (fecha, observaciones, usuario_id, recarga_id))
        
            # Actualizar stock (sumar los recibidos)
            self.cursor.execute("UPDATE stock_toner SET cantidad = cantidad + ? WHERE id_modelo=?", (cantidad, id_modelo))
        
            self.registrar_auditoria(usuario_id, 'RECIBO_RECARGA_TONER', 'recargas_toner', recarga_id, 
                                   f"Recepción de recarga: {cantidad} unidades")
            self.conn.commit()
        except sqlite3.Error as e:
            self.conn.rollback()
            raise

    # Métodos para informes
    def obtener_movimientos_toner(self, fecha_inicio, fecha_fin, id_marca=None, id_modelo=None):
        query = """
            SELECT m.nombre as marca, mo.nombre as modelo, mt.tipo, mt.cantidad, 
                   mt.responsable, mt.sector, mt.fecha, u.username
            FROM movimientos_toner mt
            JOIN modelos_toner mo ON mt.id_modelo = mo.id
            JOIN marcas_toner m ON mo.id_marca = m.id
            JOIN usuarios u ON mt.usuario_id = u.id
            WHERE mt.fecha BETWEEN ? AND ?
        """
        params = [fecha_inicio, fecha_fin]
    
        if id_marca:
            query += " AND mo.id_marca = ?"
            params.append(id_marca)
       
        if id_modelo:
            query += " AND mt.id_modelo = ?"
            params.append(id_modelo)
    
        query += " ORDER BY mt.fecha DESC"
    
        self.cursor.execute(query, params)
        return self.cursor.fetchall()

    def obtener_recargas_toner(self, año=None, mes=None, estado=None):
        query = """
            SELECT m.nombre as marca, mo.nombre as modelo, r.cantidad, r.empresa,
                   r.fecha_envio, r.fecha_recibo, r.estado, ue.username as usuario_envio,
                   ur.username as usuario_recibo
            FROM recargas_toner r
            JOIN modelos_toner mo ON r.id_modelo = mo.id
            JOIN marcas_toner m ON mo.id_marca = m.id
            JOIN usuarios ue ON r.usuario_envio_id = ue.id
            LEFT JOIN usuarios ur ON r.usuario_recibo_id = ur.id
        """
    
        conditions = []
        params = []
    
        if año:
            conditions.append("strftime('%Y', r.fecha_envio) = ?")
            params.append(str(año))
    
        if mes:
            conditions.append("strftime('%m', r.fecha_envio) = ?")
            params.append(f"{mes:02d}")
    
        if estado:
            conditions.append("r.estado = ?")
            params.append(estado)
    
        if conditions:
            query += " WHERE " + " AND ".join(conditions)
    
        query += " ORDER BY r.fecha_envio DESC"
    
        self.cursor.execute(query, params)
        return self.cursor.fetchall()





    def obtener_stock_toner_por_modelo(self, id_modelo):
        """Obtiene el stock disponible para un modelo específico"""
        self.cursor.execute(
            "SELECT cantidad FROM stock_toner WHERE id_modelo=?", 
            (id_modelo,))
        result = self.cursor.fetchone()
        return result[0] if result else 0

    def obtener_empresas_recarga(self):
        """Obtiene todas las empresas de recarga"""
        self.cursor.execute(
            "SELECT id, nombre, contacto, telefono FROM empresas_recarga ORDER BY nombre")
        return self.cursor.fetchall()

    def obtener_anios_movimientos_toner(self):
        """Obtiene los años distintos en los que hay movimientos"""
        self.cursor.execute(
            "SELECT DISTINCT strftime('%Y', fecha) FROM movimientos_toner ORDER BY fecha DESC")
        return [int(row[0]) for row in self.cursor.fetchall()]

    def obtener_movimientos_toner_para_informe(self, año, mes=None):
        """Obtiene movimientos para generar informes"""
        query = """
            SELECT m.fecha, ma.nombre, mo.nombre, mt.cantidad, mt.responsable, mt.sector
            FROM movimientos_toner mt
            JOIN modelos_toner mo ON mt.id_modelo = mo.id
            JOIN marcas_toner ma ON mo.id_marca = ma.id
            WHERE strftime('%Y', mt.fecha) = ?
        """
        params = [str(año)]
    
        if mes:
            query += " AND strftime('%m', mt.fecha) = ?"
            params.append(f"{mes:02d}")
    
        query += " ORDER BY mt.fecha DESC"
    
        self.cursor.execute(query, params)
        return self.cursor.fetchall()

    def obtener_recargas_toner_para_informe(self, año, mes=None):
        """Obtiene recargas para generar informes"""
        query = """
            SELECT r.fecha_envio, r.fecha_recibo, ma.nombre, mo.nombre, r.cantidad, e.nombre, r.estado
            FROM recargas_toner r
            JOIN modelos_toner mo ON r.id_modelo = mo.id
            JOIN marcas_toner ma ON mo.id_marca = ma.id
            JOIN empresas_recarga e ON r.id_empresa = e.id
            WHERE strftime('%Y', r.fecha_envio) = ?
        """
        params = [str(año)]
    
        if mes:
            query += " AND strftime('%m', r.fecha_envio) = ?"
            params.append(f"{mes:02d}")
    
        query += " ORDER BY r.fecha_envio DESC"
    
        self.cursor.execute(query, params)
        return self.cursor.fetchall()

          

class LoginView(ttk.Frame):
    """Vista para el sistema de login con diseño profesional"""
    
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self._setup_ui()
        self._configure_styles()
    
    def _configure_styles(self):
        style = ttk.Style()
        style.configure('Login.TFrame', 
                      background='white', 
                      borderwidth=2, 
                      relief="solid",
                      bordercolor=Config.BORDER_COLOR)
        style.configure('Login.TButton', 
                      font=('Helvetica', 10, 'bold'),
                      foreground='white', 
                      background=Config.SECONDARY_COLOR,
                      padding=10,
                      borderwidth=1,
                      relief="raised")
        style.map('Login.TButton',
                foreground=[('pressed', 'white'), ('active', 'white')],
                background=[('pressed', Config.PRIMARY_COLOR), ('active', '#2185d0')])
        style.configure('Login.TLabel', 
                      font=('Helvetica', 10), 
                      foreground=Config.TEXT_COLOR)
        style.configure('Login.TEntry',
                      fieldbackground='white',
                      foreground=Config.TEXT_COLOR,
                      bordercolor=Config.BORDER_COLOR,
                      lightcolor=Config.BORDER_COLOR,
                      darkcolor=Config.BORDER_COLOR)
    
    def _setup_ui(self):
        self.grid(row=0, column=0, sticky="nsew")
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(0, weight=1)
        self.configure(style='Login.TFrame')
        
        # Frame principal con sombra
        main_frame = ttk.Frame(self, padding=(40, 50, 40, 50), style='Login.TFrame')
        main_frame.grid(row=0, column=0, sticky="")
        main_frame.grid_columnconfigure(0, weight=1)
        
        # Logo o título
        logo_frame = ttk.Frame(main_frame)
        logo_frame.grid(row=0, column=0, pady=(0, 30))
        
        try:
            logo_img = Image.open("logo_pj.png")
            logo_img = logo_img.resize((120, 120), Image.Resampling.LANCZOS)
            self.logo = ImageTk.PhotoImage(logo_img)
            ttk.Label(logo_frame, image=self.logo).pack()
        except:
            ttk.Label(logo_frame, text="⚙️", font=('Helvetica', 48), foreground=Config.SECONDARY_COLOR).pack()
        
        ttk.Label(logo_frame, text="PODER JUDICIAL", 
                 font=('Helvetica', 16, 'bold'), 
                 foreground=Config.PRIMARY_COLOR).pack(pady=(10, 0))
        ttk.Label(logo_frame, text="Sistema de Gestión de Stock", 
                 font=('Helvetica', 12), 
                 foreground=Config.TEXT_COLOR).pack()
        
        # Campos de formulario
        form_frame = ttk.LabelFrame(main_frame, text="Ingreso al Sistema", padding="10")
        form_frame.grid(row=1, column=0, sticky="ew", pady=(20, 0))
        
        ttk.Label(form_frame, text="Usuario:", style='Login.TLabel').grid(row=0, column=0, sticky="w", pady=(0, 5))
        self.username_entry = ttk.Entry(form_frame, font=('Helvetica', 11))
        self.username_entry.grid(row=1, column=0, pady=(0, 15), sticky="ew", ipady=5)
        
        ttk.Label(form_frame, text="Contraseña:", style='Login.TLabel').grid(row=2, column=0, sticky="w", pady=(0, 5))
        self.password_entry = ttk.Entry(form_frame, show="•", font=('Helvetica', 11))
        self.password_entry.grid(row=3, column=0, pady=(0, 25), sticky="ew", ipady=5)
        
        btn_frame = ttk.Frame(form_frame)
        btn_frame.grid(row=4, column=0, sticky="ew", pady=(0, 10))
        btn_frame.grid_columnconfigure(0, weight=1)
        btn_frame.grid_columnconfigure(1, weight=1)
        
        ttk.Button(btn_frame, text="Ingresar", style='Login.TButton',
                  command=self._on_login).grid(row=0, column=0, padx=5, sticky="ew")
        ttk.Button(btn_frame, text="Cancelar", style='Login.TButton',
                  command=self.controller.cerrar_aplicacion).grid(row=0, column=1, padx=5, sticky="ew")
        
        self.password_entry.bind("<Return>", lambda e: self._on_login())
        self.username_entry.focus_set()
    
    def _on_login(self):
        username = self.username_entry.get().strip()
        password = self.password_entry.get().strip()
        
        if not username or not password:
            messagebox.showwarning("Advertencia", "Ingrese usuario y contraseña")
            return
        
        # Mostrar spinner de carga
        self.loading_label = ttk.Label(self, text="Autenticando...")
        self.loading_label.place(relx=0.5, rely=0.8, anchor="center")
        self.update()
        
        # Ejecutar autenticación en segundo plano
        threading.Thread(target=self._authenticate, args=(username, password), daemon=True).start()
    
    def _authenticate(self, username, password):
        """Método de autenticación en segundo plano"""
        try:
            result = self.controller.autenticar_usuario(username, password)
            self.after(0, self.loading_label.place_forget)
            
            if result:
                self.after(0, self.controller.mostrar_vista_principal, result)
            else:
                self.after(0, messagebox.showerror, "Error", "Credenciales inválidas")
        except Exception as e:
            self.after(0, messagebox.showerror, "Error", f"Error en autenticación: {str(e)}")


# ... (El resto del código anterior permanece igual hasta la clase MainView)


class MainView(ttk.Frame):
    """Vista principal de la aplicación"""

    def initialize(self, *args, **kwargs):
        # Verificar autenticación primero
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
    
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self.status_var = tk.StringVar()
        self._setup_ui()

    def _generar_informe_tecnico(self, equipo_id):
        """Genera un informe técnico para el equipo"""
        try:
            equipo = self.controller.db.obtener_equipo_por_id(equipo_id)
            if not equipo:
                messagebox.showerror("Error", "Equipo no encontrado")
                return
            
            repuestos = self.controller.db.obtener_repuestos_por_equipo(equipo_id)
            
            informe_data = {
                "titulo": f"Informe Técnico - Equipo {equipo['serie']}",
                "info_equipo": [
                    ["ID del Equipo", equipo['id']],
                    ["Número de Serie", equipo['serie']],
                    ["Tipo de Equipo", equipo['tipo_equipo']],
                    ["Marca", equipo['marca']],
                    ["Modelo", equipo['modelo']],
                    ["PJ", equipo['pj']],
                    ["Ubicación", equipo['ubicacion'] or "N/A"],
                    ["Fecha de Ingreso", equipo['fecha_ingreso']],
                    ["Fecha de Salida", equipo['fecha_salida'] or "N/A"],
                    ["Estado", equipo['estado']],
                    ["Falla Reportada", equipo['falla']]
                ],
                "repuestos": [["Repuesto", "Cantidad", "Costo Unitario", "Costo Total"]] + [
                    [rep[1], rep[2], f"${rep[3]:.2f}" if rep[3] else "N/A", f"${rep[2]*rep[3]:.2f}" if rep[3] else "N/A"]
                    for rep in repuestos
                ],
                "observaciones": equipo['observaciones'] or "Sin observaciones adicionales",
                "estado": equipo['estado']
            }
            
            fecha = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Informe_Tecnico_{equipo['serie']}_{fecha}.docx"
            
            if ExportManager.export_informe_tecnico(informe_data, filename):
                messagebox.showinfo("Éxito", f"Informe técnico generado:\n{filename}")
                if messagebox.askyesno("Abrir", "¿Desea abrir el documento para imprimir?"):
                    webbrowser.open(filename)
            else:
                messagebox.showerror("Error", "No se pudo generar el informe técnico")
                
        except Exception as e:
            logging.error(f"Error generando informe técnico: {e}")
            messagebox.showerror("Error", f"No se pudo generar el informe: {str(e)}")

    def _imprimir_informe_directo(self):
        """Opción directa para imprimir informe desde el menú"""
        equipo_id = simpledialog.askinteger("Imprimir Informe", "Ingrese el ID del equipo:")
        if equipo_id:
            self._generar_informe_tecnico(equipo_id)


        
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        
        # Configurar tema
        self.style = ttk.Style()
        self.style.theme_use(Config.UI_THEME)
        
        # Barra de herramientas superior
        toolbar = ttk.Frame(self)
        toolbar.pack(side=tk.TOP, fill=tk.X, padx=5, pady=5)
        
        # Botón Nuevo Equipo en la barra de herramientas
        ttk.Button(toolbar, text="Nuevo Equipo", 
                  command=lambda: self.controller.mostrar_vista("EquipmentView")).pack(side=tk.LEFT, padx=5)
        
        # Barra de estado
        status_bar = ttk.Label(self, textvariable=self.status_var, 
                             relief=tk.SUNKEN, anchor=tk.W,
                             style='Status.TLabel')
        status_bar.pack(side=tk.BOTTOM, fill=tk.X)
        
        # Barra de menú superior (menú tradicional)
        menubar = tk.Menu(self.controller.root)
        
        
        # Menú Archivo (estructura final modificada)
        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="Exportar a PDF", command=self._exportar_pdf)
        file_menu.add_command(label="Exportar a Excel", command=self._exportar_excel)
        file_menu.add_command(label="Exportar a Word", command=self._exportar_word)  # Opcional: si ya tienes esta función
        file_menu.add_separator()
        file_menu.add_command(label="Cerrar sesión", command=self._cerrar_sesion)  # ¡Nueva opción!
        file_menu.add_separator()
        file_menu.add_command(label="Salir", command=self.controller.cerrar_aplicacion)
        menubar.add_cascade(label="Archivo", menu=file_menu)
        
        
        # Menú Equipos (se puede mantener para otras opciones)
        equipos_menu = tk.Menu(menubar, tearoff=0)
        equipos_menu.add_command(label="Nuevo Equipo", 
                               command=lambda: self.controller.mostrar_vista("EquipmentView"))
        equipos_menu.add_command(label="Editar Equipo", 
                               command=self._editar_equipo)
        equipos_menu.add_command(label="Reparar Equipo", 
                               command=self._iniciar_reparacion)
        equipos_menu.add_command(label="Imprimir Informe", 
                               command=self._imprimir_informe_directo)  # Nueva opción
        
        equipos_menu.add_command(label="Listar Equipos", 
                               command=self._mostrar_lista_equipos)
        menubar.add_cascade(label="Equipos", menu=equipos_menu)
        


        # Resto del código de menús permanece igual...


        # Menú Reportes
        reportes_menu = tk.Menu(menubar, tearoff=0)
        reportes_menu.add_command(label="Reporte de Equipos", 
                                command=self._mostrar_reporte_equipos)
        reportes_menu.add_command(label="Reporte de Repuestos", 
                                command=self._mostrar_reporte_repuestos)
        menubar.add_cascade(label="Reportes", menu=reportes_menu)

        reportes_menu.add_command(label="Resumen de Repuestos", 
                          command=lambda: self.controller.mostrar_vista("ReporteResumenRepuestosView"))
        
        # Menú Configuración
        config_menu = tk.Menu(menubar, tearoff=0)
               
        config_menu.add_command(label="Gestionar Marcas", 
                              command=lambda: self.controller.mostrar_vista("MarcasView"))
        config_menu.add_command(label="Gestionar Modelos", 
                              command=lambda: self.controller.mostrar_vista("ModelosView"))
        menubar.add_cascade(label="Configuración", menu=config_menu)

        config_menu.add_command(label="Gestionar Usuarios", 
                      command=lambda: self.controller.mostrar_vista("UsuariosView"))




        # Menú Toner (nuevo menú)
        toner_menu = tk.Menu(menubar, tearoff=0)
        toner_menu.add_command(label="Gestión de Toner", 
                             command=lambda: self.controller.mostrar_vista("TonerView"))
        toner_menu.add_command(label="Marcas de Toner", 
                             command=lambda: self.controller.mostrar_vista("MarcasTonerView"))
        toner_menu.add_command(label="Modelos de Toner", 
                             command=lambda: self.controller.mostrar_vista("ModelosTonerView"))
        menubar.add_cascade(label="Toner", menu=toner_menu)

                


        # MODIFICAR ESTA PARTE - Añadir condición para mostrar opción de usuarios  
        if self.controller.current_user and self.controller.current_user['rol'] == 'admin':
            config_menu.add_command(label="Gestionar Usuarios", 
                                  command=lambda: self.controller.mostrar_vista("UsuariosView"))
        
             
        self.controller.root.config(menu=menubar)
        
        # Panel principal con pestañas
        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Pestaña de Equipos
        self.equipos_frame = ttk.Frame(self.notebook)
        self._setup_equipos_ui()
        self.notebook.add(self.equipos_frame, text="Gestión de Equipos")
        
        # Pestaña de Reportes
        self.reportes_frame = ttk.Frame(self.notebook)
        self._setup_reportes_ui()
        self.notebook.add(self.reportes_frame, text="Reportes")
        
        self.actualizar_status("Sistema listo")

    # ... (El resto de los métodos de MainView permanecen igual)

# ... (El resto del código posterior permanece igual)




    def _setup_equipos_ui(self):
        """Configura la interfaz de la pestaña de equipos"""
        self.equipos_treeview = ttk.Treeview(self.equipos_frame, columns=(
            "ID", "PJ", "Tipo", "Marca", "Modelo", "Ubicación", "Ingreso", "Salida", "Estado"
        ), show="headings")
        
        # Configurar columnas
        columns = {
            "ID": {"width": 50, "anchor": tk.CENTER},
            "PJ": {"width": 100, "anchor": tk.W},
            "Tipo": {"width": 100, "anchor": tk.W},
            "Marca": {"width": 100, "anchor": tk.W},
            "Modelo": {"width": 120, "anchor": tk.W},
            "Ubicación": {"width": 120, "anchor": tk.W},
            "Ingreso": {"width": 100, "anchor": tk.CENTER},
            "Salida": {"width": 100, "anchor": tk.CENTER},
            "Estado": {"width": 120, "anchor": tk.W}
        }
        
        for col, config in columns.items():
            self.equipos_treeview.heading(col, text=col)
            self.equipos_treeview.column(col, **config)
        
        # Scrollbar
        scrollbar = ttk.Scrollbar(self.equipos_frame, orient="vertical", 
                                command=self.equipos_treeview.yview)
        self.equipos_treeview.configure(yscrollcommand=scrollbar.set)
        
        # Posicionamiento
        self.equipos_treeview.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Cargar datos iniciales
        self._cargar_equipos()
    
    def _setup_reportes_ui(self):
        """Configura la interfaz de la pestaña de reportes"""
        label = ttk.Label(self.reportes_frame, text="Panel de Reportes - Seleccione una opción del menú")
        label.pack(pady=50)
    
    def _cargar_equipos(self, filtro=None):
        """Carga los equipos desde la base de datos"""
        try:
            # Limpiar treeview
            for item in self.equipos_treeview.get_children():
                self.equipos_treeview.delete(item)
            
            # Obtener datos
            equipos = self.controller.db.obtener_equipos(filtro)
            
            # Insertar datos
            for equipo in equipos:
                self.equipos_treeview.insert("", tk.END, values=equipo)
            
            self.actualizar_status(f"{len(equipos)} equipos cargados")
        except Exception as e:
            logging.error(f"Error cargando equipos: {e}")
            messagebox.showerror("Error", f"No se pudieron cargar los equipos: {str(e)}")
            self.actualizar_status("Error cargando equipos")
    
    def _editar_equipo(self):
        """Inicia el proceso de edición de un equipo"""
        equipo_id = simpledialog.askinteger("Editar Equipo", "Ingrese el ID del equipo a editar:")
    
        if equipo_id:
            try:
                equipo = self.controller.db.obtener_equipo_por_id(equipo_id)
                if not equipo:
                    messagebox.showerror("Error", "Equipo no encontrado")
                    return
            
                # Verificar si ya está reparado y confirmar edición
                if equipo['estado'] == "Reparado":
                    messagebox.showwarning("Advertencia", 
                         "No se recomienda editar equipos ya reparados.\n"
                         "¿Desea continuar?", icon='warning')
                    if not messagebox.askyesno("Confirmar", "Continuar con la edición?"):
                        return
            
                # Mostrar directamente la vista de edición
                self.controller.mostrar_vista("EditarEquipoView", equipo_id)
                
            except Exception as e:
                logging.error(f"Error iniciando edición: {e}")
                messagebox.showerror("Error", f"No se pudo iniciar la edición:\n{str(e)}")
    
    def _iniciar_reparacion(self):
        """Inicia el proceso de reparación de un equipo"""
        equipo_id = simpledialog.askinteger("Reparar Equipo", "Ingrese el ID del equipo a reparar:")
        
        if equipo_id:
            try:
                # Verificar que el equipo existe
                equipo = self.controller.db.obtener_equipo_por_id(equipo_id)
                if not equipo:
                    messagebox.showerror("Error", "Equipo no encontrado")
                    return
                
                # Verificar que no esté ya reparado
                if equipo['estado'] == "Reparado":
                    messagebox.showwarning("Advertencia", 
                        "Este equipo ya está marcado como Reparado.\n"
                        "No se puede volver a reparar.")
                    return
                    
                # Mostrar vista de reparación
                self.controller.mostrar_vista("ReparacionView", equipo_id)
                
            except Exception as e:
                logging.error(f"Error iniciando reparación: {e}")
                messagebox.showerror("Error", f"No se pudo iniciar la reparación:\n{str(e)}")
    
    def _mostrar_lista_equipos(self):
        """Muestra la lista de equipos actualizada"""
        self._cargar_equipos()
        self.notebook.select(self.equipos_frame)
    
    def _mostrar_reporte_equipos(self):
        """Muestra el reporte de equipos"""
        self.controller.mostrar_vista("ReporteEquiposView")
        self.notebook.select(self.reportes_frame)
    
    def _mostrar_reporte_repuestos(self):
        """Muestra el reporte de repuestos"""
        self.controller.mostrar_vista("ReporteRepuestosView")
        self.notebook.select(self.reportes_frame)
    
    def _exportar_pdf(self):
        """Exporta los datos actuales a PDF"""
        if not self.controller.current_user:
            messagebox.showwarning("Acceso denegado", "Debe iniciar sesión para exportar")
            self.controller.mostrar_vista("LoginView")
            return
     

        
        try:
            current_tab = self.notebook.tab(self.notebook.select(), "text")
            
            if current_tab == "Gestión de Equipos":
                data = []
                headers = [self.equipos_treeview.heading(col)["text"] 
                          for col in self.equipos_treeview["columns"]]
                data.append(headers)
                
                for item in self.equipos_treeview.get_children():
                    data.append(self.equipos_treeview.item(item)["values"])
                
                title = "Reporte de Equipos"
            else:
                messagebox.showwarning("Advertencia", "No hay datos para exportar en esta pestaña")
                return
            
            filename = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf")],
                title="Guardar reporte como PDF"
            )
            
            if filename:
                if ExportManager.export_to_pdf(data, filename, title):
                    messagebox.showinfo("Éxito", f"Reporte exportado a PDF:\n{filename}")
                    if messagebox.askyesno("Abrir", "¿Desea abrir el archivo PDF?"):
                        webbrowser.open(filename)
                else:
                    messagebox.showerror("Error", "No se pudo exportar el PDF")
                    
        except Exception as e:
            logging.error(f"Error exportando a PDF: {e}")
            messagebox.showerror("Error", f"Error al exportar PDF: {str(e)}")
    
    def _exportar_excel(self):
        """Exporta los datos actuales a Excel"""


        if not self.controller.current_user:
            messagebox.showwarning("Acceso denegado", "Debe iniciar sesión para exportar")
            self.controller.mostrar_vista("LoginView")
            return

        try:
            current_tab = self.notebook.tab(self.notebook.select(), "text")
            
            if current_tab == "Gestión de Equipos":
                data = []
                headers = [self.equipos_treeview.heading(col)["text"] 
                          for col in self.equipos_treeview["columns"]]
                data.append(headers)
                
                for item in self.equipos_treeview.get_children():
                    data.append(self.equipos_treeview.item(item)["values"])
                
                sheet_name = "Equipos"
            else:
                messagebox.showwarning("Advertencia", "No hay datos para exportar en esta pestaña")
                return
            
            filename = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx")],
                title="Guardar reporte como Excel"
            )
            
            if filename:
                if ExportManager.export_to_excel(data, filename, sheet_name):
                    messagebox.showinfo("Éxito", f"Reporte exportado a Excel:\n{filename}")
                else:
                    messagebox.showerror("Error", "No se pudo exportar el Excel")
                    
        except Exception as e:
            logging.error(f"Error exportando a Excel: {e}")
            messagebox.showerror("Error", f"Error al exportar Excel: {str(e)}")
    
    def _exportar_word(self):
        """Exporta los datos actuales a Word"""


        if not self.controller.current_user:
            messagebox.showwarning("Acceso denegado", "Debe iniciar sesión para exportar")
            self.controller.mostrar_vista("LoginView")
            return

        try:
            current_tab = self.notebook.tab(self.notebook.select(), "text")
            
            if current_tab == "Gestión de Equipos":
                data = []
                headers = [self.equipos_treeview.heading(col)["text"] 
                          for col in self.equipos_treeview["columns"]]
                data.append(headers)
                
                for item in self.equipos_treeview.get_children():
                    data.append(self.equipos_treeview.item(item)["values"])
                
                title = "Reporte de Equipos"
            else:
                messagebox.showwarning("Advertencia", "No hay datos para exportar en esta pestaña")
                return
            
            filename = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx")],
                title="Guardar reporte como Word"
            )
            
            if filename:
                if ExportManager.export_to_word(data, filename, title):
                    messagebox.showinfo("Éxito", f"Reporte exportado a Word:\n{filename}")
                    if messagebox.askyesno("Abrir", "¿Desea abrir el archivo Word?"):
                        webbrowser.open(filename)
                else:
                    messagebox.showerror("Error", "No se pudo exportar el documento Word")
                    
        except Exception as e:
            logging.error(f"Error exportando a Word: {e}")
            messagebox.showerror("Error", f"Error al exportar Word: {str(e)}")
    
    def actualizar_status(self, mensaje):
        """Actualiza la barra de estado"""
        self.status_var.set(mensaje)


        # --- MÉTODO NUEVO QUE DEBES AGREGAR ---
    def _cerrar_sesion(self):

        """Cierra sesión de forma segura"""
        # Confirmar cierre de sesión
        if messagebox.askyesno("Cerrar sesión", "¿Está seguro que desea cerrar la sesión?"):
            # 1. Limpiar datos de usuario
            self.controller.current_user = None
        
            # 2. Obtener la vista de login
            login_view = self.controller.views["LoginView"]
        
            # 3. Limpiar campos
            login_view.username_entry.delete(0, tk.END)
            login_view.password_entry.delete(0, tk.END)
        
            # 4. Mostrar vista de login
            self.controller.mostrar_vista("LoginView")
        
            # 5. Enfocar campo de usuario
            login_view.username_entry.focus_set()
         



class EquipmentView(ttk.Frame):
    """Vista solo para ingreso de nuevos equipos"""


    def initialize(self, *args, **kwargs):
        # Verificar autenticación primero
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
    




    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self._setup_ui()
    
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(0, weight=1)
        
        main_frame = ttk.Frame(self, padding="20")
        main_frame.grid(row=0, column=0, sticky="nsew")
        main_frame.grid_columnconfigure(0, weight=1)
        
        form_frame = ttk.LabelFrame(main_frame, text="Ingreso de Nuevo Equipo", padding="10")
        form_frame.grid(row=0, column=0, sticky="nsew", pady=(0, 20))
        form_frame.grid_columnconfigure(1, weight=1)
        
        # Campos del formulario (sin sección de repuestos)
        ttk.Label(form_frame, text="PJ:").grid(row=0, column=0, sticky="w", pady=5)
        self.pj_entry = ttk.Entry(form_frame)
        self.pj_entry.grid(row=0, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(form_frame, text="Tipo de Equipo:").grid(row=1, column=0, sticky="w", pady=5)
        self.tipo_combobox = ttk.Combobox(form_frame, state="readonly")
        self.tipo_combobox.grid(row=1, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(form_frame, text="Marca:").grid(row=2, column=0, sticky="w", pady=5)
        self.marca_combobox = ttk.Combobox(form_frame, state="readonly")
        self.marca_combobox.grid(row=2, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(form_frame, text="Modelo:").grid(row=3, column=0, sticky="w", pady=5)
        self.modelo_combobox = ttk.Combobox(form_frame, state="readonly")
        self.modelo_combobox.grid(row=3, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(form_frame, text="Número de Serie:").grid(row=4, column=0, sticky="w", pady=5)
        self.serie_entry = ttk.Entry(form_frame)
        self.serie_entry.grid(row=4, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(form_frame, text="Ubicación:").grid(row=5, column=0, sticky="w", pady=5)
        self.ubicacion_entry = ttk.Entry(form_frame)
        self.ubicacion_entry.grid(row=5, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(form_frame, text="Fecha de Ingreso:").grid(row=6, column=0, sticky="w", pady=5)
        self.fecha_ingreso_entry = DateEntry(form_frame, date_pattern="yyyy-mm-dd")
        self.fecha_ingreso_entry.grid(row=6, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(form_frame, text="Falla:").grid(row=7, column=0, sticky="w", pady=5)
        self.falla_entry = ttk.Entry(form_frame)
        self.falla_entry.grid(row=7, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(form_frame, text="Observaciones:").grid(row=8, column=0, sticky="nw", pady=5)
        self.observaciones_text = tk.Text(form_frame, height=5, width=40)
        self.observaciones_text.grid(row=8, column=1, sticky="ew", padx=5, pady=5)
        
        # Botones
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=1, column=0, sticky="e")
        
        ttk.Button(button_frame, text="Guardar", 
                  command=self._guardar_equipo).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cancelar", 
                  command=lambda: self.controller.mostrar_vista("MainView")).pack(side=tk.LEFT, padx=5)
        
        # Cargar comboboxes
        self._cargar_comboboxes()
    
    def _cargar_comboboxes(self):
        """Carga los comboboxes con datos iniciales"""
        try:
            # Tipos de equipo
            tipos = self.controller.db.obtener_tipos_equipo()
            self.tipo_combobox["values"] = [t[1] for t in tipos]
            if tipos:
                self.tipo_combobox.current(0)
            
            # Marcas
            marcas = self.controller.db.obtener_marcas()
            self.marca_combobox["values"] = [m[1] for m in marcas]
            if marcas:
                self.marca_combobox.current(0)
                self._actualizar_modelos()
            
            # Configurar eventos
            self.marca_combobox.bind("<<ComboboxSelected>>", 
                                   lambda e: self._actualizar_modelos())
            
            # Fecha actual por defecto
            self.fecha_ingreso_entry.set_date(datetime.now().date())
            
        except Exception as e:
            logging.error(f"Error cargando comboboxes: {e}")
            messagebox.showerror("Error", f"No se pudieron cargar los datos: {str(e)}")
    
    def _limpiar_campos(self):
        """Limpia todos los campos del formulario"""
        self.pj_entry.delete(0, tk.END)
        self.tipo_combobox.set('')
        self.marca_combobox.set('')
        self.modelo_combobox.set('')
        self.serie_entry.delete(0, tk.END)
        self.ubicacion_entry.delete(0, tk.END)
        self.falla_entry.delete(0, tk.END)
        self.observaciones_text.delete("1.0", tk.END)
        self.fecha_ingreso_entry.set_date(datetime.now().date())
    
    def _actualizar_modelos(self):
        """Actualiza los modelos según la marca seleccionada"""
        try:
            marca_seleccionada = self.marca_combobox.get()
            if marca_seleccionada:
                marcas = self.controller.db.obtener_marcas()
                id_marca = next(m[0] for m in marcas if m[1] == marca_seleccionada)
                
                modelos = self.controller.db.obtener_modelos(id_marca)
                self.modelo_combobox["values"] = [m[1] for m in modelos]
                if modelos:
                    self.modelo_combobox.current(0)
        except Exception as e:
            logging.error(f"Error actualizando modelos: {e}")
            messagebox.showerror("Error", f"No se pudieron cargar los modelos: {str(e)}")
    
    def _guardar_equipo(self):
        """Guarda el nuevo equipo en la base de datos"""
        try:
            # Validar campos obligatorios
            campos_obligatorios = {
                "PJ": self.pj_entry.get(),
                "Tipo de equipo": self.tipo_combobox.get(),
                "Marca": self.marca_combobox.get(),
                "Modelo": self.modelo_combobox.get(),
                "Número de serie": self.serie_entry.get(),
                "Falla": self.falla_entry.get()
            }
            
            faltantes = [campo for campo, valor in campos_obligatorios.items() if not valor]
            if faltantes:
                raise ValueError(f"Complete los campos obligatorios:\n- " + "\n- ".join(faltantes))
            
            # Validar número de serie único
            serie = self.serie_entry.get().strip()
            equipos = self.controller.db.obtener_equipos(f"e.serie = '{serie}'")
            if equipos:
                raise ValueError(f"El número de serie '{serie}' ya existe en el sistema")
            
            # Obtener IDs de las selecciones
            tipo_id = next(t[0] for t in self.controller.db.obtener_tipos_equipo() 
                         if t[1] == self.tipo_combobox.get())
            marca_id = next(m[0] for m in self.controller.db.obtener_marcas() 
                          if m[1] == self.marca_combobox.get())
            modelo_id = next(m[0] for m in self.controller.db.obtener_modelos(marca_id) 
                         if m[1] == self.modelo_combobox.get())
            
            # Preparar datos
            datos = (
                self.pj_entry.get().strip(),
                tipo_id,
                self.serie_entry.get().strip(),
                marca_id,
                modelo_id,
                self.ubicacion_entry.get().strip() or None,
                self.fecha_ingreso_entry.get_date().strftime("%Y-%m-%d"),
                None,  # fecha_salida
                self.falla_entry.get().strip(),
                "En reparación",  # estado
                self.observaciones_text.get("1.0", tk.END).strip() or None
            )
            
            # Guardar en base de datos
            equipo_id = self.controller.db.agregar_equipo(
                datos, self.controller.current_user['id'])
            
            messagebox.showinfo("Éxito", f"Equipo registrado con ID: {equipo_id}")
            
            # Limpiar campos para nuevo ingreso
            self._limpiar_campos()
            
        except ValueError as e:
            messagebox.showerror("Error", str(e))
        except Exception as e:
            logging.error(f"Error guardando equipo: {e}")
            messagebox.showerror("Error", f"No se pudo guardar el equipo:\n{str(e)}")

class EditarEquipoView(ttk.Frame):
    """Vista para editar equipos existentes"""

    def initialize(self, *args, **kwargs):
        # Verificar autenticación primero
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
    
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self.equipo_id = None
        self._setup_ui()
    
    def initialize(self, equipo_id):
        """Inicializa la vista con el ID del equipo a editar"""
        self.equipo_id = equipo_id
        self._cargar_datos_equipo()
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(0, weight=1)
        
        main_frame = ttk.Frame(self, padding="20")
        main_frame.grid(row=0, column=0, sticky="nsew")
        main_frame.grid_columnconfigure(0, weight=1)
        
        form_frame = ttk.LabelFrame(main_frame, text="Editar Equipo", padding="10")
        form_frame.grid(row=0, column=0, sticky="nsew", pady=(0, 20))
        form_frame.grid_columnconfigure(1, weight=1)
        
        # Campos del formulario
        ttk.Label(form_frame, text="PJ:").grid(row=0, column=0, sticky="w", pady=5)
        self.pj_entry = ttk.Entry(form_frame)
        self.pj_entry.grid(row=0, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(form_frame, text="Tipo de Equipo:").grid(row=1, column=0, sticky="w", pady=5)
        self.tipo_combobox = ttk.Combobox(form_frame, state="readonly")
        self.tipo_combobox.grid(row=1, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(form_frame, text="Marca:").grid(row=2, column=0, sticky="w", pady=5)
        self.marca_combobox = ttk.Combobox(form_frame, state="readonly")
        self.marca_combobox.grid(row=2, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(form_frame, text="Modelo:").grid(row=3, column=0, sticky="w", pady=5)
        self.modelo_combobox = ttk.Combobox(form_frame, state="readonly")
        self.modelo_combobox.grid(row=3, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(form_frame, text="Número de Serie:").grid(row=4, column=0, sticky="w", pady=5)
        self.serie_entry = ttk.Entry(form_frame)
        self.serie_entry.grid(row=4, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(form_frame, text="Ubicación:").grid(row=5, column=0, sticky="w", pady=5)
        self.ubicacion_entry = ttk.Entry(form_frame)
        self.ubicacion_entry.grid(row=5, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(form_frame, text="Fecha de Ingreso:").grid(row=6, column=0, sticky="w", pady=5)
        self.fecha_ingreso_entry = DateEntry(form_frame, date_pattern="yyyy-mm-dd")
        self.fecha_ingreso_entry.grid(row=6, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(form_frame, text="Falla:").grid(row=7, column=0, sticky="w", pady=5)
        self.falla_entry = ttk.Entry(form_frame)
        self.falla_entry.grid(row=7, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(form_frame, text="Observaciones:").grid(row=8, column=0, sticky="nw", pady=5)
        self.observaciones_text = tk.Text(form_frame, height=5, width=40)
        self.observaciones_text.grid(row=8, column=1, sticky="ew", padx=5, pady=5)
        
        # Botones
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=1, column=0, sticky="e")
        
        ttk.Button(button_frame, text="Actualizar", 
                  command=self._actualizar_equipo).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cancelar", 
                  command=lambda: self.controller.mostrar_vista("MainView")).pack(side=tk.LEFT, padx=5)
        
        # Configurar eventos
        self.marca_combobox.bind("<<ComboboxSelected>>", 
                               lambda e: self._actualizar_modelos())
    
    def _cargar_datos_equipo(self):
        """Carga los datos del equipo a editar"""
        equipo = self.controller.db.obtener_equipo_por_id(self.equipo_id)
        if equipo:
            # Limpiar campos primero
            self.pj_entry.delete(0, tk.END)
            self.tipo_combobox.set('')
            self.marca_combobox.set('')
            self.modelo_combobox.set('')
            self.serie_entry.delete(0, tk.END)
            self.ubicacion_entry.delete(0, tk.END)
            self.falla_entry.delete(0, tk.END)
            self.observaciones_text.delete("1.0", tk.END)
            
            # Cargar datos del equipo
            self.pj_entry.insert(0, equipo['pj'])
            
            # Cargar tipos de equipo
            tipos = [t[1] for t in self.controller.db.obtener_tipos_equipo()]
            self.tipo_combobox['values'] = tipos
            if equipo['tipo_equipo'] in tipos:
                self.tipo_combobox.set(equipo['tipo_equipo'])
            
            # Cargar marcas
            marcas = [m[1] for m in self.controller.db.obtener_marcas()]
            self.marca_combobox['values'] = marcas
            if equipo['marca'] in marcas:
                self.marca_combobox.set(equipo['marca'])
                # Actualizar modelos para la marca seleccionada
                self._actualizar_modelos()
                
                # Cargar modelo
                modelos = [m[1] for m in self.controller.db.obtener_modelos(
                    next(m[0] for m in self.controller.db.obtener_marcas() 
                        if m[1] == equipo['marca'])
                )]
                if equipo['modelo'] in modelos:
                    self.modelo_combobox.set(equipo['modelo'])
            
            self.serie_entry.insert(0, equipo['serie'])
            self.ubicacion_entry.insert(0, equipo['ubicacion'] or "")
            
            # Convertir fecha de string a objeto date
            fecha_ingreso = datetime.strptime(equipo['fecha_ingreso'], "%Y-%m-%d").date()
            self.fecha_ingreso_entry.set_date(fecha_ingreso)
            
            self.falla_entry.insert(0, equipo['falla'])
            self.observaciones_text.insert("1.0", equipo['observaciones'] or "")
    
    def _actualizar_modelos(self):
        """Actualiza los modelos según la marca seleccionada"""
        try:
            marca_seleccionada = self.marca_combobox.get()
            if marca_seleccionada:
                marcas = self.controller.db.obtener_marcas()
                id_marca = next(m[0] for m in marcas if m[1] == marca_seleccionada)
                
                modelos = self.controller.db.obtener_modelos(id_marca)
                self.modelo_combobox["values"] = [m[1] for m in modelos]
                if modelos:
                    self.modelo_combobox.current(0)
        except Exception as e:
            logging.error(f"Error actualizando modelos: {e}")
            messagebox.showerror("Error", f"No se pudieron cargar los modelos: {str(e)}")
    
    def _actualizar_equipo(self):
        """Actualiza el equipo en la base de datos"""
        try:
            # Validar campos obligatorios
            campos_obligatorios = {
                "PJ": self.pj_entry.get(),
                "Tipo de equipo": self.tipo_combobox.get(),
                "Marca": self.marca_combobox.get(),
                "Modelo": self.modelo_combobox.get(),
                "Número de serie": self.serie_entry.get(),
                "Falla": self.falla_entry.get()
            }
            
            faltantes = [campo for campo, valor in campos_obligatorios.items() if not valor]
            if faltantes:
                raise ValueError(f"Complete los campos obligatorios:\n- " + "\n- ".join(faltantes))
            
            # Validar número de serie único (excluyendo el equipo actual)
            serie = self.serie_entry.get().strip()
            equipos = self.controller.db.obtener_equipos(f"e.serie = '{serie}' AND e.id != {self.equipo_id}")
            if equipos:
                raise ValueError(f"El número de serie '{serie}' ya existe en otro equipo")
            
            # Obtener IDs de las selecciones
            tipo_id = next(t[0] for t in self.controller.db.obtener_tipos_equipo() 
                         if t[1] == self.tipo_combobox.get())
            marca_id = next(m[0] for m in self.controller.db.obtener_marcas() 
                          if m[1] == self.marca_combobox.get())
            modelo_id = next(m[0] for m in self.controller.db.obtener_modelos(marca_id) 
                         if m[1] == self.modelo_combobox.get())
            
            # Preparar datos
            datos = (
                self.pj_entry.get().strip(),
                tipo_id,
                serie,
                marca_id,
                modelo_id,
                self.ubicacion_entry.get().strip() or None,
                self.fecha_ingreso_entry.get_date().strftime("%Y-%m-%d"),
                self.falla_entry.get().strip(),
                self.observaciones_text.get("1.0", tk.END).strip() or None
            )
            
            # Actualizar en base de datos
            self.controller.db.actualizar_equipo(
                self.equipo_id, datos, self.controller.current_user['id'])
            
            messagebox.showinfo("Éxito", "Equipo actualizado correctamente")
            self.controller.mostrar_vista("MainView")
            
        except ValueError as e:
            messagebox.showerror("Error", str(e))
        except Exception as e:
            logging.error(f"Error actualizando equipo: {e}")
            messagebox.showerror("Error", f"No se pudo actualizar el equipo:\n{str(e)}")

class MarcasView(ttk.Frame):
    """Vista para gestionar marcas de equipos"""


    def initialize(self, *args, **kwargs):
        # Verificar autenticación primero
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
    
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self._setup_ui()
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        
        # Frame superior con botones
        button_frame = ttk.Frame(self)
        button_frame.grid(row=0, column=0, sticky="ew", pady=5)
        
        ttk.Button(button_frame, text="Agregar Marca", 
                  command=self._agregar_marca).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Modificar Marca", 
                  command=self._modificar_marca).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Eliminar Marca", 
                  command=self._eliminar_marca).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Volver", 
                  command=lambda: self.controller.mostrar_vista("MainView")).pack(side=tk.RIGHT, padx=5)
        
        # Treeview para mostrar las marcas
        self.treeview = ttk.Treeview(self, columns=("ID", "Nombre"), show="headings")
        self.treeview.heading("ID", text="ID")
        self.treeview.heading("Nombre", text="Nombre")
        self.treeview.column("ID", width=50, anchor=tk.CENTER)
        self.treeview.column("Nombre", width=200, anchor=tk.W)
        
        scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.treeview.yview)
        self.treeview.configure(yscrollcommand=scrollbar.set)
        
        self.treeview.grid(row=1, column=0, sticky="nsew")
        scrollbar.grid(row=1, column=1, sticky="ns")
        
        # Cargar datos iniciales
        self._cargar_marcas()
    
    def _cargar_marcas(self):
        """Carga las marcas desde la base de datos"""
        for item in self.treeview.get_children():
            self.treeview.delete(item)
            
        marcas = self.controller.db.obtener_marcas()
        for marca in marcas:
            self.treeview.insert("", tk.END, values=marca)
    
    def _agregar_marca(self):
        """Agrega una nueva marca"""
        nombre = simpledialog.askstring("Agregar Marca", "Ingrese el nombre de la nueva marca:")
        if nombre:
            try:
                self.controller.db.agregar_marca(nombre, self.controller.current_user['id'])
                self._cargar_marcas()
                messagebox.showinfo("Éxito", "Marca agregada correctamente")
            except ValueError as e:
                messagebox.showerror("Error", str(e))
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo agregar la marca: {str(e)}")
    
    def _modificar_marca(self):
        """Modifica una marca existente"""
        seleccion = self.treeview.selection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione una marca para modificar")
            return
            
        item = self.treeview.item(seleccion[0])
        marca_id, nombre_actual = item['values']
        
        nuevo_nombre = simpledialog.askstring("Modificar Marca", "Ingrese el nuevo nombre:", 
                                            initialvalue=nombre_actual)
        if nuevo_nombre and nuevo_nombre != nombre_actual:
            try:
                # Primero verificamos que no exista ya una marca con ese nombre
                marcas = self.controller.db.obtener_marcas()
                if any(nuevo_nombre.lower() == m[1].lower() for m in marcas):
                    raise ValueError("Ya existe una marca con ese nombre")
                
                # Actualizamos la marca
                self.controller.db.cursor.execute(
                    "UPDATE marcas SET nombre = ? WHERE id = ?",
                    (nuevo_nombre, marca_id))
                
                self.controller.db.registrar_auditoria(
                    self.controller.current_user['id'], 'MODIFICAR_MARCA', 'marcas', marca_id,
                    f"Modificación de marca: {nombre_actual} -> {nuevo_nombre}")
                
                self.controller.db.conn.commit()
                self._cargar_marcas()
                messagebox.showinfo("Éxito", "Marca modificada correctamente")
            except ValueError as e:
                messagebox.showerror("Error", str(e))
            except Exception as e:
                self.controller.db.conn.rollback()
                messagebox.showerror("Error", f"No se pudo modificar la marca: {str(e)}")
    
    def _eliminar_marca(self):
        """Elimina una marca existente"""
        seleccion = self.treeview.selection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione una marca para eliminar")
            return
            
        item = self.treeview.item(seleccion[0])
        marca_id, nombre = item['values']
        
        # Verificar si la marca tiene modelos asociados
        self.controller.db.cursor.execute(
            "SELECT COUNT(*) FROM modelos WHERE id_marca = ?", (marca_id,))
        count_modelos = self.controller.db.cursor.fetchone()[0]
        
        if count_modelos > 0:
            messagebox.showerror("Error", 
                               "No se puede eliminar la marca porque tiene modelos asociados.\n"
                               "Elimine primero los modelos relacionados.")
            return
            
        if messagebox.askyesno("Confirmar", f"¿Está seguro que desea eliminar la marca '{nombre}'?"):
            try:
                self.controller.db.cursor.execute(
                    "DELETE FROM marcas WHERE id = ?", (marca_id,))
                
                self.controller.db.registrar_auditoria(
                    self.controller.current_user['id'], 'ELIMINAR_MARCA', 'marcas', marca_id,
                    f"Marca eliminada: {nombre}")
                
                self.controller.db.conn.commit()
                self._cargar_marcas()
                messagebox.showinfo("Éxito", "Marca eliminada correctamente")
            except Exception as e:
                self.controller.db.conn.rollback()
                messagebox.showerror("Error", f"No se pudo eliminar la marca: {str(e)}")

class ModelosView(ttk.Frame):
    """Vista para gestionar modelos de equipos"""


    def initialize(self, *args, **kwargs):
        # Verificar autenticación primero
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
    
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self._setup_ui()
        
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        
        # Frame superior con controles
        control_frame = ttk.Frame(self)
        control_frame.grid(row=0, column=0, sticky="ew", pady=5)
        
        # Combobox para seleccionar marca
        ttk.Label(control_frame, text="Marca:").pack(side=tk.LEFT, padx=5)
        self.marca_combobox = ttk.Combobox(control_frame, state="readonly")
        self.marca_combobox.pack(side=tk.LEFT, padx=5)
        self.marca_combobox.bind("<<ComboboxSelected>>", lambda e: self._cargar_modelos())
        
        # Botones
        button_frame = ttk.Frame(control_frame)
        button_frame.pack(side=tk.RIGHT, padx=5)
        
        ttk.Button(button_frame, text="Agregar Modelo", 
                  command=self._agregar_modelo).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Modificar Modelo", 
                  command=self._modificar_modelo).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Eliminar Modelo", 
                  command=self._eliminar_modelo).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Volver", 
                  command=lambda: self.controller.mostrar_vista("MainView")).pack(side=tk.LEFT, padx=5)
        
        # Treeview para mostrar los modelos
        self.treeview = ttk.Treeview(self, columns=("ID", "Marca", "Nombre"), show="headings")
        self.treeview.heading("ID", text="ID")
        self.treeview.heading("Marca", text="Marca")
        self.treeview.heading("Nombre", text="Nombre")
        self.treeview.column("ID", width=50, anchor=tk.CENTER)
        self.treeview.column("Marca", width=150, anchor=tk.W)
        self.treeview.column("Nombre", width=150, anchor=tk.W)
        
        scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.treeview.yview)
        self.treeview.configure(yscrollcommand=scrollbar.set)
        
        self.treeview.grid(row=1, column=0, sticky="nsew")
        scrollbar.grid(row=1, column=1, sticky="ns")
        
        # Cargar datos iniciales
        self._cargar_marcas()
    
    def _cargar_marcas(self):
        """Carga las marcas en el combobox"""
        self.marca_combobox['values'] = [m[1] for m in self.controller.db.obtener_marcas()]
        if self.marca_combobox['values']:
            self.marca_combobox.current(0)
            self._cargar_modelos()
    
    def _cargar_modelos(self):
        """Carga los modelos de la marca seleccionada"""
        marca_seleccionada = self.marca_combobox.get()
        if not marca_seleccionada:
            return
            
        # Limpiar treeview
        for item in self.treeview.get_children():
            self.treeview.delete(item)
            
        # Obtener ID de la marca seleccionada
        marcas = self.controller.db.obtener_marcas()
        marca_id = next(m[0] for m in marcas if m[1] == marca_seleccionada)
        
        # Obtener y mostrar modelos
        modelos = self.controller.db.obtener_modelos(marca_id)
        for modelo in modelos:
            self.treeview.insert("", tk.END, values=(modelo[0], marca_seleccionada, modelo[1]))
    
    def _agregar_modelo(self):
        """Agrega un nuevo modelo a la marca seleccionada"""
        marca_seleccionada = self.marca_combobox.get()
        if not marca_seleccionada:
            messagebox.showwarning("Advertencia", "Seleccione una marca primero")
            return
            
        nombre = simpledialog.askstring("Agregar Modelo", "Ingrese el nombre del nuevo modelo:")
        if nombre:
            try:
                # Obtener ID de la marca seleccionada
                marcas = self.controller.db.obtener_marcas()
                marca_id = next(m[0] for m in marcas if m[1] == marca_seleccionada)
                
                # Agregar el modelo
                self.controller.db.agregar_modelo(marca_id, nombre, self.controller.current_user['id'])
                self._cargar_modelos()
                messagebox.showinfo("Éxito", "Modelo agregado correctamente")
            except ValueError as e:
                messagebox.showerror("Error", str(e))
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo agregar el modelo: {str(e)}")
    
    def _modificar_modelo(self):
        """Modifica un modelo existente"""
        seleccion = self.treeview.selection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione un modelo para modificar")
            return
            
        item = self.treeview.item(seleccion[0])
        modelo_id, marca_nombre, nombre_actual = item['values']
        
        nuevo_nombre = simpledialog.askstring("Modificar Modelo", "Ingrese el nuevo nombre:", 
                                            initialvalue=nombre_actual)
        if nuevo_nombre and nuevo_nombre != nombre_actual:
            try:
                # Obtener ID de la marca
                marcas = self.controller.db.obtener_marcas()
                marca_id = next(m[0] for m in marcas if m[1] == marca_nombre)
                
                # Verificar que no exista ya un modelo con ese nombre para esta marca
                modelos = self.controller.db.obtener_modelos(marca_id)
                if any(nuevo_nombre.lower() == m[1].lower() for m in modelos):
                    raise ValueError("Ya existe un modelo con ese nombre para esta marca")
                
                # Actualizar el modelo
                self.controller.db.cursor.execute(
                    "UPDATE modelos SET nombre = ? WHERE id = ?",
                    (nuevo_nombre, modelo_id))
                
                self.controller.db.registrar_auditoria(
                    self.controller.current_user['id'], 'MODIFICAR_MODELO', 'modelos', modelo_id,
                    f"Modificación de modelo: {nombre_actual} -> {nuevo_nombre}")
                
                self.controller.db.conn.commit()
                self._cargar_modelos()
                messagebox.showinfo("Éxito", "Modelo modificado correctamente")
            except ValueError as e:
                messagebox.showerror("Error", str(e))
            except Exception as e:
                self.controller.db.conn.rollback()
                messagebox.showerror("Error", f"No se pudo modificar el modelo: {str(e)}")
    
    def _eliminar_modelo(self):
        """Elimina un modelo existente"""
        seleccion = self.treeview.selection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione un modelo para eliminar")
            return
            
        item = self.treeview.item(seleccion[0])
        modelo_id, marca_nombre, nombre = item['values']
        
        # Verificar si el modelo tiene equipos asociados
        self.controller.db.cursor.execute(
            "SELECT COUNT(*) FROM equipos WHERE id_modelo = ?", (modelo_id,))
        count_equipos = self.controller.db.cursor.fetchone()[0]
        
        if count_equipos > 0:
            messagebox.showerror("Error", 
                               "No se puede eliminar el modelo porque tiene equipos asociados.\n"
                               "Elimine primero los equipos relacionados.")
            return
            
        if messagebox.askyesno("Confirmar", f"¿Está seguro que desea eliminar el modelo '{nombre}'?"):
            try:
                self.controller.db.cursor.execute(
                    "DELETE FROM modelos WHERE id = ?", (modelo_id,))
                
                self.controller.db.registrar_auditoria(
                    self.controller.current_user['id'], 'ELIMINAR_MODELO', 'modelos', modelo_id,
                    f"Modelo eliminado: {nombre} (Marca: {marca_nombre})")
                
                self.controller.db.conn.commit()
                self._cargar_modelos()
                messagebox.showinfo("Éxito", "Modelo eliminado correctamente")
            except Exception as e:
                self.controller.db.conn.rollback()
                messagebox.showerror("Error", f"No se pudo eliminar el modelo: {str(e)}")

class ReparacionView(ttk.Frame):
    """Vista para gestionar la reparación de equipos"""



    def initialize(self, *args, **kwargs):
        
        # Verificar autenticación primero
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return

            
        
     
    def _imprimir_informe(self):
        """Genera e imprime el informe técnico del equipo"""
        try:
            equipo = self.controller.db.obtener_equipo_por_id(self.equipo_id)
            if not equipo:
                messagebox.showerror("Error", "No se pudo obtener información del equipo")
                return
            
            repuestos = self.controller.db.obtener_repuestos_por_equipo(self.equipo_id)
            
            informe_data = {
                "titulo": f"Informe Técnico - Equipo {equipo['serie']}",
                "info_equipo": [
                    ["ID del Equipo", equipo['id']],
                    ["Número de Serie", equipo['serie']],
                    ["Tipo de Equipo", equipo['tipo_equipo']],
                    ["Marca", equipo['marca']],
                    ["Modelo", equipo['modelo']],
                    ["PJ", equipo['pj']],
                    ["Ubicación", equipo['ubicacion'] or "N/A"],
                    ["Fecha de Ingreso", equipo['fecha_ingreso']],
                    ["Fecha de Salida", equipo['fecha_salida'] or "N/A"],
                    ["Estado", equipo['estado']],
                    ["Falla Reportada", equipo['falla']]
                ],
                "repuestos": [["Repuesto", "Cantidad", "Costo Unitario", "Costo Total"]] + [
                    [rep[1], rep[2], f"${rep[3]:.2f}" if rep[3] else "N/A", f"${rep[2]*rep[3]:.2f}" if rep[3] else "N/A"]
                    for rep in repuestos
                ],
                "observaciones": equipo['observaciones'] or "Sin observaciones adicionales",
                "estado": equipo['estado']
            }
            
            fecha = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Informe_Tecnico_{equipo['serie']}_{fecha}.docx"
            
            if ExportManager.export_informe_tecnico(informe_data, filename):
                messagebox.showinfo("Éxito", f"Informe técnico generado:\n{filename}")
                if messagebox.askyesno("Abrir", "¿Desea abrir el documento para imprimir?"):
                    webbrowser.open(filename)
            else:
                messagebox.showerror("Error", "No se pudo generar el informe técnico")
                
        except Exception as e:
            logging.error(f"Error generando informe técnico: {e}")
            messagebox.showerror("Error", f"No se pudo generar el informe: {str(e)}")
        
    



    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self.equipo_id = None
        self._setup_ui()
    
    def initialize(self, equipo_id):
        """Inicializa la vista con el ID del equipo a reparar"""
        self.equipo_id = equipo_id
        self._cargar_datos_equipo()
        self._cargar_repuestos()
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        
        # Frame superior con información del equipo
        self.info_frame = ttk.LabelFrame(self, text="Información del Equipo", padding="10")
        self.info_frame.grid(row=0, column=0, sticky="ew", pady=5)
        
        # Frame para repuestos
        repuestos_frame = ttk.LabelFrame(self, text="Repuestos Utilizados", padding="10")
        repuestos_frame.grid(row=1, column=0, sticky="nsew", pady=5)
        repuestos_frame.grid_columnconfigure(0, weight=1)
        repuestos_frame.grid_rowconfigure(1, weight=1)
        
        # Botones para repuestos
        repuestos_btn_frame = ttk.Frame(repuestos_frame)
        repuestos_btn_frame.grid(row=0, column=0, sticky="ew", pady=5)
        
        ttk.Button(repuestos_btn_frame, text="Agregar Repuesto", 
                  command=self._agregar_repuesto).pack(side=tk.LEFT, padx=5)
        ttk.Button(repuestos_btn_frame, text="Eliminar Repuesto", 
                  command=self._eliminar_repuesto).pack(side=tk.LEFT, padx=5)
        
        # Treeview para repuestos
        self.repuestos_tree = ttk.Treeview(repuestos_frame, columns=("ID", "Nombre", "Cantidad", "Costo"), show="headings")
        self.repuestos_tree.heading("ID", text="ID")
        self.repuestos_tree.heading("Nombre", text="Nombre")
        self.repuestos_tree.heading("Cantidad", text="Cantidad")
        self.repuestos_tree.heading("Costo", text="Costo")
        self.repuestos_tree.column("ID", width=50, anchor=tk.CENTER)
        self.repuestos_tree.column("Nombre", width=200, anchor=tk.W)
        self.repuestos_tree.column("Cantidad", width=80, anchor=tk.CENTER)
        self.repuestos_tree.column("Costo", width=100, anchor=tk.E)
        
        scrollbar = ttk.Scrollbar(repuestos_frame, orient="vertical", command=self.repuestos_tree.yview)
        self.repuestos_tree.configure(yscrollcommand=scrollbar.set)
        
        self.repuestos_tree.grid(row=1, column=0, sticky="nsew")
        scrollbar.grid(row=1, column=1, sticky="ns")
        
        # Frame para estado y observaciones
        estado_frame = ttk.LabelFrame(self, text="Estado y Observaciones", padding="10")
        estado_frame.grid(row=2, column=0, sticky="ew", pady=5)
        
        ttk.Label(estado_frame, text="Estado:").grid(row=0, column=0, sticky="w", pady=5)
        self.estado_combobox = ttk.Combobox(estado_frame, values=["En reparación", "Reparado", "Irreparable"], state="readonly")
        self.estado_combobox.grid(row=0, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(estado_frame, text="Observaciones:").grid(row=1, column=0, sticky="nw", pady=5)
        self.observaciones_text = tk.Text(estado_frame, height=5, width=50)
        self.observaciones_text.grid(row=1, column=1, sticky="ew", padx=5, pady=5)
        
        # Botones inferiores
        btn_frame = ttk.Frame(self)
        btn_frame.grid(row=3, column=0, sticky="e", pady=10)
        
        ttk.Button(btn_frame, text="Guardar Cambios", 
                  command=self._guardar_cambios).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Volver", 
                  command=lambda: self.controller.mostrar_vista("MainView")).pack(side=tk.LEFT, padx=5)

    def _cargar_datos_equipo(self):
        """Carga los datos del equipo a reparar"""
        equipo = self.controller.db.obtener_equipo_por_id(self.equipo_id)
        if equipo:
            # Limpiar frame de información
            for widget in self.info_frame.winfo_children():
                widget.destroy()
            
            # Mostrar información del equipo
            info_text = (f"ID: {equipo['id']} | PJ: {equipo['pj']}\n"
                        f"Tipo: {equipo['tipo_equipo']} | Marca: {equipo['marca']} | Modelo: {equipo['modelo']}\n"
                        f"Serie: {equipo['serie']} | Ubicación: {equipo['ubicacion'] or 'N/A'}\n"
                        f"Falla: {equipo['falla']}\n"
                        f"Fecha ingreso: {equipo['fecha_ingreso']}")
            
            ttk.Label(self.info_frame, text=info_text).pack(anchor="w")
            
            # Establecer estado actual
            self.estado_combobox.set(equipo['estado'])
            self.observaciones_text.insert("1.0", equipo['observaciones'] or "")
    
    def _cargar_repuestos(self):
        """Carga los repuestos utilizados en el equipo"""
        for item in self.repuestos_tree.get_children():
            self.repuestos_tree.delete(item)
            
        repuestos = self.controller.db.obtener_repuestos_por_equipo(self.equipo_id)
        for rep in repuestos:
            self.repuestos_tree.insert("", tk.END, values=rep)
    
    def _agregar_repuesto(self):
        """Agrega un nuevo repuesto al equipo"""
        # Crear una ventana de diálogo personalizada para mantener el foco
        dialog = tk.Toplevel()
        dialog.title("Agregar Repuesto")
        dialog.transient(self)  # Establece relación con la ventana principal
        dialog.grab_set()  # Captura el foco
        
        # Configurar la ventana
        ttk.Label(dialog, text="Nombre del repuesto:").grid(row=0, column=0, padx=5, pady=5)
        nombre_entry = ttk.Entry(dialog)
        nombre_entry.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(dialog, text="Cantidad:").grid(row=1, column=0, padx=5, pady=5)
        cantidad_entry = ttk.Entry(dialog)
        cantidad_entry.grid(row=1, column=1, padx=5, pady=5)
        cantidad_entry.insert(0, "1")  # Valor por defecto
        
        ttk.Label(dialog, text="Costo unitario:").grid(row=2, column=0, padx=5, pady=5)
        costo_entry = ttk.Entry(dialog)
        costo_entry.grid(row=2, column=1, padx=5, pady=5)
        costo_entry.insert(0, "0.0")  # Valor por defecto
        
        def on_ok():
            nombre = nombre_entry.get().strip()
            cantidad = cantidad_entry.get().strip()
            costo = costo_entry.get().strip()
            dialog.destroy()
            
            if nombre and cantidad and costo:
                try:
                    cantidad = int(cantidad)
                    costo = float(costo)
                    
                    if cantidad <= 0:
                        raise ValueError("La cantidad debe ser mayor a 0")
                    if costo < 0:
                        raise ValueError("El costo no puede ser negativo")
                    
                    self.controller.db.agregar_repuesto(
                        self.equipo_id, nombre, cantidad, costo, self.controller.current_user['id'])
                    self._cargar_repuestos()
                    messagebox.showinfo("Éxito", "Repuesto agregado correctamente")
                except ValueError as e:
                    messagebox.showerror("Error", f"Dato inválido: {str(e)}")
                except Exception as e:
                    messagebox.showerror("Error", f"No se pudo agregar el repuesto: {str(e)}")
        
        ttk.Button(dialog, text="Aceptar", command=on_ok).grid(row=3, column=0, columnspan=2, pady=10)
        
        # Centrar la ventana
        dialog.update_idletasks()
        width = dialog.winfo_width()
        height = dialog.winfo_height()
        x = (dialog.winfo_screenwidth() // 2) - (width // 2)
        y = (dialog.winfo_screenheight() // 2) - (height // 2)
        dialog.geometry(f'{width}x{height}+{x}+{y}')
        
        # Establecer foco en el primer campo
        nombre_entry.focus_set()
        
        # Esperar hasta que la ventana se cierre
        self.wait_window(dialog)
    
    def _eliminar_repuesto(self):
        """Elimina un repuesto seleccionado"""
        seleccion = self.repuestos_tree.selection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione un repuesto para eliminar")
            return
            
        repuesto_id = self.repuestos_tree.item(seleccion[0])['values'][0]
        
        try:
            self.controller.db.eliminar_repuesto(repuesto_id, self.controller.current_user['id'])
            self._cargar_repuestos()
            messagebox.showinfo("Éxito", "Repuesto eliminado correctamente")
        except ValueError as e:
            messagebox.showerror("Error", str(e))
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo eliminar el repuesto: {str(e)}")
    
    def _guardar_cambios(self):
        """Guarda los cambios de estado y observaciones"""
        estado = self.estado_combobox.get()
        observaciones = self.observaciones_text.get("1.0", tk.END).strip()
        
        try:
            self.controller.db.actualizar_estado_equipo(
                self.equipo_id, estado, observaciones, self.controller.current_user['id'])
            messagebox.showinfo("Éxito", "Cambios guardados correctamente")
            self.controller.mostrar_vista("MainView")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron guardar los cambios: {str(e)}")

class ReporteEquiposView(ttk.Frame):
    """Vista para generar reportes de equipos"""


    def initialize(self, *args, **kwargs):
        # Verificar autenticación primero
        if not self.controller.current_user:
           self.controller.mostrar_vista("LoginView")
           return
    
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self._setup_ui()



    def _procesar_datos_para_exportacion(self, data):
        """Procesa los datos del treeview para asegurar compatibilidad con la exportación"""
        if not data or len(data) < 1:
            return data
    
        processed_data = [data[0]]  # Mantener encabezados
    
        for row in data[1:]:
            processed_row = []
            for value in row:
                # Convertir fechas y manejar valores None
                if isinstance(value, str) and self._es_fecha(value):
                    try:
                        # Intentar parsear fecha y formatear consistentemente
                        fecha_obj = datetime.strptime(value, '%Y-%m-%d')
                        processed_value = fecha_obj.strftime('%d/%m/%Y')
                    except ValueError:
                        processed_value = value
                elif value is None:
                    processed_value = ""
                else:
                    processed_value = str(value)
            
                processed_row.append(processed_value)
            processed_data.append(processed_row)
    
        return processed_data

    def _es_fecha(self, cadena):
        """Intenta determinar si una cadena representa una fecha"""
        patrones_fecha = [
            r'\d{4}-\d{2}-\d{2}',  # YYYY-MM-DD
            r'\d{2}/\d{2}/\d{4}',   # DD/MM/YYYY
            r'\d{2}-\d{2}-\d{4}'    # DD-MM-YYYY
        ]
    
        return any(re.fullmatch(patron, cadena) for patron in patrones_fecha)   


        
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(0, weight=1)
        
        main_frame = ttk.Frame(self, padding="20")
        main_frame.grid(row=0, column=0, sticky="nsew")
        
        # Frame de filtros
        filtros_frame = ttk.LabelFrame(main_frame, text="Filtros", padding="10")
        filtros_frame.grid(row=0, column=0, sticky="ew", pady=5)
        
        ttk.Label(filtros_frame, text="Estado:").grid(row=0, column=0, sticky="w", padx=5)
        self.estado_combobox = ttk.Combobox(filtros_frame, values=["Todos", "En reparación", "Reparado", "Irreparable"], state="readonly")
        self.estado_combobox.set("Todos")
        self.estado_combobox.grid(row=0, column=1, sticky="ew", padx=5)
        
        ttk.Label(filtros_frame, text="Fecha desde:").grid(row=1, column=0, sticky="w", padx=5)
        self.fecha_desde_entry = DateEntry(filtros_frame, date_pattern="yyyy-mm-dd")
        self.fecha_desde_entry.grid(row=1, column=1, sticky="ew", padx=5)
        
        ttk.Label(filtros_frame, text="Fecha hasta:").grid(row=2, column=0, sticky="w", padx=5)
        self.fecha_hasta_entry = DateEntry(filtros_frame, date_pattern="yyyy-mm-dd")
        self.fecha_hasta_entry.grid(row=2, column=1, sticky="ew", padx=5)
        
        ttk.Button(filtros_frame, text="Aplicar Filtros", 
                  command=self._aplicar_filtros).grid(row=3, column=0, columnspan=2, pady=5)
        
        # Treeview para resultados
        self.resultados_tree = ttk.Treeview(main_frame, columns=(
            "ID", "PJ", "Tipo", "Marca", "Modelo", "Ubicación", "Ingreso", "Salida", "Estado"
        ), show="headings")
        
        columns = {
            "ID": {"width": 50, "anchor": tk.CENTER},
            "PJ": {"width": 100, "anchor": tk.W},
            "Tipo": {"width": 100, "anchor": tk.W},
            "Marca": {"width": 100, "anchor": tk.W},
            "Modelo": {"width": 120, "anchor": tk.W},
            "Ubicación": {"width": 120, "anchor": tk.W},
            "Ingreso": {"width": 100, "anchor": tk.CENTER},
            "Salida": {"width": 100, "anchor": tk.CENTER},
            "Estado": {"width": 120, "anchor": tk.W}
        }
        
        for col, config in columns.items():
            self.resultados_tree.heading(col, text=col)
            self.resultados_tree.column(col, **config)
        
        scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=self.resultados_tree.yview)
        self.resultados_tree.configure(yscrollcommand=scrollbar.set)
        
        self.resultados_tree.grid(row=1, column=0, sticky="nsew")
        scrollbar.grid(row=1, column=1, sticky="ns")
        
        # Botones de exportación
        export_frame = ttk.Frame(main_frame)
        export_frame.grid(row=2, column=0, sticky="e", pady=10)
        
        ttk.Button(export_frame, text="Exportar a PDF", 
                  command=lambda: self._exportar("pdf")).pack(side=tk.LEFT, padx=5)
        ttk.Button(export_frame, text="Exportar a Excel", 
                  command=lambda: self._exportar("excel")).pack(side=tk.LEFT, padx=5)
        ttk.Button(export_frame, text="Exportar a Word", 
                  command=lambda: self._exportar("word")).pack(side=tk.LEFT, padx=5)
        ttk.Button(export_frame, text="Volver", 
                  command=lambda: self.controller.mostrar_vista("MainView")).pack(side=tk.LEFT, padx=5)
        
        # Cargar datos iniciales
        self._cargar_datos()
    
    def _aplicar_filtros(self):
        """Aplica los filtros seleccionados"""
        self._cargar_datos()
    
    def _cargar_datos(self):
        """Carga los datos según los filtros aplicados"""
        # Construir filtro
        filtros = []
        
        estado = self.estado_combobox.get()
        if estado != "Todos":
            filtros.append(f"e.estado = '{estado}'")
        
        fecha_desde = self.fecha_desde_entry.get_date().strftime("%Y-%m-%d")
        fecha_hasta = self.fecha_hasta_entry.get_date().strftime("%Y-%m-%d")
        filtros.append(f"e.fecha_ingreso BETWEEN '{fecha_desde}' AND '{fecha_hasta}'")
        
        filtro = " AND ".join(filtros) if filtros else None
        
        # Limpiar treeview
        for item in self.resultados_tree.get_children():
            self.resultados_tree.delete(item)
            
        # Obtener y mostrar datos
        equipos = self.controller.db.obtener_equipos(filtro)
        for equipo in equipos:
            self.resultados_tree.insert("", tk.END, values=equipo)
    
    def _exportar(self, formato):
        """Exporta los resultados al formato especificado"""

        # ===== VERIFICAR AUTENTICACIÓN =====
        if not self.controller.current_user:
            messagebox.showwarning("Acceso denegado", "Debe iniciar sesión primero")
            self.controller.mostrar_vista("LoginView")
            return



        try:
            # Obtener datos del treeview
            data = []
            headers = [self.resultados_tree.heading(col)["text"] 
                      for col in self.resultados_tree["columns"]]
            data.append(headers)
    
            for item in self.resultados_tree.get_children():
                data.append(self.resultados_tree.item(item)["values"])

            # Validar que hay datos para exportar
            if len(data) <= 1:  # Solo encabezados
                messagebox.showwarning("Advertencia", "No hay datos para exportar")
                return

            # Procesar datos para asegurar compatibilidad
            processed_data = self._procesar_datos_para_exportacion(data)

            # Generar nombre de archivo con marca de tiempo
            fecha = datetime.now().strftime("%Y%m%d_%H%M%S")
            default_filename = f"reporte_equipos_{fecha}"

            # Configuración según formato
            format_config = {
                "pdf": {
                    "ext": "pdf",
                    "filetypes": [("PDF files", "*.pdf")],
                    "export_func": ExportManager.export_to_pdf,
                    "title": "Reporte de Equipos - " + datetime.now().strftime('%d/%m/%Y %H:%M')
                },
                "excel": {
                    "ext": "xlsx",
                    "filetypes": [("Excel files", "*.xlsx")],
                    "export_func": ExportManager.export_to_excel,
                    "title": "Reporte de Equipos",
                    "sheet_name": "Equipos"
                },
                "word": {
                    "ext": "docx",
                    "filetypes": [("Word files", "*.docx")],
                    "export_func": ExportManager.export_to_word,
                    "title": "Reporte de Equipos - " + datetime.now().strftime('%d/%m/%Y %H:%M')
                }
            }

            config = format_config.get(formato.lower())
            if not config:
                raise ValueError(f"Formato no soportado: {formato}")

            filename = filedialog.asksaveasfilename(
                defaultextension=f".{config['ext']}",
                filetypes=config['filetypes'],
                initialfile=default_filename,
                title=f"Guardar reporte como {formato.upper()}")

            if filename:
                # Mostrar diálogo de progreso
                progress = tk.Toplevel()
                progress.title("Exportando...")
                progress.geometry("300x100")
                tk.Label(progress, text=f"Generando archivo {formato.upper()}...").pack(pady=20)
                progress.grab_set()
                self.update()

                try:
                    # Configurar parámetros específicos para cada formato
                    kwargs = {'title': config['title']}
                    if formato.lower() == 'excel':
                        success = ExportManager.export_to_excel(
                            data=processed_data,
                            filename=filename,
                            title=f"Reporte de Equipos - {datetime.now().strftime('%d/%m/%Y')}",  # Título añadido
                            sheet_name="Equipos"
                       )

                    # Ejecutar exportación
                    success = config['export_func'](processed_data, filename, **kwargs)
                
                    if success:
                       messagebox.showinfo("Éxito", f"Reporte exportado correctamente a:\n{filename}")
                       if messagebox.askyesno("Abrir", f"¿Desea abrir el archivo {config['ext'].upper()}?"):
                           try:
                               webbrowser.open(filename)
                           except Exception as e:
                               logging.warning(f"No se pudo abrir el archivo: {str(e)}")
                    else:
                        messagebox.showerror("Error", f"No se pudo exportar el reporte a {formato.upper()}")
                finally:
                    progress.destroy()
            
        except PermissionError:
            messagebox.showerror("Error de permisos", 
                               "No tiene permisos para guardar en la ubicación seleccionada.\n"
                               "Por favor, elija otra ubicación o cierre el archivo si está abierto.")
        except Exception as e:
            logging.error(f"Error en exportación: {str(e)}", exc_info=True)
            messagebox.showerror("Error", 
                               f"No se pudo completar la exportación:\n{str(e)}\n"
                               "Consulte el archivo de logs para más detalles.")



      
 

       


        
class ReporteRepuestosView(ttk.Frame):
    """Vista para generar reportes de repuestos utilizados"""


    def initialize(self, *args, **kwargs):
        # Verificar autenticación primero
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
    
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self._setup_ui()


    def _procesar_datos_para_exportacion(self, data):
        """Prepara los datos para exportación, convirtiendo fechas y valores especiales"""
        processed = [data[0]]  # Mantener encabezados
    
        for row in data[1:]:
            processed_row = []
            for value in row:
                if value is None:
                   processed_value = ""
                elif isinstance(value, str) and self._es_fecha(value):
                    try:
                        processed_value = datetime.strptime(value, '%Y-%m-%d').strftime('%d/%m/%Y')
                    except ValueError:
                        processed_value = value
                else:
                    processed_value = str(value)
                processed_row.append(processed_value)
            processed.append(processed_row)
    
        return processed

    def _es_fecha(self, cadena):
        """Determina si una cadena parece ser una fecha"""
        try:
            datetime.strptime(cadena, '%Y-%m-%d')
            return True
        except ValueError:
            return False   
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(0, weight=1)
        
        main_frame = ttk.Frame(self, padding="20")
        main_frame.grid(row=0, column=0, sticky="nsew")
        
        # Frame de filtros
        filtros_frame = ttk.LabelFrame(main_frame, text="Filtros por Fecha", padding="10")
        filtros_frame.grid(row=0, column=0, sticky="ew", pady=5)
        
        ttk.Label(filtros_frame, text="Fecha desde:").grid(row=0, column=0, sticky="w", padx=5)
        self.fecha_desde_entry = DateEntry(filtros_frame, date_pattern="yyyy-mm-dd")
        self.fecha_desde_entry.grid(row=0, column=1, sticky="ew", padx=5)
        
        ttk.Label(filtros_frame, text="Fecha hasta:").grid(row=1, column=0, sticky="w", padx=5)
        self.fecha_hasta_entry = DateEntry(filtros_frame, date_pattern="yyyy-mm-dd")
        self.fecha_hasta_entry.grid(row=1, column=1, sticky="ew", padx=5)
        
        ttk.Button(filtros_frame, text="Generar Reporte", 
                  command=self._generar_reporte).grid(row=2, column=0, columnspan=2, pady=5)
        
        # Treeview para resultados
        self.resultados_tree = ttk.Treeview(main_frame, columns=(
            "Repuesto", "Cantidad", "Costo Total", "Equipo", "Tipo"
        ), show="headings")
        
        columns = {
            "Repuesto": {"width": 150, "anchor": tk.W},
            "Cantidad": {"width": 80, "anchor": tk.CENTER},
            "Costo Total": {"width": 100, "anchor": tk.E},
            "Equipo": {"width": 120, "anchor": tk.W},
            "Tipo": {"width": 120, "anchor": tk.W}
        }
        
        for col, config in columns.items():
            self.resultados_tree.heading(col, text=col)
            self.resultados_tree.column(col, **config)
        
        scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=self.resultados_tree.yview)
        self.resultados_tree.configure(yscrollcommand=scrollbar.set)
        
        self.resultados_tree.grid(row=1, column=0, sticky="nsew")
        scrollbar.grid(row=1, column=1, sticky="ns")
        
        # Botones de exportación
        export_frame = ttk.Frame(main_frame)
        export_frame.grid(row=2, column=0, sticky="e", pady=10)
        
        ttk.Button(export_frame, text="Exportar a PDF", 
                  command=lambda: self._exportar("pdf")).pack(side=tk.LEFT, padx=5)
        ttk.Button(export_frame, text="Exportar a Excel", 
                  command=lambda: self._exportar("excel")).pack(side=tk.LEFT, padx=5)
        ttk.Button(export_frame, text="Exportar a Word", 
                  command=lambda: self._exportar("word")).pack(side=tk.LEFT, padx=5)
        ttk.Button(export_frame, text="Volver", 
                  command=lambda: self.controller.mostrar_vista("MainView")).pack(side=tk.LEFT, padx=5)
    
    def _generar_reporte(self):
        """Genera el reporte con los filtros aplicados"""
        fecha_desde = self.fecha_desde_entry.get_date().strftime("%Y-%m-%d")
        fecha_hasta = self.fecha_hasta_entry.get_date().strftime("%Y-%m-%d")
        
        # Limpiar treeview
        for item in self.resultados_tree.get_children():
            self.resultados_tree.delete(item)
            
        # Obtener y mostrar datos
        repuestos = self.controller.db.obtener_repuestos_por_periodo(fecha_desde, fecha_hasta)
        for rep in repuestos:
            self.resultados_tree.insert("", tk.END, values=rep)
    
    def _exportar(self, formato):
        """Exporta los resultados al formato especificado"""



        # ===== VERIFICAR AUTENTICACIÓN =====
        if not self.controller.current_user:
            messagebox.showwarning("Acceso denegado", "Debe iniciar sesión primero")
            self.controller.mostrar_vista("LoginView")
            return


    
        try:
            # ===== OBTENER DATOS DEL TREEVIEW =====
            data = []
            headers = [self.resultados_tree.heading(col)["text"] 
                      for col in self.resultados_tree["columns"]]
            data.append(headers)
        
            for item in self.resultados_tree.get_children():
                data.append(self.resultados_tree.item(item)["values"])

            # ===== VALIDAR DATOS =====
            if len(data) <= 1:  # Solo encabezados sin datos
                messagebox.showwarning("Advertencia", "No hay datos para exportar")
                return

            # ===== PROCESAR DATOS =====
            # Convertir fechas y manejar valores nulos
            processed_data = self._procesar_datos_para_exportacion(data)

            # ===== CONFIGURACIÓN COMÚN =====
            fecha_reporte = self.fecha_desde_entry.get_date()
            periodo = f"{fecha_reporte.strftime('%d-%m-%Y')} al {self.fecha_hasta_entry.get_date().strftime('%d-%m-%Y')}"
            default_filename = f"reporte_repuestos_{datetime.now().strftime('%Y%m%d_%H%M%S')}"



            




            # ===== EXPORTACIÓN A EXCEL =====
            if formato.lower() == "excel":
                filename = filedialog.asksaveasfilename(  # <-- ÚNICA LLAMADA A filedialog
                    defaultextension=".xlsx",
                    filetypes=[("Excel files", "*.xlsx")],
                    initialfile=default_filename,
                    title="Guardar reporte como Excel")
                
                if filename:
                    # Mostrar progreso
                    progress = tk.Toplevel()
                    progress.title("Exportando a Excel...")
                    tk.Label(progress, text="Generando archivo Excel...").pack(pady=20)
                    progress.grab_set()
                    self.update()

                    try:
                        # Configurar nombre seguro para la hoja
                        sheet_name = f"Repuestos {periodo.split(' al ')[0]}"
                    
                        if ExportManager.export_to_excel(
                            data=processed_data,
                            filename=filename,
                            title=f"Reporte de Repuestos ({periodo})",
                            sheet_name=sheet_name
                        ):
                            messagebox.showinfo("Éxito", 
                                          f"Reporte exportado correctamente:\n{filename}")
                            if messagebox.askyesno("Abrir", "¿Desea abrir el archivo Excel?"):
                                try:
                                    webbrowser.open(filename)
                                except Exception as e:
                                    logging.warning(f"No se pudo abrir el archivo: {e}")
                                    messagebox.showwarning("Aviso", 
                                                         f"Archivo creado pero no se pudo abrir:\n{filename}")
                        else:
                            messagebox.showerror("Error", "No se pudo generar el archivo Excel")
                    finally:
                        progress.destroy()

            # ===== EXPORTACIÓN A PDF =====
            elif formato.lower() == "pdf":
                filename = filedialog.asksaveasfilename(
                    defaultextension=".pdf",
                    filetypes=[("PDF files", "*.pdf")],
                    initialfile=default_filename,
                    title="Guardar reporte como PDF")

                if filename:
                    progress = tk.Toplevel()
                    progress.title("Exportando a PDF...")
                    tk.Label(progress, text="Generando archivo PDF...").pack(pady=20)
                    progress.grab_set()
                    self.update()

                    try:
                        if ExportManager.export_to_pdf(
                            data=processed_data,
                            filename=filename,
                            title=f"Reporte de Repuestos ({periodo})"
                        ):
                            messagebox.showinfo("Éxito", 
                                          f"Reporte exportado correctamente:\n{filename}")
                            if messagebox.askyesno("Abrir", "¿Desea abrir el archivo PDF?"):
                                webbrowser.open(filename)
                        else:
                            messagebox.showerror("Error", "No se pudo generar el archivo PDF")
                    finally:
                        progress.destroy()

            # ===== EXPORTACIÓN A WORD =====
            elif formato.lower() == "word":
                filename = filedialog.asksaveasfilename(
                    defaultextension=".docx",
                    filetypes=[("Word files", "*.docx")],
                    initialfile=default_filename,
                    title="Guardar reporte como Word")

                if filename:
                    progress = tk.Toplevel()
                    progress.title("Exportando a Word...")
                    tk.Label(progress, text="Generando archivo Word...").pack(pady=20)
                    progress.grab_set()
                    self.update()

                    try:
                        if ExportManager.export_to_word(
                            data=processed_data,
                            filename=filename,
                            title=f"Reporte de Repuestos ({periodo})"
                        ):
                            messagebox.showinfo("Éxito", 
                                              f"Reporte exportado correctamente:\n{filename}")
                            if messagebox.askyesno("Abrir", "¿Desea abrir el archivo Word?"):
                                webbrowser.open(filename)
                        else:
                            messagebox.showerror("Error", "No se pudo generar el archivo Word")
                    finally:
                        progress.destroy()

            else:
                messagebox.showerror("Error", f"Formato no soportado: {formato}")

        except Exception as e:
            logging.error(f"Error durante la exportación: {str(e)}", exc_info=True)
            messagebox.showerror(
                 "Error crítico",
                 f"No se pudo completar la exportación:\n{str(e)}\n"
                 "Revise el archivo de logs para más detalles"
            )
        
class UsuariosView(ttk.Frame):
    """Vista para gestionar usuarios del sistema"""

    def initialize(self, *args, **kwargs):
        # Verificar autenticación primero
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
           
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self._setup_ui()
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        
        # Frame superior con botones
        button_frame = ttk.Frame(self)
        button_frame.grid(row=0, column=0, sticky="ew", pady=5)
        
        ttk.Button(button_frame, text="Agregar Usuario", 
                  command=self._agregar_usuario).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Modificar Usuario", 
                  command=self._modificar_usuario).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cambiar Contraseña", 
                  command=self._cambiar_password).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Activar/Desactivar", 
                  command=self._toggle_activo).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Volver", 
                  command=lambda: self.controller.mostrar_vista("MainView")).pack(side=tk.RIGHT, padx=5)
        
        # Treeview para mostrar los usuarios
        self.treeview = ttk.Treeview(self, columns=("ID", "Username", "Rol", "Activo"), show="headings")
        self.treeview.heading("ID", text="ID")
        self.treeview.heading("Username", text="Usuario")
        self.treeview.heading("Rol", text="Rol")
        self.treeview.heading("Activo", text="Activo")
        self.treeview.column("ID", width=50, anchor=tk.CENTER)
        self.treeview.column("Username", width=150, anchor=tk.W)
        self.treeview.column("Rol", width=100, anchor=tk.W)
        self.treeview.column("Activo", width=70, anchor=tk.CENTER)
        
        scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.treeview.yview)
        self.treeview.configure(yscrollcommand=scrollbar.set)
        
        self.treeview.grid(row=1, column=0, sticky="nsew")
        scrollbar.grid(row=1, column=1, sticky="ns")
        
        # Cargar datos iniciales
        self._cargar_usuarios()
    
    def _cargar_usuarios(self):
        """Carga los usuarios desde la base de datos"""
        for item in self.treeview.get_children():
            self.treeview.delete(item)
            
        try:
            self.controller.db.cursor.execute(
                "SELECT id, username, rol, activo FROM usuarios ORDER BY username")
            usuarios = self.controller.db.cursor.fetchall()
            
            for usuario in usuarios:
                estado = "Sí" if usuario[3] else "No"
                self.treeview.insert("", tk.END, values=(usuario[0], usuario[1], usuario[2], estado))
        except sqlite3.Error as e:
            logging.error(f"Error cargando usuarios: {e}")
            messagebox.showerror("Error", f"No se pudieron cargar los usuarios: {str(e)}")
    
    def _agregar_usuario(self):
        """Agrega un nuevo usuario al sistema"""
        dialog = tk.Toplevel()
        dialog.title("Agregar Usuario")
        dialog.transient(self)
        dialog.grab_set()
        
        # Variables
        username_var = tk.StringVar()
        password_var = tk.StringVar()
        confirm_var = tk.StringVar()
        rol_var = tk.StringVar(value="tecnico")  # Valor por defecto
        
        # Frame principal
        main_frame = ttk.Frame(dialog, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Campos del formulario
        ttk.Label(main_frame, text="Nombre de usuario:").grid(row=0, column=0, sticky="w", pady=5)
        username_entry = ttk.Entry(main_frame, textvariable=username_var)
        username_entry.grid(row=0, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(main_frame, text="Contraseña:").grid(row=1, column=0, sticky="w", pady=5)
        password_entry = ttk.Entry(main_frame, textvariable=password_var, show="•")
        password_entry.grid(row=1, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(main_frame, text="Confirmar contraseña:").grid(row=2, column=0, sticky="w", pady=5)
        confirm_entry = ttk.Entry(main_frame, textvariable=confirm_var, show="•")
        confirm_entry.grid(row=2, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(main_frame, text="Rol:").grid(row=3, column=0, sticky="w", pady=5)
        rol_combobox = ttk.Combobox(main_frame, textvariable=rol_var, 
                                   values=["admin", "tecnico", "consulta"], state="readonly")
        rol_combobox.grid(row=3, column=1, sticky="ew", padx=5, pady=5)
        
        # Botones
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=4, column=0, columnspan=2, pady=10)
        
        def on_ok():
            username = username_var.get().strip()
            password = password_var.get()
            confirm = confirm_var.get()
            rol = rol_var.get()
            
            if not username or not password:
                messagebox.showwarning("Advertencia", "Complete todos los campos")
                return
                
            if password != confirm:
                messagebox.showwarning("Advertencia", "Las contraseñas no coinciden")
                return
                
            try:
                hashed_pw = hashlib.sha256(password.encode('utf-8')).hexdigest()
                
                self.controller.db.cursor.execute(
                    "INSERT INTO usuarios (username, password, rol, activo) VALUES (?, ?, ?, 1)",
                    (username, hashed_pw, rol))
                
                user_id = self.controller.db.cursor.lastrowid
                
                self.controller.db.registrar_auditoria(
                    self.controller.current_user['id'], 'ALTA_USUARIO', 'usuarios', user_id,
                    f"Nuevo usuario creado: {username} (Rol: {rol})")
                
                self.controller.db.conn.commit()
                self._cargar_usuarios()
                dialog.destroy()
                messagebox.showinfo("Éxito", "Usuario creado correctamente")
            except sqlite3.IntegrityError:
                messagebox.showerror("Error", "El nombre de usuario ya existe")
            except Exception as e:
                self.controller.db.conn.rollback()
                messagebox.showerror("Error", f"No se pudo crear el usuario: {str(e)}")
        
        ttk.Button(button_frame, text="Aceptar", command=on_ok).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cancelar", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        
        # Centrar la ventana
        dialog.update_idletasks()
        width = dialog.winfo_width()
        height = dialog.winfo_height()
        x = (dialog.winfo_screenwidth() // 2) - (width // 2)
        y = (dialog.winfo_screenheight() // 2) - (height // 2)
        dialog.geometry(f'{width}x{height}+{x}+{y}')
        
        username_entry.focus_set()
    
    def _modificar_usuario(self):
        """Modifica el rol de un usuario existente"""
        seleccion = self.treeview.selection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione un usuario para modificar")
            return
            
        item = self.treeview.item(seleccion[0])
        user_id, username, rol_actual, _ = item['values']
        
        # Proteger al usuario admin
        if username == "admin":
            messagebox.showwarning("Advertencia", "No se puede modificar el usuario administrador")
            return
            
        dialog = tk.Toplevel()
        dialog.title("Modificar Usuario")
        dialog.transient(self)
        dialog.grab_set()
        
        # Variables
        rol_var = tk.StringVar(value=rol_actual)
        
        # Frame principal
        main_frame = ttk.Frame(dialog, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Campos del formulario
        ttk.Label(main_frame, text=f"Usuario: {username}").grid(row=0, column=0, sticky="w", pady=5)
        ttk.Label(main_frame, text="Nuevo rol:").grid(row=1, column=0, sticky="w", pady=5)
        rol_combobox = ttk.Combobox(main_frame, textvariable=rol_var, 
                                   values=["admin", "tecnico", "consulta"], state="readonly")
        rol_combobox.grid(row=1, column=1, sticky="ew", padx=5, pady=5)
        
        # Botones
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=2, column=0, columnspan=2, pady=10)
        
        def on_ok():
            nuevo_rol = rol_var.get()
            
            if nuevo_rol == rol_actual:
                dialog.destroy()
                return
                
            try:
                self.controller.db.cursor.execute(
                    "UPDATE usuarios SET rol = ? WHERE id = ?",
                    (nuevo_rol, user_id))
                
                self.controller.db.registrar_auditoria(
                    self.controller.current_user['id'], 'MODIFICAR_USUARIO', 'usuarios', user_id,
                    f"Rol de usuario modificado: {username} ({rol_actual} -> {nuevo_rol})")
                
                self.controller.db.conn.commit()
                self._cargar_usuarios()
                dialog.destroy()
                messagebox.showinfo("Éxito", "Usuario modificado correctamente")
            except Exception as e:
                self.controller.db.conn.rollback()
                messagebox.showerror("Error", f"No se pudo modificar el usuario: {str(e)}")
        
        ttk.Button(button_frame, text="Aceptar", command=on_ok).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cancelar", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        
        # Centrar la ventana
        dialog.update_idletasks()
        width = dialog.winfo_width()
        height = dialog.winfo_height()
        x = (dialog.winfo_screenwidth() // 2) - (width // 2)
        y = (dialog.winfo_screenheight() // 2) - (height // 2)
        dialog.geometry(f'{width}x{height}+{x}+{y}')
    
    def _cambiar_password(self):
        """Cambia la contraseña de un usuario"""
        seleccion = self.treeview.selection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione un usuario para cambiar contraseña")
            return
            
        item = self.treeview.item(seleccion[0])
        user_id, username, _, _ = item['values']
        
        dialog = tk.Toplevel()
        dialog.title("Cambiar Contraseña")
        dialog.transient(self)
        dialog.grab_set()
        
        # Variables
        password_var = tk.StringVar()
        confirm_var = tk.StringVar()
        
        # Frame principal
        main_frame = ttk.Frame(dialog, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Campos del formulario
        ttk.Label(main_frame, text=f"Usuario: {username}").grid(row=0, column=0, sticky="w", pady=5)
        ttk.Label(main_frame, text="Nueva contraseña:").grid(row=1, column=0, sticky="w", pady=5)
        password_entry = ttk.Entry(main_frame, textvariable=password_var, show="•")
        password_entry.grid(row=1, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(main_frame, text="Confirmar contraseña:").grid(row=2, column=0, sticky="w", pady=5)
        confirm_entry = ttk.Entry(main_frame, textvariable=confirm_var, show="•")
        confirm_entry.grid(row=2, column=1, sticky="ew", padx=5, pady=5)
        
        # Botones
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=3, column=0, columnspan=2, pady=10)
        
        def on_ok():
            password = password_var.get()
            confirm = confirm_var.get()
            
            if not password:
                messagebox.showwarning("Advertencia", "Ingrese una contraseña")
                return
                
            if password != confirm:
                messagebox.showwarning("Advertencia", "Las contraseñas no coinciden")
                return
                
            try:
                hashed_pw = hashlib.sha256(password.encode('utf-8')).hexdigest()
                
                self.controller.db.cursor.execute(
                    "UPDATE usuarios SET password = ? WHERE id = ?",
                    (hashed_pw, user_id))
                
                self.controller.db.registrar_auditoria(
                    self.controller.current_user['id'], 'CAMBIAR_PASSWORD', 'usuarios', user_id,
                    f"Contraseña cambiada para usuario: {username}")
                
                self.controller.db.conn.commit()
                dialog.destroy()
                messagebox.showinfo("Éxito", "Contraseña cambiada correctamente")
            except Exception as e:
                self.controller.db.conn.rollback()
                messagebox.showerror("Error", f"No se pudo cambiar la contraseña: {str(e)}")
        
        ttk.Button(button_frame, text="Aceptar", command=on_ok).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cancelar", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        
        # Centrar la ventana
        dialog.update_idletasks()
        width = dialog.winfo_width()
        height = dialog.winfo_height()
        x = (dialog.winfo_screenwidth() // 2) - (width // 2)
        y = (dialog.winfo_screenheight() // 2) - (height // 2)
        dialog.geometry(f'{width}x{height}+{x}+{y}')
        
        password_entry.focus_set()
    
    def _toggle_activo(self):
        """Activa o desactiva un usuario"""
        seleccion = self.treeview.selection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione un usuario para activar/desactivar")
            return
            
        item = self.treeview.item(seleccion[0])
        user_id, username, _, estado_actual = item['values']
        
        # Proteger al usuario admin
        if username == "admin":
            messagebox.showwarning("Advertencia", "No se puede desactivar el usuario administrador")
            return
            
        nuevo_estado = 0 if estado_actual == "Sí" else 1
        accion = "desactivar" if nuevo_estado == 0 else "activar"
        
        if messagebox.askyesno("Confirmar", f"¿Está seguro que desea {accion} al usuario '{username}'?"):
            try:
                self.controller.db.cursor.execute(
                    "UPDATE usuarios SET activo = ? WHERE id = ?",
                    (nuevo_estado, user_id))
                
                self.controller.db.registrar_auditoria(
                    self.controller.current_user['id'], 'CAMBIAR_ESTADO', 'usuarios', user_id,
                    f"Usuario {accion}do: {username}")
                
                self.controller.db.conn.commit()
                self._cargar_usuarios()
                messagebox.showinfo("Éxito", f"Usuario {accion}do correctamente")
            except Exception as e:
                self.controller.db.conn.rollback()
                messagebox.showerror("Error", f"No se pudo {accion} el usuario: {str(e)}")


class ReporteResumenRepuestosView(ttk.Frame):
      """Vista para mostrar resumen de repuestos utilizados por mes/año"""
    
      def initialize(self, *args, **kwargs):
          if not self.controller.current_user:
              self.controller.mostrar_vista("LoginView")
              return
    
      def __init__(self, parent, controller):
          super().__init__(parent)
          self.controller = controller
          self._setup_ui()
    
      def _setup_ui(self):
          self.grid_columnconfigure(0, weight=1)
          self.grid_rowconfigure(1, weight=1)
        
          # Frame de controles
          control_frame = ttk.Frame(self)
          control_frame.grid(row=0, column=0, sticky="ew", pady=5)
        
          # Filtros
          ttk.Label(control_frame, text="Año:").pack(side=tk.LEFT, padx=5)
          self.anio_combobox = ttk.Combobox(control_frame, values=[str(y) for y in range(2020, datetime.now().year + 1)])
          self.anio_combobox.set(str(datetime.now().year))
          self.anio_combobox.pack(side=tk.LEFT, padx=5)
        
          ttk.Label(control_frame, text="Mes:").pack(side=tk.LEFT, padx=5)
          self.mes_combobox = ttk.Combobox(control_frame, state="readonly", 
                                          values=["Todos", "Enero", "Febrero", "Marzo", "Abril", 
                                                 "Mayo", "Junio", "Julio", "Agosto", 
                                                 "Septiembre", "Octubre", "Noviembre", "Diciembre"])
          self.mes_combobox.set("Todos")
          self.mes_combobox.pack(side=tk.LEFT, padx=5)
        
          ttk.Button(control_frame, text="Generar Reporte", 
                   command=self._generar_reporte).pack(side=tk.LEFT, padx=5)
        
          ttk.Button(control_frame, text="Volver", 
                   command=lambda: self.controller.mostrar_vista("MainView")).pack(side=tk.RIGHT, padx=5)
        
          # Treeview para mostrar los datos
          self.treeview = ttk.Treeview(self, columns=("Repuesto", "Cantidad", "Mes", "Año"), show="headings")
          self.treeview.heading("Repuesto", text="Repuesto")
          self.treeview.heading("Cantidad", text="Cantidad Total")
          self.treeview.heading("Mes", text="Mes")
          self.treeview.heading("Año", text="Año")
        
          self.treeview.column("Repuesto", width=200)
          self.treeview.column("Cantidad", width=100, anchor=tk.CENTER)
          self.treeview.column("Mes", width=100, anchor=tk.CENTER)
          self.treeview.column("Año", width=80, anchor=tk.CENTER)
        
          scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.treeview.yview)
          self.treeview.configure(yscrollcommand=scrollbar.set)
        
          self.treeview.grid(row=1, column=0, sticky="nsew")
          scrollbar.grid(row=1, column=1, sticky="ns")
        
          # Botón de exportación
          export_frame = ttk.Frame(self)
          export_frame.grid(row=2, column=0, sticky="e", pady=5)
        
          ttk.Button(export_frame, text="Exportar a Excel", 
                   command=self._exportar_excel).pack(side=tk.LEFT, padx=5)


          ttk.Button(export_frame, text="Exportar a Word",
                   command=self._exportar_word).pack(side=tk.LEFT, padx=5)  # Nuevo botón
        
          # Cargar datos iniciales
          self._generar_reporte()
    
      def _generar_reporte(self):
          """Genera el reporte con los filtros aplicados"""
          try:
              # Limpiar treeview
              for item in self.treeview.get_children():
                  self.treeview.delete(item)
            
              # Obtener parámetros
              anio = int(self.anio_combobox.get())
              mes = None if self.mes_combobox.get() == "Todos" else self._get_month_number(self.mes_combobox.get())
            
              # Obtener datos
              repuestos = self.controller.db.obtener_resumen_repuestos(anio, mes)
            
              # Mostrar datos
              meses = ["", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", 
                      "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
            
              for rep in repuestos:
                  nombre, cantidad, mes_num, año = rep
                  mes_nombre = meses[int(mes_num)] if mes_num else "Todos"
                  self.treeview.insert("", tk.END, values=(nombre, cantidad, mes_nombre, año))
                
          except Exception as e:
              messagebox.showerror("Error", f"No se pudo generar el reporte: {str(e)}")
    
      def _get_month_number(self, month_name):
          """Convierte nombre de mes a número"""
          meses = {"Enero": 1, "Febrero": 2, "Marzo": 3, "Abril": 4, "Mayo": 5, "Junio": 6,
                  "Julio": 7, "Agosto": 8, "Septiembre": 9, "Octubre": 10, "Noviembre": 11, "Diciembre": 12}
          return meses.get(month_name, 0)
    
      def _exportar_excel(self):
          """Exporta el reporte a Excel"""
          try:
              # Obtener datos del treeview
              data = []
              headers = [self.treeview.heading(col)["text"] for col in self.treeview["columns"]]
              data.append(headers)
            
              for item in self.treeview.get_children():
                  data.append(self.treeview.item(item)["values"])
            
              # Generar nombre de archivo
              anio = self.anio_combobox.get()
              mes = self.mes_combobox.get()
              default_filename = f"Resumen_repuestos_{mes}_{anio}.xlsx"
            
              filename = filedialog.asksaveasfilename(
                  defaultextension=".xlsx",
                  filetypes=[("Excel files", "*.xlsx")],
                  initialfile=default_filename,
                  title="Guardar reporte como Excel")
            
              if filename:
                  if ExportManager.export_to_excel(
                      data, filename, 
                      title=f"Resumen de repuestos - {mes} {anio}",
                      sheet_name=f"Repuestos {mes[:3]}"):
                    
                      messagebox.showinfo("Éxito", f"Reporte exportado correctamente:\n{filename}")
                      if messagebox.askyesno("Abrir", "¿Desea abrir el archivo Excel?"):
                          webbrowser.open(filename)
                  else:
                      messagebox.showerror("Error", "No se pudo exportar el reporte")
                    
          except Exception as e:
              messagebox.showerror("Error", f"No se pudo exportar: {str(e)}")



      def _exportar_word(self):
          """Exporta el reporte a Word con formato profesional"""
          try:
              # 1. Obtener datos del treeview
              data = []
              headers = [self.treeview.heading(col)["text"] for col in self.treeview["columns"]]
              data.append(headers)
        
              for item in self.treeview.get_children():
                  data.append(self.treeview.item(item)["values"])
        
              # 2. Definir título y nombre de archivo
              anio = self.anio_combobox.get()
              mes = self.mes_combobox.get()
              titulo = f"Resumen de Repuestos - {mes} {anio}"
              nombre_archivo = f"Resumen_repuestos_{mes}_{anio}.docx"
        
              # 3. Pedir ubicación para guardar
              filename = filedialog.asksaveasfilename(
                  defaultextension=".docx",
                  filetypes=[("Word files", "*.docx")],
                  initialfile=nombre_archivo,
                  title="Guardar reporte como Word"
              )
        
              if filename:
                  # 4. Mostrar progreso (opcional)
                  progress = tk.Toplevel()
                  progress.title("Generando Word...")
                  tk.Label(progress, text="Exportando a Word...").pack(pady=20)
                  progress.grab_set()
                  self.update()
            
                  # 5. Llamar al ExportManager
                  if ExportManager.export_to_word(
                      data, 
                      filename, 
                      title=titulo,
                      subtitle=f"Período: {mes} {anio}"
                  ):
                      messagebox.showinfo("Éxito", f"Documento guardado:\n{filename}")
                      if messagebox.askyesno("Abrir", "¿Abrir el documento?"):
                          webbrowser.open(filename)
                  else:
                      messagebox.showerror("Error", "No se pudo generar el documento")
            
                  progress.destroy()
            
          except Exception as e:
              messagebox.showerror("Error", f"Error al exportar: {str(e)}")       


class TonerView(ttk.Frame):
    """Vista principal de gestión de toner"""
    
    def initialize(self, *args, **kwargs):
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
    
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self._setup_ui()
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        



        # Barra de herramientas
        toolbar = ttk.Frame(self)
        toolbar.pack(side=tk.TOP, fill=tk.X, padx=5, pady=5)
        
        ttk.Button(toolbar, text="Registrar Retiro", 
                  command=lambda: self.controller.mostrar_vista("RetiroTonerView")).pack(side=tk.LEFT, padx=5)
        ttk.Button(toolbar, text="Registrar Recarga", 
                  command=lambda: self.controller.mostrar_vista("RecargaTonerView")).pack(side=tk.LEFT, padx=5)
        ttk.Button(toolbar, text="Cargar Stock", 
                  command=self._cargar_stock).pack(side=tk.LEFT, padx=5)
        ttk.Button(toolbar, text="Volver", 
                  command=lambda: self.controller.mostrar_vista("MainView")).pack(side=tk.RIGHT, padx=5)
        ttk.Button(toolbar, text="Empresas Recarga", 
                  command=lambda: self.controller.mostrar_vista("EmpresasRecargaView")).pack(side=tk.LEFT, padx=5)
        
        # Notebook (pestañas)
        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Pestaña de Stock
        self.stock_frame = ttk.Frame(self.notebook)
        self._setup_stock_ui()  # Esto ahora funcionará correctamente
        self.notebook.add(self.stock_frame, text="Stock de Toner")
        
        # Pestaña de Movimientos
        self.movimientos_frame = ttk.Frame(self.notebook)
        self._setup_movimientos_ui()
        self.notebook.add(self.movimientos_frame, text="Movimientos")
        
        # Pestaña de Recargas
        self.recargas_frame = ttk.Frame(self.notebook)
        self._setup_recargas_ui()
        self.notebook.add(self.recargas_frame, text="Recargas")
        
        # Pestaña de Informes
        self.informes_frame = ttk.Frame(self.notebook)
        self._setup_informes_ui()
        self.notebook.add(self.informes_frame, text="Informes")
        
        # Cargar datos iniciales
        #self._cargar_stock()

    # Estos métodos deben estar al mismo nivel que __init__ y _setup_ui
    def _setup_stock_ui(self):
        """Configura la interfaz de la pestaña de stock"""
        self.stock_tree = ttk.Treeview(self.stock_frame, columns=("Marca", "Modelo", "Cantidad"), show="headings")
        self.stock_tree.heading("Marca", text="Marca")
        self.stock_tree.heading("Modelo", text="Modelo")
        self.stock_tree.heading("Cantidad", text="Cantidad")
        self.stock_tree.column("Marca", width=150, anchor=tk.W)
        self.stock_tree.column("Modelo", width=150, anchor=tk.W)
        self.stock_tree.column("Cantidad", width=100, anchor=tk.CENTER)
    
        scrollbar = ttk.Scrollbar(self.stock_frame, orient="vertical", command=self.stock_tree.yview)
        self.stock_tree.configure(yscrollcommand=scrollbar.set)
    
        self.stock_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    def _setup_movimientos_ui(self):
        """Configura la interfaz de la pestaña de movimientos"""
        # Frame de filtros
        filtros_frame = ttk.LabelFrame(self.movimientos_frame, text="Filtros", padding="10")
        filtros_frame.pack(fill=tk.X, padx=5, pady=5)
    
        ttk.Label(filtros_frame, text="Fecha desde:").grid(row=0, column=0, padx=5, pady=5)
        self.mov_fecha_desde = DateEntry(filtros_frame, date_pattern="yyyy-mm-dd")
        self.mov_fecha_desde.grid(row=0, column=1, padx=5, pady=5)
    
        ttk.Label(filtros_frame, text="Fecha hasta:").grid(row=0, column=2, padx=5, pady=5)
        self.mov_fecha_hasta = DateEntry(filtros_frame, date_pattern="yyyy-mm-dd")
        self.mov_fecha_hasta.grid(row=0, column=3, padx=5, pady=5)
    
        ttk.Label(filtros_frame, text="Marca:").grid(row=1, column=0, padx=5, pady=5)
        self.mov_marca_combo = ttk.Combobox(filtros_frame, state="readonly")
        self.mov_marca_combo.grid(row=1, column=1, padx=5, pady=5)
    
        ttk.Label(filtros_frame, text="Modelo:").grid(row=1, column=2, padx=5, pady=5)
        self.mov_modelo_combo = ttk.Combobox(filtros_frame, state="readonly")
        self.mov_modelo_combo.grid(row=1, column=3, padx=5, pady=5)
    
        ttk.Button(filtros_frame, text="Aplicar Filtros", 
                  command=self._aplicar_filtros_movimientos).grid(row=1, column=4, padx=5, pady=5)
    
        # Treeview de movimientos
        self.movimientos_tree = ttk.Treeview(
            self.movimientos_frame, 
            columns=("Fecha", "Tipo", "Marca", "Modelo", "Cantidad", "Responsable", "Sector"), 
            show="headings"
        )
        columns = {
            "Fecha": {"width": 120, "anchor": tk.CENTER},
            "Tipo": {"width": 100, "anchor": tk.CENTER},
            "Marca": {"width": 120, "anchor": tk.W},
            "Modelo": {"width": 120, "anchor": tk.W},
            "Cantidad": {"width": 80, "anchor": tk.CENTER},
            "Responsable": {"width": 150, "anchor": tk.W},
            "Sector": {"width": 150, "anchor": tk.W}
        }
    
        for col, config in columns.items():
            self.movimientos_tree.heading(col, text=col)
            self.movimientos_tree.column(col, **config)
    
        scrollbar = ttk.Scrollbar(self.movimientos_frame, orient="vertical", command=self.movimientos_tree.yview)
        self.movimientos_tree.configure(yscrollcommand=scrollbar.set)
    
        self.movimientos_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    
        # Cargar combobox de marcas
        self._cargar_marcas_movimientos()
        self.mov_marca_combo.bind("<<ComboboxSelected>>", lambda e: self._cargar_modelos_movimientos())

    
    def _setup_recargas_ui(self):
        """Configura la interfaz de la pestaña de recargas"""
        # Frame de filtros
        filtros_frame = ttk.LabelFrame(self.recargas_frame, text="Filtros", padding="10")
        filtros_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(filtros_frame, text="Estado:").grid(row=0, column=0, padx=5, pady=5)
        self.rec_estado_combo = ttk.Combobox(
            filtros_frame, 
            values=["Todos", "Enviado", "Recibido"], 
            state="readonly"
        )
        self.rec_estado_combo.set("Todos")
        self.rec_estado_combo.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Button(filtros_frame, text="Aplicar Filtros", 
                  command=self._aplicar_filtros_recargas).grid(row=0, column=2, padx=5, pady=5)
        
        # Treeview de recargas
        self.recargas_tree = ttk.Treeview(
            self.recargas_frame, 
            columns=("Fecha Envío", "Fecha Recibo", "Marca", "Modelo", "Cantidad", "Empresa", "Estado"), 
            show="headings"
        )
        columns = {
            "Fecha Envío": {"width": 120, "anchor": tk.CENTER},
            "Fecha Recibo": {"width": 120, "anchor": tk.CENTER},
            "Marca": {"width": 120, "anchor": tk.W},
            "Modelo": {"width": 120, "anchor": tk.W},
            "Cantidad": {"width": 80, "anchor": tk.CENTER},
            "Empresa": {"width": 150, "anchor": tk.W},
            "Estado": {"width": 100, "anchor": tk.CENTER}
        }
        
        for col, config in columns.items():
            self.recargas_tree.heading(col, text=col)
            self.recargas_tree.column(col, **config)
        
        scrollbar = ttk.Scrollbar(self.recargas_frame, orient="vertical", command=self.recargas_tree.yview)
        self.recargas_tree.configure(yscrollcommand=scrollbar.set)
        
        self.recargas_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Botón para marcar como recibido
        btn_frame = ttk.Frame(self.recargas_frame)
        btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(btn_frame, text="Marcar como Recibido", 
                  command=self._marcar_recibido).pack(side=tk.LEFT, padx=5)
    
    def _setup_informes_ui(self):
        """Configura la interfaz de la pestaña de informes"""
        # Frame de controles
        controles_frame = ttk.Frame(self.informes_frame)
        controles_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(controles_frame, text="Año:").pack(side=tk.LEFT, padx=5)
        self.inf_anio_combo = ttk.Combobox(
            controles_frame, 
            values=[str(y) for y in range(2020, datetime.now().year + 1)], 
            state="readonly"
        )
        self.inf_anio_combo.set(str(datetime.now().year))
        self.inf_anio_combo.pack(side=tk.LEFT, padx=5)
        
        ttk.Label(controles_frame, text="Mes:").pack(side=tk.LEFT, padx=5)
        self.inf_mes_combo = ttk.Combobox(
            controles_frame, 
            values=["Todos", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", 
                  "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"], 
            state="readonly"
        )
        self.inf_mes_combo.set("Todos")
        self.inf_mes_combo.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(controles_frame, text="Generar Informe de Consumos", 
                    command=self._generar_informe_consumos).pack(side=tk.LEFT, padx=5)
        ttk.Button(controles_frame, text="Generar Informe de Recargas", 
                    command=self._generar_informe_recargas).pack(side=tk.LEFT, padx=5)
    
    def _cargar_stock(self):
        """Muestra el diálogo para cargar stock de toner"""
        dialog = tk.Toplevel()
        dialog.title("Cargar Stock de Toner")
        dialog.transient(self)
        dialog.grab_set()
        """Carga los datos de stock de toner"""
        for item in self.stock_tree.get_children():
            self.stock_tree.delete(item)
            
        try:
            stock = self.controller.db.obtener_stock_toner()
            for item in stock:
                self.stock_tree.insert("", tk.END, values=(item[0], item[1], item[2]))
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar los datos: {str(e)}")
    
    def _cargar_marcas_movimientos(self):
        """Carga las marcas en el combobox de movimientos"""
        try:
            marcas = self.controller.db.obtener_marcas_toner()
            self.mov_marca_combo['values'] = ["Todas"] + [m[1] for m in marcas]
            self.mov_marca_combo.set("Todas")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar las marcas: {str(e)}")
    
    def _cargar_modelos_movimientos(self):
        """Carga los modelos en el combobox de movimientos según la marca seleccionada"""
        marca = self.mov_marca_combo.get()
        if marca == "Todas":
            self.mov_modelo_combo['values'] = ["Todos"]
            self.mov_modelo_combo.set("Todos")
            return
            
        try:
            marcas = self.controller.db.obtener_marcas_toner()
            id_marca = next(m[0] for m in marcas if m[1] == marca)
            
            modelos = self.controller.db.obtener_modelos_toner(id_marca)
            self.mov_modelo_combo['values'] = ["Todos"] + [m[1] for m in modelos]
            self.mov_modelo_combo.set("Todos")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar los modelos: {str(e)}")
    
    def _aplicar_filtros_movimientos(self):
        """Aplica los filtros y carga los movimientos"""
        for item in self.movimientos_tree.get_children():
            self.movimientos_tree.delete(item)
            
        try:
            fecha_desde = self.mov_fecha_desde.get_date().strftime("%Y-%m-%d")
            fecha_hasta = self.mov_fecha_hasta.get_date().strftime("%Y-%m-%d")
            
            marca = self.mov_marca_combo.get()
            modelo = self.mov_modelo_combo.get()
            
            id_marca = None
            if marca != "Todas":
                marcas = self.controller.db.obtener_marcas_toner()
                id_marca = next(m[0] for m in marcas if m[1] == marca)
            
            id_modelo = None
            if modelo != "Todos" and marca != "Todas":
                modelos = self.controller.db.obtener_modelos_toner(id_marca)
                id_modelo = next(m[0] for m in modelos if m[1] == modelo)
            
            movimientos = self.controller.db.obtener_movimientos_toner(fecha_desde, fecha_hasta, id_marca, id_modelo)
            
            for mov in movimientos:
                self.movimientos_tree.insert("", tk.END, values=(
                    mov[6],  # fecha
                    mov[2],  # tipo
                    mov[0],  # marca
                    mov[1],  # modelo
                    mov[3],  # cantidad
                    mov[4],  # responsable
                    mov[5]   # sector
                ))
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar los movimientos: {str(e)}")
    
    def _aplicar_filtros_recargas(self):
        """Aplica los filtros y carga las recargas"""
        for item in self.recargas_tree.get_children():
            self.recargas_tree.delete(item)
            
        try:
            estado = None if self.rec_estado_combo.get() == "Todos" else self.rec_estado_combo.get()
            
            recargas = self.controller.db.obtener_recargas_toner(estado=estado)
             
            for rec in recargas:
                self.recargas_tree.insert("", tk.END, values=(
                    rec[4],  # fecha_envio
                    rec[5] if rec[5] else "N/A",  # fecha_recibo
                    rec[0],  # marca
                    rec[1],  # modelo
                    rec[2],  # cantidad
                    rec[3],  # empresa
                    rec[6]   # estado
                ))
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar las recargas: {str(e)}")
    
    def _marcar_recibido(self):
        """Marca una recarga como recibida"""
        seleccion = self.recargas_tree.selection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione una recarga para marcar como recibida")
            return
            
        item = self.recargas_tree.item(seleccion[0])
        valores = item['values']
        
        if valores[6] == "Recibido":
            messagebox.showwarning("Advertencia", "Esta recarga ya está marcada como recibida")
            return
            
        # Mostrar diálogo para ingresar observaciones
        dialog = tk.Toplevel()
        dialog.title("Marcar como Recibido")
        dialog.transient(self)
        dialog.grab_set()
        
        ttk.Label(dialog, text="Observaciones:").pack(padx=10, pady=5)
        observaciones_text = tk.Text(dialog, height=5, width=40)
        observaciones_text.pack(padx=10, pady=5)
        
        def on_ok():
            observaciones = observaciones_text.get("1.0", tk.END).strip()
            
            # Obtener ID de la recarga (necesitamos modificar la consulta para incluir el ID)
            # Esto es un ejemplo, necesitarías ajustar según tu implementación
            try:
                # En una implementación real, deberías tener el ID de la recarga
                # Aquí asumimos que el primer valor es el ID (necesitarías modificar tu consulta SQL)
                recarga_id = item['values'][0]
                
                self.controller.db.recibir_recarga_toner(
                    recarga_id, 
                    observaciones, 
                    self.controller.current_user['id']
                )
               
                messagebox.showinfo("Éxito", "Recarga marcada como recibida")
                dialog.destroy()
                self._aplicar_filtros_recargas()
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo marcar como recibida: {str(e)}")
        
        ttk.Button(dialog, text="Aceptar", command=on_ok).pack(side=tk.LEFT, padx=10, pady=10)
        ttk.Button(dialog, text="Cancelar", command=dialog.destroy).pack(side=tk.RIGHT, padx=10, pady=10)
    
    def _cargar_stock(self):
        """Muestra el diálogo para cargar stock de toner"""
        dialog = tk.Toplevel()
        dialog.title("Cargar Stock de Toner")
        dialog.transient(self)
        dialog.grab_set()
        
        # Frame principal
        main_frame = ttk.Frame(dialog, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Combobox para marca
        ttk.Label(main_frame, text="Marca:").grid(row=0, column=0, sticky="w", pady=5)
        marca_combo = ttk.Combobox(main_frame, state="readonly")
        marca_combo.grid(row=0, column=1, sticky="ew", padx=5, pady=5)
        
        # Combobox para modelo (se actualiza según la marca seleccionada)
        ttk.Label(main_frame, text="Modelo:").grid(row=1, column=0, sticky="w", pady=5)
        modelo_combo = ttk.Combobox(main_frame, state="readonly")
        modelo_combo.grid(row=1, column=1, sticky="ew", padx=5, pady=5)
        
        # Campo para cantidad
        ttk.Label(main_frame, text="Cantidad:").grid(row=2, column=0, sticky="w", pady=5)
        cantidad_entry = ttk.Entry(main_frame)
        cantidad_entry.grid(row=2, column=1, sticky="ew", padx=5, pady=5)
        
        # Campo para observaciones
        ttk.Label(main_frame, text="Observaciones:").grid(row=3, column=0, sticky="nw", pady=5)
        observaciones_text = tk.Text(main_frame, height=5, width=30)
        observaciones_text.grid(row=3, column=1, sticky="ew", padx=5, pady=5)
        
        # Botones
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=4, column=0, columnspan=2, pady=10)
        
        def on_ok():
            marca = marca_combo.get()
            modelo = modelo_combo.get()
            cantidad = cantidad_entry.get()
            observaciones = observaciones_text.get("1.0", tk.END).strip()
            
            if not marca or not modelo or not cantidad:
                messagebox.showwarning("Advertencia", "Complete todos los campos obligatorios")
                return
                
            try:
                cantidad = int(cantidad)
                if cantidad <= 0:
                    raise ValueError("La cantidad debe ser mayor a cero")
                
                # Obtener ID del modelo
                marcas = self.controller.db.obtener_marcas_toner()
                id_marca = next(m[0] for m in marcas if m[1] == marca)
                
                modelos = self.controller.db.obtener_modelos_toner(id_marca)
                id_modelo = next(m[0] for m in modelos if m[1] == modelo)
                
                # Registrar movimiento de ingreso
                self.controller.db.registrar_movimiento_toner(
                    id_modelo, 
                    'ingreso', 
                    cantidad, 
                    self.controller.current_user['username'], 
                    "Almacén", 
                    None, 
                    observaciones, 
                    self.controller.current_user['id']
                )
                
                messagebox.showinfo("Éxito", "Stock cargado correctamente")
                dialog.destroy()
                self._cargar_stock()
            except ValueError as e:
                messagebox.showerror("Error", f"Cantidad inválida: {str(e)}")
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo cargar el stock: {str(e)}")
        
        ttk.Button(button_frame, text="Aceptar", command=on_ok).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cancelar", command=dialog.destroy).pack(side=tk.RIGHT, padx=5)
        
        # Cargar marcas
        try:
            marcas = self.controller.db.obtener_marcas_toner()
            marca_combo['values'] = [m[1] for m in marcas]
            if marcas:
                marca_combo.current(0)
                # Cargar modelos para la primera marca
                modelos = self.controller.db.obtener_modelos_toner(marcas[0][0])
                modelo_combo['values'] = [m[1] for m in modelos]
                if modelos:
                    modelo_combo.current(0)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar las marcas: {str(e)}")
        
        # Configurar evento para cambiar modelos cuando cambia la marca
        marca_combo.bind("<<ComboboxSelected>>", lambda e: self._actualizar_modelos_combo(marca_combo, modelo_combo))
    
    def _actualizar_modelos_combo(self, marca_combo, modelo_combo):
        """Actualiza los modelos cuando se selecciona una marca"""
        marca = marca_combo.get()
        if not marca:
            modelo_combo['values'] = []
            return
            
        try:
            marcas = self.controller.db.obtener_marcas_toner()
            id_marca = next(m[0] for m in marcas if m[1] == marca)
            
            modelos = self.controller.db.obtener_modelos_toner(id_marca)
            modelo_combo['values'] = [m[1] for m in modelos]
            modelo_combo.set('')
            #if modelos:
                #modelo_combo.current(0)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar los modelos: {str(e)}")
            modelo_combo['values'] = []
    
    def _generar_informe_consumos(self):
        """Genera informe de consumos de toner"""
        try:
            año = int(self.inf_anio_combo.get())
            mes = None if self.inf_mes_combo.get() == "Todos" else self._get_month_number(self.inf_mes_combo.get())
            
            # Obtener datos
            movimientos = self.controller.db.obtener_movimientos_toner(
                f"{año}-01-01", 
                f"{año}-12-31", 
                None, 
                None
            )
            
            if mes:
                movimientos = [m for m in movimientos if datetime.strptime(m[6], "%Y-%m-%d %H:%M:%S").month == mes]
            
            # Procesar datos para el informe
            data = {
                "titulo": f"Informe de Consumos de Toner - {self.inf_anio_combo.get()}",
                "subtitulo": f"Mes: {self.inf_mes_combo.get()}" if mes else "Anual",
                "encabezados": ["Fecha", "Marca", "Modelo", "Cantidad", "Responsable", "Sector"],
                "datos": [[m[6][:10], m[0], m[1], m[3], m[4], m[5]] for m in movimientos if m[2] == 'retiro'],
                "resumen": self._generar_resumen_consumos(movimientos)
            }
           
            # Generar nombre de archivo
            fecha = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Informe_Consumos_Toner_{año}_{self.inf_mes_combo.get()}_{fecha}.docx"
            
            if ExportManager.export_informe_toner(data, filename):
                messagebox.showinfo("Éxito", f"Informe generado:\n{filename}")
                if messagebox.askyesno("Abrir", "¿Desea abrir el documento para imprimir?"):
                    webbrowser.open(filename)
            else:
                messagebox.showerror("Error", "No se pudo generar el informe")
                
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo generar el informe: {str(e)}")
     
    def _generar_informe_recargas(self):
        """Genera informe de recargas de toner"""
        try:
            año = int(self.inf_anio_combo.get())
            mes = None if self.inf_mes_combo.get() == "Todos" else self._get_month_number(self.inf_mes_combo.get())
            
            # Obtener datos
            recargas = self.controller.db.obtener_recargas_toner(año, mes)
            
            # Procesar datos para el informe
            data = {
                "titulo": f"Informe de Recargas de Toner - {self.inf_anio_combo.get()}",
                "subtitulo": f"Mes: {self.inf_mes_combo.get()}" if mes else "Anual",
                "encabezados": ["Fecha Envío", "Fecha Recibo", "Marca", "Modelo", "Cantidad", "Empresa", "Estado"],
                "datos": [[r[4], r[5] if r[5] else "N/A", r[0], r[1], r[2], r[3], r[6]] for r in recargas],
                "resumen": self._generar_resumen_recargas(recargas)
            }
            
            # Generar nombre de archivo
            fecha = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Informe_Recargas_Toner_{año}_{self.inf_mes_combo.get()}_{fecha}.docx"
          
            if ExportManager.export_informe_toner(data, filename):
                messagebox.showinfo("Éxito", f"Informe generado:\n{filename}")
                if messagebox.askyesno("Abrir", "¿Desea abrir el documento para imprimir?"):
                    webbrowser.open(filename)
            else:
                messagebox.showerror("Error", "No se pudo generar el informe")
                
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo generar el informe: {str(e)}")
    
    def _generar_resumen_consumos(self, movimientos):
        """Genera un resumen de consumos por marca y modelo"""
        resumen = {}
        
        for mov in movimientos:
            if mov[2] != 'retiro':
               continue
                
            key = (mov[0], mov[1])  # (marca, modelo)
            if key not in resumen:
               resumen[key] = 0
            resumen[key] += mov[3]  # cantidad
        
        # Convertir a lista ordenada
        return sorted([(k[0], k[1], v) for k, v in resumen.items()], key=lambda x: x[2], reverse=True)
    
    def _generar_resumen_recargas(self, recargas):
        """Genera un resumen de recargas por empresa y estado"""
        resumen = {}
        
        for rec in recargas:
            key = (rec[3], rec[6])  # (empresa, estado)
            if key not in resumen:
                resumen[key] = 0
            resumen[key] += rec[2]  # cantidad
        
        # Convertir a lista ordenada
        return sorted([(k[0], k[1], v) for k, v in resumen.items()], key=lambda x: x[2], reverse=True)
    
    def _get_month_number(self, month_name):
        """Convierte nombre de mes a número"""
        meses = {
            "Enero": 1, "Febrero": 2, "Marzo": 3, "Abril": 4, "Mayo": 5, "Junio": 6,
            "Julio": 7, "Agosto": 8, "Septiembre": 9, "Octubre": 10, "Noviembre": 11, "Diciembre": 12
        }
        return meses.get(month_name, 0)

    

class MarcasTonerView(ttk.Frame):
    """Vista para gestionar marcas de toner"""
    
    def initialize(self, *args, **kwargs):
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
        self._cargar_marcas()
    
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self._setup_ui()
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        
        # Frame superior con botones
        button_frame = ttk.Frame(self)
        button_frame.grid(row=0, column=0, sticky="ew", pady=5)
        
        ttk.Button(button_frame, text="Agregar Marca", 
                 command=self._agregar_marca).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Modificar Marca", 
                 command=self._modificar_marca).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Eliminar Marca", 
                 command=self._eliminar_marca).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Volver", 
                 command=lambda: self.controller.mostrar_vista("TonerView")).pack(side=tk.RIGHT, padx=5)
        
        # Treeview para mostrar las marcas
        self.treeview = ttk.Treeview(self, columns=("ID", "Nombre"), show="headings")
        self.treeview.heading("ID", text="ID")
        self.treeview.heading("Nombre", text="Nombre")
        self.treeview.column("ID", width=50, anchor=tk.CENTER)
        self.treeview.column("Nombre", width=200, anchor=tk.W)
        
        scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.treeview.yview)
        self.treeview.configure(yscrollcommand=scrollbar.set)
        
        self.treeview.grid(row=1, column=0, sticky="nsew")
        scrollbar.grid(row=1, column=1, sticky="ns")
        
        # Cargar datos iniciales
        self._cargar_marcas()
    
    def _cargar_marcas(self):
        """Carga las marcas desde la base de datos"""
        for item in self.treeview.get_children():
            self.treeview.delete(item)
            
        try:
            marcas = self.controller.db.obtener_marcas_toner()
            for marca in marcas:
                self.treeview.insert("", tk.END, values=marca)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar las marcas: {str(e)}")
    



    
    def _agregar_marca(self):
        """Agrega una nueva marca de toner"""
        nombre = simpledialog.askstring("Agregar Marca", "Ingrese el nombre de la nueva marca de toner:")
        if nombre:
            try:
                self.controller.db.agregar_marca_toner(nombre, self.controller.current_user['id'])
                self._cargar_marcas()
                messagebox.showinfo("Éxito", "Marca agregada correctamente")
            except ValueError as e:
                messagebox.showerror("Error", str(e))
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo agregar la marca: {str(e)}")
    
    def _modificar_marca(self):
        """Modifica una marca existente"""
        seleccion = self.treeview.selection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione una marca para modificar")
            return
            
        item = self.treeview.item(seleccion[0])
        marca_id, nombre_actual = item['values']
        
        nuevo_nombre = simpledialog.askstring("Modificar Marca", "Ingrese el nuevo nombre:", 
                                            initialvalue=nombre_actual)
        if nuevo_nombre and nuevo_nombre != nombre_actual:
            try:
                # Actualizar en la base de datos
                self.controller.db.cursor.execute(
                    "UPDATE marcas_toner SET nombre = ? WHERE id = ?",
                    (nuevo_nombre, marca_id))
                
                self.controller.db.registrar_auditoria(
                    self.controller.current_user['id'], 
                    'MODIFICAR_MARCA_TONER', 
                    'marcas_toner', 
                    marca_id,
                    f"Modificación de marca: {nombre_actual} -> {nuevo_nombre}"
                )
                
                self.controller.db.conn.commit()
                self._cargar_marcas()
                messagebox.showinfo("Éxito", "Marca modificada correctamente")
            except Exception as e:
                self.controller.db.conn.rollback()
                messagebox.showerror("Error", f"No se pudo modificar la marca: {str(e)}")
    
    def _eliminar_marca(self):
        """Elimina una marca existente"""
        seleccion = self.treeview.selection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione una marca para eliminar")
            return
            
        item = self.treeview.item(seleccion[0])
        marca_id, nombre = item['values']
        
        # Verificar si tiene modelos asociados
        try:
            self.controller.db.cursor.execute(
                "SELECT COUNT(*) FROM modelos_toner WHERE id_marca = ?", 
                (marca_id,))
            if self.controller.db.cursor.fetchone()[0] > 0:
                messagebox.showerror("Error", 
                                   "No se puede eliminar: la marca tiene modelos asociados")
                return
                
            if messagebox.askyesno("Confirmar", f"¿Eliminar la marca '{nombre}'?"):
                self.controller.db.cursor.execute(
                    "DELETE FROM marcas_toner WHERE id = ?", 
                    (marca_id,))
                
                self.controller.db.registrar_auditoria(
                    self.controller.current_user['id'], 
                    'ELIMINAR_MARCA_TONER', 
                    'marcas_toner', 
                    marca_id,
                    f"Marca eliminada: {nombre}"
                )
                
                self.controller.db.conn.commit()
                self._cargar_marcas()
                messagebox.showinfo("Éxito", "Marca eliminada correctamente")
        except Exception as e:
            self.controller.db.conn.rollback()
            messagebox.showerror("Error", f"No se pudo eliminar la marca: {str(e)}")
          

class ModelosTonerView(ttk.Frame):
    """Vista para gestionar modelos de toner"""
    
    def initialize(self, *args, **kwargs):
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
    
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self._setup_ui()
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        
        # Frame de controles
        control_frame = ttk.Frame(self)
        control_frame.grid(row=0, column=0, sticky="ew", pady=5)
        
        # Combobox para seleccionar marca
        ttk.Label(control_frame, text="Marca:").pack(side=tk.LEFT, padx=5)
        self.marca_combobox = ttk.Combobox(control_frame, state="readonly")
        self.marca_combobox.pack(side=tk.LEFT, padx=5)
        self.marca_combobox.bind("<<ComboboxSelected>>", lambda e: self._cargar_modelos())
        
        # Botones
        button_frame = ttk.Frame(control_frame)
        button_frame.pack(side=tk.RIGHT, padx=5)
        
        ttk.Button(button_frame, text="Agregar Modelo", 
                 command=self._agregar_modelo).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Modificar Modelo", 
                 command=self._modificar_modelo).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Eliminar Modelo", 
                 command=self._eliminar_modelo).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Volver", 
                 command=lambda: self.controller.mostrar_vista("TonerView")).pack(side=tk.LEFT, padx=5)
        
        # Treeview para mostrar los modelos
        self.treeview = ttk.Treeview(self, columns=("ID", "Marca", "Nombre"), show="headings")
        self.treeview.heading("ID", text="ID")
        self.treeview.heading("Marca", text="Marca")
        self.treeview.heading("Nombre", text="Modelo")
        self.treeview.column("ID", width=50, anchor=tk.CENTER)
        self.treeview.column("Marca", width=150, anchor=tk.W)
        self.treeview.column("Nombre", width=150, anchor=tk.W)
        
        scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.treeview.yview)
        self.treeview.configure(yscrollcommand=scrollbar.set)
        
        self.treeview.grid(row=1, column=0, sticky="nsew")
        scrollbar.grid(row=1, column=1, sticky="ns")

    def _cargar_marcas(self):
        """Carga las marcas en el combobox"""
        try:
            marcas = self.controller.db.obtener_marcas_toner()
            self.marca_combobox['values'] = [m[1] for m in marcas]
            if marcas:
                self.marca_combobox.current(0)
                self._cargar_modelos()
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar las marcas: {str(e)}")

    def _cargar_modelos(self):
        """Carga los modelos según la marca seleccionada"""
        marca = self.marca_combo.get()
        if not marca:
            return
        
        try:
            # Obtener ID de la marca seleccionada
            marcas = self.controller.db.obtener_marcas_toner()
            marca_id = next(m[0] for m in marcas if m[1] == marca)
        
            # Obtener y mostrar modelos
            modelos = self.controller.db.obtener_modelos_toner(marca_id)
            self.modelo_combo['values'] = [m[1] for m in modelos]
            if modelos:
                self.modelo_combo.current(0)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar los modelos: {str(e)}")


    def _agregar_modelo(self):
        """Agrega un nuevo modelo a la marca seleccionada"""
        marca = self.marca_combobox.get()
        if not marca:
            messagebox.showwarning("Advertencia", "Seleccione una marca primero")
            return
            
        nombre = simpledialog.askstring("Agregar Modelo", "Ingrese el nombre del nuevo modelo:")
        if nombre:
            try:
                # Obtener ID de la marca
                marcas = self.controller.db.obtener_marcas_toner()
                marca_id = next(m[0] for m in marcas if m[1] == marca)
                
                # Agregar el modelo
                self.controller.db.agregar_modelo_toner(marca_id, nombre, self.controller.current_user['id'])
                self._cargar_modelos()
                messagebox.showinfo("Éxito", "Modelo agregado correctamente")
            except ValueError as e:
                messagebox.showerror("Error", str(e))
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo agregar el modelo: {str(e)}")


            
         
    def _modificar_modelo(self):
        """Modifica un modelo existente"""
        seleccion = self.treeview.selection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione un modelo para modificar")
            return
            
        item = self.treeview.item(seleccion[0])
        modelo_id, marca_nombre, nombre_actual = item['values']
        
        nuevo_nombre = simpledialog.askstring("Modificar Modelo", "Ingrese el nuevo nombre:", 
                                            initialvalue=nombre_actual)
        if nuevo_nombre and nuevo_nombre != nombre_actual:
            try:
                # Obtener ID de la marca
                marcas = self.controller.db.obtener_marcas_toner()
                marca_id = next(m[0] for m in marcas if m[1] == marca_nombre)
                
                # Verificar que no exista ya
                modelos = self.controller.db.obtener_modelos_toner(marca_id)
                if any(nuevo_nombre.lower() == m[1].lower() for m in modelos):
                    raise ValueError("Ya existe un modelo con ese nombre para esta marca")
                
                # Actualizar
                self.controller.db.cursor.execute(
                    "UPDATE modelos_toner SET nombre = ? WHERE id = ?",
                    (nuevo_nombre, modelo_id))
                
                self.controller.db.registrar_auditoria(
                    self.controller.current_user['id'], 
                    'MODIFICAR_MODELO_TONER', 
                    'modelos_toner', 
                    modelo_id,
                    f"Modificación de modelo: {nombre_actual} -> {nuevo_nombre}"
                )
                
                self.controller.db.conn.commit()
                self._cargar_modelos()
                messagebox.showinfo("Éxito", "Modelo modificado correctamente")
            except ValueError as e:
                messagebox.showerror("Error", str(e))
            except Exception as e:
                self.controller.db.conn.rollback()
                messagebox.showerror("Error", f"No se pudo modificar el modelo: {str(e)}")

    def _eliminar_modelo(self):
        """Elimina un modelo existente"""
        seleccion = self.treeview.selection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione un modelo para eliminar")
            return
            
        item = self.treeview.item(seleccion[0])
        modelo_id, marca_nombre, nombre = item['values']
        
        # Verificar si tiene stock asociado
        try:
            self.controller.db.cursor.execute(
                "SELECT COUNT(*) FROM stock_toner WHERE id_modelo = ?", 
                (modelo_id,))
            if self.controller.db.cursor.fetchone()[0] > 0:
                messagebox.showerror("Error", 
                                   "No se puede eliminar: el modelo tiene stock asociado")
                return
                
            if messagebox.askyesno("Confirmar", f"¿Eliminar el modelo '{nombre}'?"):
                self.controller.db.cursor.execute(
                    "DELETE FROM modelos_toner WHERE id = ?", 
                    (modelo_id,))
                
                self.controller.db.registrar_auditoria(
                    self.controller.current_user['id'], 
                    'ELIMINAR_MODELO_TONER', 
                    'modelos_toner', 
                    modelo_id,
                    f"Modelo eliminado: {nombre} (Marca: {marca_nombre})"
                )
                
                self.controller.db.conn.commit()
                self._cargar_modelos()
                messagebox.showinfo("Éxito", "Modelo eliminado correctamente")
        except Exception as e:
            self.controller.db.conn.rollback()
            messagebox.showerror("Error", f"No se pudo eliminar el modelo: {str(e)}")
            
class RetiroTonerView(ttk.Frame):
    """Vista para registrar retiro de toner"""
    
    def initialize(self, *args, **kwargs):
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
        self._cargar_marcas()
    
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self._setup_ui()
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        
        # Frame principal
        main_frame = ttk.Frame(self, padding="10")
        main_frame.grid(row=0, column=0, sticky="nsew")
        
        # Combobox para marca
        ttk.Label(main_frame, text="Marca:").grid(row=0, column=0, sticky="w", pady=5)
        self.marca_combo = ttk.Combobox(main_frame, state="readonly")
        self.marca_combo.grid(row=0, column=1, sticky="ew", padx=5, pady=5)
        self.marca_combo.bind("<<ComboboxSelected>>", lambda e: self._cargar_modelos())
        
        # Combobox para modelo
        ttk.Label(main_frame, text="Modelo:").grid(row=1, column=0, sticky="w", pady=5)
        self.modelo_combo = ttk.Combobox(main_frame, state="readonly")
        self.modelo_combo.grid(row=1, column=1, sticky="ew", padx=5, pady=5)
        
        # Campo para cantidad
        ttk.Label(main_frame, text="Cantidad:").grid(row=2, column=0, sticky="w", pady=5)
        self.cantidad_entry = ttk.Entry(main_frame)
        self.cantidad_entry.grid(row=2, column=1, sticky="ew", padx=5, pady=5)
        
        # Campo para responsable
        ttk.Label(main_frame, text="Responsable:").grid(row=3, column=0, sticky="w", pady=5)
        self.responsable_entry = ttk.Entry(main_frame)
        self.responsable_entry.grid(row=3, column=1, sticky="ew", padx=5, pady=5)
        
        # Campo para sector
        ttk.Label(main_frame, text="Sector:").grid(row=4, column=0, sticky="w", pady=5)
        self.sector_entry = ttk.Entry(main_frame)
        self.sector_entry.grid(row=4, column=1, sticky="ew", padx=5, pady=5)
        
        # Campo para observaciones
        ttk.Label(main_frame, text="Observaciones:").grid(row=5, column=0, sticky="nw", pady=5)
        self.observaciones_text = tk.Text(main_frame, height=5, width=30)
        self.observaciones_text.grid(row=5, column=1, sticky="ew", padx=5, pady=5)
        
        # Botones
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=6, column=0, columnspan=2, pady=10)
        
        ttk.Button(button_frame, text="Registrar Retiro", 
                 command=self._registrar_retiro).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cancelar", 
                 command=lambda: self.controller.mostrar_vista("TonerView")).pack(side=tk.RIGHT, padx=5)
    
    def _cargar_marcas(self):
        """Carga las marcas en el combobox"""
        try:
            marcas = self.controller.db.obtener_marcas_toner()
            self.marca_combo['values'] = [m[1] for m in marcas]
            if marcas:
                self.marca_combo.current(0)
                self._cargar_modelos()
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar las marcas: {str(e)}")
    
    def _cargar_modelos(self):
        """Carga los modelos según la marca seleccionada"""
        marca = self.marca_combo.get()
        if not marca:
            return
            
        try:
            # Obtener ID de la marca seleccionada
            marcas = self.controller.db.obtener_marcas_toner()
            marca_id = next(m[0] for m in marcas if m[1] == marca)
            
            # Obtener y mostrar modelos
            modelos = self.controller.db.obtener_modelos_toner(marca_id)
            self.modelo_combo['values'] = [m[1] for m in modelos]
            if modelos:
                self.modelo_combo.current(0)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar los modelos: {str(e)}")
    
    def _registrar_retiro(self):
        """Registra el retiro de toner"""
        marca = self.marca_combo.get()
        modelo = self.modelo_combo.get()
        cantidad = self.cantidad_entry.get()
        responsable = self.responsable_entry.get()
        sector = self.sector_entry.get()
        observaciones = self.observaciones_text.get("1.0", tk.END).strip()
        
        if not marca or not modelo or not cantidad or not responsable or not sector:
            messagebox.showwarning("Advertencia", "Complete todos los campos obligatorios")
            return
            
        try:
            cantidad = int(cantidad)
            if cantidad <= 0:
                raise ValueError("La cantidad debe ser mayor a cero")
            
            # Obtener ID del modelo
            marcas = self.controller.db.obtener_marcas_toner()
            id_marca = next(m[0] for m in marcas if m[1] == marca)
            
            modelos = self.controller.db.obtener_modelos_toner(id_marca)
            id_modelo = next(m[0] for m in modelos if m[1] == modelo)
            
            # Verificar stock disponible
            stock = self.controller.db.obtener_stock_toner_por_modelo(id_modelo)
            if stock < cantidad:
                raise ValueError(f"No hay suficiente stock. Disponible: {stock}")
            
            # Registrar movimiento de retiro
            self.controller.db.registrar_movimiento_toner(
                id_modelo, 
                'retiro', 
                cantidad, 
                responsable, 
                sector, 
                observaciones, 
                self.controller.current_user['id']
            )
            
            messagebox.showinfo("Éxito", "Retiro registrado correctamente")
            self.controller.mostrar_vista("TonerView")
        except ValueError as e:
            messagebox.showerror("Error", f"Dato inválido: {str(e)}")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo registrar el retiro: {str(e)}")

            

class RecargaTonerView(ttk.Frame):
    """Vista para gestionar recargas de toner"""
    
    def initialize(self, *args, **kwargs):
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
        self._cargar_marcas()
        self._cargar_empresas()
    
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self._setup_ui()
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        
        # Frame principal
        main_frame = ttk.Frame(self, padding="10")
        main_frame.grid(row=0, column=0, sticky="nsew")
        
        # Combobox para marca
        ttk.Label(main_frame, text="Marca:").grid(row=0, column=0, sticky="w", pady=5)
        self.marca_combo = ttk.Combobox(main_frame, state="readonly")
        self.marca_combo.grid(row=0, column=1, sticky="ew", padx=5, pady=5)
        self.marca_combo.bind("<<ComboboxSelected>>", lambda e: self._cargar_modelos())
        
        # Combobox para modelo
        ttk.Label(main_frame, text="Modelo:").grid(row=1, column=0, sticky="w", pady=5)
        self.modelo_combo = ttk.Combobox(main_frame, state="readonly")
        self.modelo_combo.grid(row=1, column=1, sticky="ew", padx=5, pady=5)
        
        # Campo para cantidad
        ttk.Label(main_frame, text="Cantidad:").grid(row=2, column=0, sticky="w", pady=5)
        self.cantidad_entry = ttk.Entry(main_frame)
        self.cantidad_entry.grid(row=2, column=1, sticky="ew", padx=5, pady=5)
        
        # Combobox para empresa
        ttk.Label(main_frame, text="Empresa:").grid(row=3, column=0, sticky="w", pady=5)
        self.empresa_combo = ttk.Combobox(main_frame, state="readonly")
        self.empresa_combo.grid(row=3, column=1, sticky="ew", padx=5, pady=5)
        
        # Campo para observaciones
        ttk.Label(main_frame, text="Observaciones:").grid(row=4, column=0, sticky="nw", pady=5)
        self.observaciones_text = tk.Text(main_frame, height=5, width=30)
        self.observaciones_text.grid(row=4, column=1, sticky="ew", padx=5, pady=5)
        
        # Botones
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=5, column=0, columnspan=2, pady=10)
        
        ttk.Button(button_frame, text="Registrar Envío", 
                 command=self._registrar_envio).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cancelar", 
                 command=lambda: self.controller.mostrar_vista("TonerView")).pack(side=tk.RIGHT, padx=5)
    
    def _cargar_marcas(self):
        """Carga las marcas en el combobox"""
        try:
            marcas = self.controller.db.obtener_marcas_toner()
            self.marca_combo['values'] = [m[1] for m in marcas]
            if marcas:
                self.marca_combo.current(0)
                self._cargar_modelos()
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar las marcas: {str(e)}")
    
    def _cargar_modelos(self):
        """Carga los modelos según la marca seleccionada"""
        marca = self.marca_combo.get()
        if not marca:
            return
            
        try:
            # Obtener ID de la marca seleccionada
            marcas = self.controller.db.obtener_marcas_toner()
            marca_id = next(m[0] for m in marcas if m[1] == marca)
            
            # Obtener y mostrar modelos
            modelos = self.controller.db.obtener_modelos_toner(marca_id)
            self.modelo_combo['values'] = [m[1] for m in modelos]
            if modelos:
                self.modelo_combo.current(0)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar los modelos: {str(e)}")
    
    def _cargar_empresas(self):
        """Carga las empresas de recarga en el combobox"""
        try:
            empresas = self.controller.db.obtener_empresas_recarga()
            self.empresa_combo['values'] = [e[1] for e in empresas]
            if empresas:
                self.empresa_combo.current(0)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar las empresas: {str(e)}")
    
    def _registrar_envio(self):
        """Registra el envío de toner para recarga"""
        marca = self.marca_combo.get()
        modelo = self.modelo_combo.get()
        cantidad = self.cantidad_entry.get()
        empresa = self.empresa_combo.get()
        observaciones = self.observaciones_text.get("1.0", tk.END).strip()
        
        if not marca or not modelo or not cantidad or not empresa:
            messagebox.showwarning("Advertencia", "Complete todos los campos obligatorios")
            return
            
        try:
            cantidad = int(cantidad)
            if cantidad <= 0:
                raise ValueError("La cantidad debe ser mayor a cero")
            
            # Obtener ID del modelo
            marcas = self.controller.db.obtener_marcas_toner()
            id_marca = next(m[0] for m in marcas if m[1] == marca)
            
            modelos = self.controller.db.obtener_modelos_toner(id_marca)
            id_modelo = next(m[0] for m in modelos if m[1] == modelo)
            
            # Obtener ID de la empresa
            empresas = self.controller.db.obtener_empresas_recarga()
            id_empresa = next(e[0] for e in empresas if e[1] == empresa)
            
            # Registrar recarga
            self.controller.db.registrar_recarga_toner(
                id_modelo, 
                cantidad, 
                id_empresa, 
                observaciones, 
                self.controller.current_user['id']
            )
            
            messagebox.showinfo("Éxito", "Envío para recarga registrado correctamente")
            self.controller.mostrar_vista("TonerView")
        except ValueError as e:
            messagebox.showerror("Error", f"Dato inválido: {str(e)}")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo registrar el envío: {str(e)}")

class InformesTonerView(ttk.Frame):
    """Vista para generar informes de toner"""
    
    def initialize(self, *args, **kwargs):
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
        self._cargar_anios()
    
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self._setup_ui()
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        
        # Frame de controles
        control_frame = ttk.Frame(self)
        control_frame.grid(row=0, column=0, sticky="ew", pady=5)
        
        ttk.Label(control_frame, text="Año:").pack(side=tk.LEFT, padx=5)
        self.anio_combo = ttk.Combobox(control_frame, state="readonly")
        self.anio_combo.pack(side=tk.LEFT, padx=5)
        
        ttk.Label(control_frame, text="Mes:").pack(side=tk.LEFT, padx=5)
        self.mes_combo = ttk.Combobox(
            control_frame, 
            values=["Todos", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", 
                  "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"], 
            state="readonly"
        )
        self.mes_combo.set("Todos")
        self.mes_combo.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(control_frame, text="Generar Informe de Consumos", 
                  command=self._generar_informe_consumos).pack(side=tk.LEFT, padx=5)
        ttk.Button(control_frame, text="Generar Informe de Recargas", 
                  command=self._generar_informe_recargas).pack(side=tk.LEFT, padx=5)
        ttk.Button(control_frame, text="Volver", 
                  command=lambda: self.controller.mostrar_vista("TonerView")).pack(side=tk.RIGHT, padx=5)
    
    def _cargar_anios(self):
        """Carga los años disponibles en el combobox"""
        try:
            anios = self.controller.db.obtener_anios_movimientos_toner()
            self.anio_combo['values'] = [str(a) for a in anios]
            if anios:
                self.anio_combo.set(str(datetime.now().year))
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar los años: {str(e)}")
    
    def _generar_informe_consumos(self):
        """Genera informe de consumos de toner"""
        try:
            año = int(self.anio_combo.get())
            mes = None if self.mes_combo.get() == "Todos" else self._get_month_number(self.mes_combo.get())
            
            # Obtener datos
            movimientos = self.controller.db.obtener_movimientos_toner_para_informe(año, mes)
            
            # Procesar datos para el informe
            data = {
                "titulo": f"Informe de Consumos de Toner - {self.anio_combo.get()}",
                "subtitulo": f"Mes: {self.mes_combo.get()}" if mes else "Anual",
                "encabezados": ["Fecha", "Marca", "Modelo", "Cantidad", "Responsable", "Sector"],
                "datos": [[m[0], m[1], m[2], m[3], m[4], m[5]] for m in movimientos],
                "resumen": self._generar_resumen_consumos(movimientos)
            }
            
            # Generar nombre de archivo
            fecha = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Informe_Consumos_Toner_{año}_{self.mes_combo.get()}_{fecha}.docx"
            
            if ExportManager.export_informe_toner(data, filename):
                messagebox.showinfo("Éxito", f"Informe generado:\n{filename}")
                if messagebox.askyesno("Abrir", "¿Desea abrir el documento para imprimir?"):
                    webbrowser.open(filename)
            else:
                messagebox.showerror("Error", "No se pudo generar el informe")
                
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo generar el informe: {str(e)}")
    
    def _generar_informe_recargas(self):
        """Genera informe de recargas de toner"""
        try:
            año = int(self.anio_combo.get())
            mes = None if self.mes_combo.get() == "Todos" else self._get_month_number(self.mes_combo.get())
            
            # Obtener datos
            recargas = self.controller.db.obtener_recargas_toner_para_informe(año, mes)
            
            # Procesar datos para el informe
            data = {
                "titulo": f"Informe de Recargas de Toner - {self.anio_combo.get()}",
                "subtitulo": f"Mes: {self.mes_combo.get()}" if mes else "Anual",
                "encabezados": ["Fecha Envío", "Fecha Recibo", "Marca", "Modelo", "Cantidad", "Empresa", "Estado"],
                "datos": [[r[0], r[1] if r[1] else "N/A", r[2], r[3], r[4], r[5], r[6]] for r in recargas],
                "resumen": self._generar_resumen_recargas(recargas)
            }
            
            # Generar nombre de archivo
            fecha = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Informe_Recargas_Toner_{año}_{self.mes_combo.get()}_{fecha}.docx"
            
            if ExportManager.export_informe_toner(data, filename):
                messagebox.showinfo("Éxito", f"Informe generado:\n{filename}")
                if messagebox.askyesno("Abrir", "¿Desea abrir el documento para imprimir?"):
                    webbrowser.open(filename)
            else:
                messagebox.showerror("Error", "No se pudo generar el informe")
                
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo generar el informe: {str(e)}")
    
    def _generar_resumen_consumos(self, movimientos):
        """Genera un resumen de consumos por marca y modelo"""
        resumen = {}
        
        for mov in movimientos:
            key = (mov[1], mov[2])  # (marca, modelo)
            if key not in resumen:
                resumen[key] = 0
            resumen[key] += mov[3]  # cantidad
        
        # Convertir a lista ordenada
        return sorted([(k[0], k[1], v) for k, v in resumen.items()], key=lambda x: x[2], reverse=True)
    
    def _generar_resumen_recargas(self, recargas):
        """Genera un resumen de recargas por empresa y estado"""
        resumen = {}
        
        for rec in recargas:
            key = (rec[5], rec[6])  # (empresa, estado)
            if key not in resumen:
                resumen[key] = 0
            resumen[key] += rec[4]  # cantidad
        
        # Convertir a lista ordenada
        return sorted([(k[0], k[1], v) for k, v in resumen.items()], key=lambda x: x[2], reverse=True)
    
    def _get_month_number(self, month_name):
        """Convierte nombre de mes a número"""
        meses = {
            "Enero": 1, "Febrero": 2, "Marzo": 3, "Abril": 4, "Mayo": 5, "Junio": 6,
            "Julio": 7, "Agosto": 8, "Septiembre": 9, "Octubre": 10, "Noviembre": 11, "Diciembre": 12
        }
        return meses.get(month_name, 0)

class MarcasTonerView(ttk.Frame):
    """Vista para gestionar marcas de equipos"""


    def initialize(self, *args, **kwargs):
        # Verificar autenticación primero
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
    
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self._setup_ui()
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        
        # Frame superior con botones
        button_frame = ttk.Frame(self)
        button_frame.grid(row=0, column=0, sticky="ew", pady=5)
        
        ttk.Button(button_frame, text="Agregar Marca", 
                  command=self._agregar_marca).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Modificar Marca", 
                  command=self._modificar_marca).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Eliminar Marca", 
                  command=self._eliminar_marca).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Volver", 
                  command=lambda: self.controller.mostrar_vista("MainView")).pack(side=tk.RIGHT, padx=5)
        
        # Treeview para mostrar las marcas
        self.treeview = ttk.Treeview(self, columns=("ID", "Nombre"), show="headings")
        self.treeview.heading("ID", text="ID")
        self.treeview.heading("Nombre", text="Nombre")
        self.treeview.column("ID", width=50, anchor=tk.CENTER)
        self.treeview.column("Nombre", width=200, anchor=tk.W)
        
        scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.treeview.yview)
        self.treeview.configure(yscrollcommand=scrollbar.set)
        
        self.treeview.grid(row=1, column=0, sticky="nsew")
        scrollbar.grid(row=1, column=1, sticky="ns")
        
        # Cargar datos iniciales
        self._cargar_marcas()
    
    def _cargar_marcas(self):
        """Carga las marcas desde la base de datos"""
        for item in self.treeview.get_children():
            self.treeview.delete(item)
            
        marcas = self.controller.db.obtener_marcas()
        for marca in marcas:
            self.treeview.insert("", tk.END, values=marca)
    
    def _agregar_marca(self):
        """Agrega una nueva marca"""
        nombre = simpledialog.askstring("Agregar Marca", "Ingrese el nombre de la nueva marca:")
        if nombre:
            try:
                self.controller.db.agregar_marca(nombre, self.controller.current_user['id'])
                self._cargar_marcas()
                messagebox.showinfo("Éxito", "Marca agregada correctamente")
            except ValueError as e:
                messagebox.showerror("Error", str(e))
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo agregar la marca: {str(e)}")
    
    def _modificar_marca(self):
        """Modifica una marca existente"""
        seleccion = self.treeview.selection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione una marca para modificar")
            return
            
        item = self.treeview.item(seleccion[0])
        marca_id, nombre_actual = item['values']
        
        nuevo_nombre = simpledialog.askstring("Modificar Marca", "Ingrese el nuevo nombre:", 
                                            initialvalue=nombre_actual)
        if nuevo_nombre and nuevo_nombre != nombre_actual:
            try:
                # Primero verificamos que no exista ya una marca con ese nombre
                marcas = self.controller.db.obtener_marcas()
                if any(nuevo_nombre.lower() == m[1].lower() for m in marcas):
                    raise ValueError("Ya existe una marca con ese nombre")
                
                # Actualizamos la marca
                self.controller.db.cursor.execute(
                    "UPDATE marcas SET nombre = ? WHERE id = ?",
                    (nuevo_nombre, marca_id))
                
                self.controller.db.registrar_auditoria(
                    self.controller.current_user['id'], 'MODIFICAR_MARCA', 'marcas', marca_id,
                    f"Modificación de marca: {nombre_actual} -> {nuevo_nombre}")
                
                self.controller.db.conn.commit()
                self._cargar_marcas()
                messagebox.showinfo("Éxito", "Marca modificada correctamente")
            except ValueError as e:
                messagebox.showerror("Error", str(e))
            except Exception as e:
                self.controller.db.conn.rollback()
                messagebox.showerror("Error", f"No se pudo modificar la marca: {str(e)}")
    
    def _eliminar_marca(self):
        """Elimina una marca existente"""
        seleccion = self.treeview.selection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione una marca para eliminar")
            return
            
        item = self.treeview.item(seleccion[0])
        marca_id, nombre = item['values']
        
        # Verificar si la marca tiene modelos asociados
        self.controller.db.cursor.execute(
            "SELECT COUNT(*) FROM modelos WHERE id_marca = ?", (marca_id,))
        count_modelos = self.controller.db.cursor.fetchone()[0]
        
        if count_modelos > 0:
            messagebox.showerror("Error", 
                               "No se puede eliminar la marca porque tiene modelos asociados.\n"
                               "Elimine primero los modelos relacionados.")
            return
            
        if messagebox.askyesno("Confirmar", f"¿Está seguro que desea eliminar la marca '{nombre}'?"):
            try:
                self.controller.db.cursor.execute(
                    "DELETE FROM marcas WHERE id = ?", (marca_id,))
                
                self.controller.db.registrar_auditoria(
                    self.controller.current_user['id'], 'ELIMINAR_MARCA', 'marcas', marca_id,
                    f"Marca eliminada: {nombre}")
                
                self.controller.db.conn.commit()
                self._cargar_marcas()
                messagebox.showinfo("Éxito", "Marca eliminada correctamente")
            except Exception as e:
                self.controller.db.conn.rollback()
                messagebox.showerror("Error", f"No se pudo eliminar la marca: {str(e)}")
        

class ModelosTonerView(ttk.Frame):
    """Vista para gestionar modelos de toner"""
    
    def initialize(self, *args, **kwargs):
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
    
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self._setup_ui()
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        
        # Frame superior con controles
        control_frame = ttk.Frame(self)
        control_frame.grid(row=0, column=0, sticky="ew", pady=5)
        
        # Combobox para seleccionar marca
        ttk.Label(control_frame, text="Marca:").pack(side=tk.LEFT, padx=5)
        self.marca_combobox = ttk.Combobox(control_frame, state="readonly")
        self.marca_combobox.pack(side=tk.LEFT, padx=5)
        self.marca_combobox.bind("<<ComboboxSelected>>", lambda e: self._cargar_modelos())
        
        # Botones
        button_frame = ttk.Frame(control_frame)
        button_frame.pack(side=tk.RIGHT, padx=5)
        
        ttk.Button(button_frame, text="Agregar Modelo", 
                  command=self._agregar_modelo).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Modificar Modelo", 
                  command=self._modificar_modelo).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Eliminar Modelo", 
                  command=self._eliminar_modelo).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Volver", 
                  command=lambda: self.controller.mostrar_vista("MainView")).pack(side=tk.LEFT, padx=5)
        
        # Treeview para mostrar los modelos
        self.treeview = ttk.Treeview(self, columns=("ID", "Marca", "Nombre"), show="headings")
        self.treeview.heading("ID", text="ID")
        self.treeview.heading("Marca", text="Marca")
        self.treeview.heading("Nombre", text="Nombre")
        self.treeview.column("ID", width=50, anchor=tk.CENTER)
        self.treeview.column("Marca", width=150, anchor=tk.W)
        self.treeview.column("Nombre", width=150, anchor=tk.W)
        
        scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.treeview.yview)
        self.treeview.configure(yscrollcommand=scrollbar.set)
        
        self.treeview.grid(row=1, column=0, sticky="nsew")
        scrollbar.grid(row=1, column=1, sticky="ns")
        
        # Cargar datos iniciales
        self._cargar_marcas()
    
    def _cargar_marcas(self):
        """Carga las marcas en el combobox"""
        self.marca_combobox['values'] = [m[1] for m in self.controller.db.obtener_marcas()]
        if self.marca_combobox['values']:
            self.marca_combobox.current(0)
            self._cargar_modelos()
    
    def _cargar_modelos(self):
        """Carga los modelos de la marca seleccionada"""
        marca_seleccionada = self.marca_combobox.get()
        if not marca_seleccionada:
            return
            
        # Limpiar treeview
        for item in self.treeview.get_children():
            self.treeview.delete(item)
            
        # Obtener ID de la marca seleccionada
        marcas = self.controller.db.obtener_marcas()
        marca_id = next(m[0] for m in marcas if m[1] == marca_seleccionada)
        
        # Obtener y mostrar modelos
        modelos = self.controller.db.obtener_modelos(marca_id)
        for modelo in modelos:
            self.treeview.insert("", tk.END, values=(modelo[0], marca_seleccionada, modelo[1]))
    
    def _agregar_modelo(self):
        """Agrega un nuevo modelo a la marca seleccionada"""
        marca_seleccionada = self.marca_combobox.get()
        if not marca_seleccionada:
            messagebox.showwarning("Advertencia", "Seleccione una marca primero")
            return
            
        nombre = simpledialog.askstring("Agregar Modelo", "Ingrese el nombre del nuevo modelo:")
        if nombre:
            try:
                # Obtener ID de la marca seleccionada
                marcas = self.controller.db.obtener_marcas()
                marca_id = next(m[0] for m in marcas if m[1] == marca_seleccionada)
                
                # Agregar el modelo
                self.controller.db.agregar_modelo(marca_id, nombre, self.controller.current_user['id'])
                self._cargar_modelos()
                messagebox.showinfo("Éxito", "Modelo agregado correctamente")
            except ValueError as e:
                messagebox.showerror("Error", str(e))
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo agregar el modelo: {str(e)}")
    
    def _modificar_modelo(self):
        """Modifica un modelo existente"""
        seleccion = self.treeview.selection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione un modelo para modificar")
            return
            
        item = self.treeview.item(seleccion[0])
        modelo_id, marca_nombre, nombre_actual = item['values']
        
        nuevo_nombre = simpledialog.askstring("Modificar Modelo", "Ingrese el nuevo nombre:", 
                                            initialvalue=nombre_actual)
        if nuevo_nombre and nuevo_nombre != nombre_actual:
            try:
                # Obtener ID de la marca
                marcas = self.controller.db.obtener_marcas()
                marca_id = next(m[0] for m in marcas if m[1] == marca_nombre)
                
                # Verificar que no exista ya un modelo con ese nombre para esta marca
                modelos = self.controller.db.obtener_modelos(marca_id)
                if any(nuevo_nombre.lower() == m[1].lower() for m in modelos):
                    raise ValueError("Ya existe un modelo con ese nombre para esta marca")
                
                # Actualizar el modelo
                self.controller.db.cursor.execute(
                    "UPDATE modelos SET nombre = ? WHERE id = ?",
                    (nuevo_nombre, modelo_id))
                
                self.controller.db.registrar_auditoria(
                    self.controller.current_user['id'], 'MODIFICAR_MODELO', 'modelos', modelo_id,
                    f"Modificación de modelo: {nombre_actual} -> {nuevo_nombre}")
                
                self.controller.db.conn.commit()
                self._cargar_modelos()
                messagebox.showinfo("Éxito", "Modelo modificado correctamente")
            except ValueError as e:
                messagebox.showerror("Error", str(e))
            except Exception as e:
                self.controller.db.conn.rollback()
                messagebox.showerror("Error", f"No se pudo modificar el modelo: {str(e)}")
    
    def _eliminar_modelo(self):
        """Elimina un modelo existente"""
        seleccion = self.treeview.selection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione un modelo para eliminar")
            return
            
        item = self.treeview.item(seleccion[0])
        modelo_id, marca_nombre, nombre = item['values']
        
        # Verificar si el modelo tiene equipos asociados
        self.controller.db.cursor.execute(
            "SELECT COUNT(*) FROM equipos WHERE id_modelo = ?", (modelo_id,))
        count_equipos = self.controller.db.cursor.fetchone()[0]
        
        if count_equipos > 0:
            messagebox.showerror("Error", 
                               "No se puede eliminar el modelo porque tiene equipos asociados.\n"
                               "Elimine primero los equipos relacionados.")
            return
            
        if messagebox.askyesno("Confirmar", f"¿Está seguro que desea eliminar el modelo '{nombre}'?"):
            try:
                self.controller.db.cursor.execute(
                    "DELETE FROM modelos WHERE id = ?", (modelo_id,))
                
                self.controller.db.registrar_auditoria(
                    self.controller.current_user['id'], 'ELIMINAR_MODELO', 'modelos', modelo_id,
                    f"Modelo eliminado: {nombre} (Marca: {marca_nombre})")
                
                self.controller.db.conn.commit()
                self._cargar_modelos()
                messagebox.showinfo("Éxito", "Modelo eliminado correctamente")
            except Exception as e:
                self.controller.db.conn.rollback()
                messagebox.showerror("Error", f"No se pudo eliminar el modelo: {str(e)}")

class RetiroTonerView(ttk.Frame):
    """Vista para registrar retiro de toner"""
    
    def initialize(self, *args, **kwargs):
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
    
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self._setup_ui()
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        
        # Frame principal
        main_frame = ttk.Frame(self, padding="10")
        main_frame.grid(row=0, column=0, sticky="nsew")
        
        # Combobox para marca
        ttk.Label(main_frame, text="Marca:").grid(row=0, column=0, sticky="w", pady=5)
        self.marca_combo = ttk.Combobox(main_frame, state="readonly")
        self.marca_combo.grid(row=0, column=1, sticky="ew", padx=5, pady=5)
        self.marca_combo.bind("<<ComboboxSelected>>", lambda e: self._cargar_modelos())
        
        # Combobox para modelo
        ttk.Label(main_frame, text="Modelo:").grid(row=1, column=0, sticky="w", pady=5)
        self.modelo_combo = ttk.Combobox(main_frame, state="readonly")
        self.modelo_combo.grid(row=1, column=1, sticky="ew", padx=5, pady=5)
        
        # Campo para cantidad
        ttk.Label(main_frame, text="Cantidad:").grid(row=2, column=0, sticky="w", pady=5)
        self.cantidad_entry = ttk.Entry(main_frame)
        self.cantidad_entry.grid(row=2, column=1, sticky="ew", padx=5, pady=5)
        
        # Campo para responsable
        ttk.Label(main_frame, text="Responsable:").grid(row=3, column=0, sticky="w", pady=5)
        self.responsable_entry = ttk.Entry(main_frame)
        self.responsable_entry.grid(row=3, column=1, sticky="ew", padx=5, pady=5)
        
        # Campo para sector
        ttk.Label(main_frame, text="Sector:").grid(row=4, column=0, sticky="w", pady=5)
        self.sector_entry = ttk.Entry(main_frame)
        self.sector_entry.grid(row=4, column=1, sticky="ew", padx=5, pady=5)
        
        # Campo para observaciones
        ttk.Label(main_frame, text="Observaciones:").grid(row=5, column=0, sticky="nw", pady=5)
        self.observaciones_text = tk.Text(main_frame, height=5, width=30)
        self.observaciones_text.grid(row=5, column=1, sticky="ew", padx=5, pady=5)
        
        # Botones
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=6, column=0, columnspan=2, pady=10)
        
        ttk.Button(button_frame, text="Registrar Retiro", 
                 command=self._registrar_retiro).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cancelar", 
                 command=lambda: self.controller.mostrar_vista("TonerView")).pack(side=tk.RIGHT, padx=5)
    
    def _cargar_marcas(self):
        """Carga las marcas en el combobox"""
        try:
            marcas = self.controller.db.obtener_marcas_toner()
            self.marca_combo['values'] = [m[1] for m in marcas]
            if marcas:
                self.marca_combo.current(0)
                self._cargar_modelos()
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar las marcas: {str(e)}")
    
    def _cargar_modelos(self):
        """Carga los modelos según la marca seleccionada"""
        marca = self.marca_combo.get()
        if not marca:
            return
            
        try:
            # Obtener ID de la marca seleccionada
            marcas = self.controller.db.obtener_marcas_toner()
            marca_id = next(m[0] for m in marcas if m[1] == marca)
            
            # Obtener y mostrar modelos
            modelos = self.controller.db.obtener_modelos_toner(marca_id)
            self.modelo_combo['values'] = [m[1] for m in modelos]
            if modelos:
                self.modelo_combo.current(0)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar los modelos: {str(e)}")
    
    def _registrar_retiro(self):
        """Registra el retiro de toner"""
        marca = self.marca_combo.get()
        modelo = self.modelo_combo.get()
        cantidad = self.cantidad_entry.get()
        responsable = self.responsable_entry.get()
        sector = self.sector_entry.get()
        observaciones = self.observaciones_text.get("1.0", tk.END).strip()
        
        if not marca or not modelo or not cantidad or not responsable or not sector:
            messagebox.showwarning("Advertencia", "Complete todos los campos obligatorios")
            return
            
        try:
            cantidad = int(cantidad)
            if cantidad <= 0:
                raise ValueError("La cantidad debe ser mayor a cero")
            
            # Obtener ID del modelo
            marcas = self.controller.db.obtener_marcas_toner()
            id_marca = next(m[0] for m in marcas if m[1] == marca)
            
            modelos = self.controller.db.obtener_modelos_toner(id_marca)
            id_modelo = next(m[0] for m in modelos if m[1] == modelo)
            
            # Verificar stock disponible
            stock = self.controller.db.obtener_stock_toner_por_modelo(id_modelo)
            if stock < cantidad:
                raise ValueError(f"No hay suficiente stock. Disponible: {stock}")
            
            # Registrar movimiento de retiro
            self.controller.db.registrar_movimiento_toner(
                id_modelo, 
                'retiro', 
                cantidad, 
                responsable, 
                sector, 
                observaciones, 
                self.controller.current_user['id']
            )
            
            messagebox.showinfo("Éxito", "Retiro registrado correctamente")
            self.controller.mostrar_vista("TonerView")
        except ValueError as e:
            messagebox.showerror("Error", f"Dato inválido: {str(e)}")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo registrar el retiro: {str(e)}")

class RecargaTonerView(ttk.Frame):
    """Vista para gestionar recargas de toner"""
    
    def initialize(self, *args, **kwargs):
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
    
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self._setup_ui()
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        
        # Frame principal
        main_frame = ttk.Frame(self, padding="10")
        main_frame.grid(row=0, column=0, sticky="nsew")
        
        # Combobox para marca
        ttk.Label(main_frame, text="Marca:").grid(row=0, column=0, sticky="w", pady=5)
        self.marca_combo = ttk.Combobox(main_frame, state="readonly")
        self.marca_combo.grid(row=0, column=1, sticky="ew", padx=5, pady=5)
        self.marca_combo.bind("<<ComboboxSelected>>", lambda e: self._cargar_modelos())
        
        # Combobox para modelo
        ttk.Label(main_frame, text="Modelo:").grid(row=1, column=0, sticky="w", pady=5)
        self.modelo_combo = ttk.Combobox(main_frame, state="readonly")
        self.modelo_combo.grid(row=1, column=1, sticky="ew", padx=5, pady=5)
        
        # Campo para cantidad
        ttk.Label(main_frame, text="Cantidad:").grid(row=2, column=0, sticky="w", pady=5)
        self.cantidad_entry = ttk.Entry(main_frame)
        self.cantidad_entry.grid(row=2, column=1, sticky="ew", padx=5, pady=5)
        
        # Combobox para empresa
        ttk.Label(main_frame, text="Empresa:").grid(row=3, column=0, sticky="w", pady=5)
        self.empresa_combo = ttk.Combobox(main_frame, state="readonly")
        self.empresa_combo.grid(row=3, column=1, sticky="ew", padx=5, pady=5)
        
        # Campo para observaciones
        ttk.Label(main_frame, text="Observaciones:").grid(row=4, column=0, sticky="nw", pady=5)
        self.observaciones_text = tk.Text(main_frame, height=5, width=30)
        self.observaciones_text.grid(row=4, column=1, sticky="ew", padx=5, pady=5)
        
        # Botones
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=5, column=0, columnspan=2, pady=10)
        
        ttk.Button(button_frame, text="Registrar Envío", 
                 command=self._registrar_envio).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cancelar", 
                 command=lambda: self.controller.mostrar_vista("TonerView")).pack(side=tk.RIGHT, padx=5)
    
    def _cargar_marcas(self):
        """Carga las marcas en el combobox"""
        try:
            marcas = self.controller.db.obtener_marcas_toner()
            self.marca_combo['values'] = [m[1] for m in marcas]
            if marcas:
                self.marca_combo.current(0)
                self._cargar_modelos()
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar las marcas: {str(e)}")
    
    def _cargar_modelos(self):
        """Carga los modelos según la marca seleccionada"""
        marca = self.marca_combo.get()
        if not marca:
            return
            
        try:
            # Obtener ID de la marca seleccionada
            marcas = self.controller.db.obtener_marcas_toner()
            marca_id = next(m[0] for m in marcas if m[1] == marca)
            
            # Obtener y mostrar modelos
            modelos = self.controller.db.obtener_modelos_toner(marca_id)
            self.modelo_combo['values'] = [m[1] for m in modelos]
            if modelos:
                self.modelo_combo.current(0)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar los modelos: {str(e)}")
    
    def _cargar_empresas(self):
        """Carga las empresas de recarga en el combobox"""
        try:
            empresas = self.controller.db.obtener_empresas_recarga()
            self.empresa_combo['values'] = [e[1] for e in empresas]
            if empresas:
                self.empresa_combo.current(0)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar las empresas: {str(e)}")
    
    def _registrar_envio(self):
        """Registra el envío de toner para recarga"""
        marca = self.marca_combo.get()
        modelo = self.modelo_combo.get()
        cantidad = self.cantidad_entry.get()
        empresa = self.empresa_combo.get()
        observaciones = self.observaciones_text.get("1.0", tk.END).strip()
        
        if not marca or not modelo or not cantidad or not empresa:
            messagebox.showwarning("Advertencia", "Complete todos los campos obligatorios")
            return
            
        try:
            cantidad = int(cantidad)
            if cantidad <= 0:
                raise ValueError("La cantidad debe ser mayor a cero")
            
            # Obtener ID del modelo
            marcas = self.controller.db.obtener_marcas_toner()
            id_marca = next(m[0] for m in marcas if m[1] == marca)
            
            modelos = self.controller.db.obtener_modelos_toner(id_marca)
            id_modelo = next(m[0] for m in modelos if m[1] == modelo)
            
            # Obtener ID de la empresa
            empresas = self.controller.db.obtener_empresas_recarga()
            id_empresa = next(e[0] for e in empresas if e[1] == empresa)
            
            # Registrar recarga
            self.controller.db.registrar_recarga_toner(
                id_modelo, 
                cantidad, 
                id_empresa, 
                observaciones, 
                self.controller.current_user['id']
            )
            
            messagebox.showinfo("Éxito", "Envío para recarga registrado correctamente")
            self.controller.mostrar_vista("TonerView")
        except ValueError as e:
            messagebox.showerror("Error", f"Dato inválido: {str(e)}")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo registrar el envío: {str(e)}")

class InformesTonerView(ttk.Frame):
    """Vista para generar informes de toner"""
    
    def initialize(self, *args, **kwargs):
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
    
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self._setup_ui()
    
    def _setup_ui(self):
        # Formulario para generar informes
        pass


class EmpresasRecargaView(ttk.Frame):
    """Vista para gestionar empresas de recarga de toner"""
    
    def initialize(self, *args, **kwargs):
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
        self._cargar_empresas()
    
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self._setup_ui()
    
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        
        # Frame superior con botones
        button_frame = ttk.Frame(self)
        button_frame.grid(row=0, column=0, sticky="ew", pady=5)
        
        ttk.Button(button_frame, text="Agregar Empresa", 
                 command=self._agregar_empresa).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Modificar Empresa", 
                 command=self._modificar_empresa).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Eliminar Empresa", 
                 command=self._eliminar_empresa).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Volver", 
                 command=lambda: self.controller.mostrar_vista("TonerView")).pack(side=tk.RIGHT, padx=5)
        
        # Treeview para mostrar las empresas
        self.treeview = ttk.Treeview(self, columns=("ID", "Nombre", "Contacto", "Teléfono"), show="headings")
        self.treeview.heading("ID", text="ID")
        self.treeview.heading("Nombre", text="Nombre")
        self.treeview.heading("Contacto", text="Contacto")
        self.treeview.heading("Teléfono", text="Teléfono")
        self.treeview.column("ID", width=50, anchor=tk.CENTER)
        self.treeview.column("Nombre", width=150, anchor=tk.W)
        self.treeview.column("Contacto", width=150, anchor=tk.W)
        self.treeview.column("Teléfono", width=100, anchor=tk.W)
        
        scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.treeview.yview)
        self.treeview.configure(yscrollcommand=scrollbar.set)
        
        self.treeview.grid(row=1, column=0, sticky="nsew")
        scrollbar.grid(row=1, column=1, sticky="ns")
        
        # Cargar datos iniciales
        self._cargar_empresas()
    
    def _cargar_empresas(self):
        """Carga las empresas desde la base de datos"""
        for item in self.treeview.get_children():
            self.treeview.delete(item)
            
        try:
            empresas = self.controller.db.obtener_empresas_recarga()
            for empresa in empresas:
                self.treeview.insert("", tk.END, values=empresa)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar las empresas: {str(e)}")
    
    def _agregar_empresa(self):
        """Agrega una nueva empresa de recarga"""
        dialog = tk.Toplevel()
        dialog.title("Agregar Empresa")
        dialog.transient(self)
        dialog.grab_set()
        
        # Variables
        nombre_var = tk.StringVar()
        contacto_var = tk.StringVar()
        telefono_var = tk.StringVar()
        
        # Frame principal
        main_frame = ttk.Frame(dialog, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Campos del formulario
        ttk.Label(main_frame, text="Nombre:").grid(row=0, column=0, sticky="w", pady=5)
        nombre_entry = ttk.Entry(main_frame, textvariable=nombre_var)
        nombre_entry.grid(row=0, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(main_frame, text="Contacto:").grid(row=1, column=0, sticky="w", pady=5)
        contacto_entry = ttk.Entry(main_frame, textvariable=contacto_var)
        contacto_entry.grid(row=1, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(main_frame, text="Teléfono:").grid(row=2, column=0, sticky="w", pady=5)
        telefono_entry = ttk.Entry(main_frame, textvariable=telefono_var)
        telefono_entry.grid(row=2, column=1, sticky="ew", padx=5, pady=5)
        
        # Botones
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=3, column=0, columnspan=2, pady=10)
        
        def on_ok():
            nombre = nombre_var.get().strip()
            contacto = contacto_var.get().strip()
            telefono = telefono_var.get().strip()
            
            if not nombre:
                messagebox.showwarning("Advertencia", "El nombre es obligatorio")
                return
                
            try:
                self.controller.db.cursor.execute(
                    "INSERT INTO empresas_recarga (nombre, contacto, telefono) VALUES (?, ?, ?)",
                    (nombre, contacto or None, telefono or None))
                
                empresa_id = self.controller.db.cursor.lastrowid
                
                self.controller.db.registrar_auditoria(
                    self.controller.current_user['id'], 'ALTA_EMPRESA_RECARGA', 'empresas_recarga', empresa_id,
                    f"Nueva empresa agregada: {nombre}")
                
                self.controller.db.conn.commit()
                self._cargar_empresas()
                dialog.destroy()
                messagebox.showinfo("Éxito", "Empresa agregada correctamente")
            except sqlite3.IntegrityError:
                messagebox.showerror("Error", "Ya existe una empresa con ese nombre")
            except Exception as e:
                self.controller.db.conn.rollback()
                messagebox.showerror("Error", f"No se pudo agregar la empresa: {str(e)}")
        
        ttk.Button(button_frame, text="Aceptar", command=on_ok).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cancelar", command=dialog.destroy).pack(side=tk.RIGHT, padx=5)
        
        # Centrar la ventana
        dialog.update_idletasks()
        width = dialog.winfo_width()
        height = dialog.winfo_height()
        x = (dialog.winfo_screenwidth() // 2) - (width // 2)
        y = (dialog.winfo_screenheight() // 2) - (height // 2)
        dialog.geometry(f'{width}x{height}+{x}+{y}')
        
        nombre_entry.focus_set()
    
    def _modificar_empresa(self):
        """Modifica una empresa existente"""
        seleccion = self.treeview.selection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione una empresa para modificar")
            return
            
        item = self.treeview.item(seleccion[0])
        empresa_id, nombre_actual, contacto_actual, telefono_actual = item['values']
        
        dialog = tk.Toplevel()
        dialog.title("Modificar Empresa")
        dialog.transient(self)
        dialog.grab_set()
        
        # Variables
        nombre_var = tk.StringVar(value=nombre_actual)
        contacto_var = tk.StringVar(value=contacto_actual or "")
        telefono_var = tk.StringVar(value=telefono_actual or "")
        
        # Frame principal
        main_frame = ttk.Frame(dialog, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Campos del formulario
        ttk.Label(main_frame, text="Nombre:").grid(row=0, column=0, sticky="w", pady=5)
        nombre_entry = ttk.Entry(main_frame, textvariable=nombre_var)
        nombre_entry.grid(row=0, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(main_frame, text="Contacto:").grid(row=1, column=0, sticky="w", pady=5)
        contacto_entry = ttk.Entry(main_frame, textvariable=contacto_var)
        contacto_entry.grid(row=1, column=1, sticky="ew", padx=5, pady=5)
        
        ttk.Label(main_frame, text="Teléfono:").grid(row=2, column=0, sticky="w", pady=5)
        telefono_entry = ttk.Entry(main_frame, textvariable=telefono_var)
        telefono_entry.grid(row=2, column=1, sticky="ew", padx=5, pady=5)
        
        # Botones
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=3, column=0, columnspan=2, pady=10)
        
        def on_ok():
            nuevo_nombre = nombre_var.get().strip()
            nuevo_contacto = contacto_var.get().strip()
            nuevo_telefono = telefono_var.get().strip()
            
            if not nuevo_nombre:
                messagebox.showwarning("Advertencia", "El nombre es obligatorio")
                return
                
            try:
                # Verificar si el nombre cambió y si ya existe
                if nuevo_nombre != nombre_actual:
                    self.controller.db.cursor.execute(
                        "SELECT COUNT(*) FROM empresas_recarga WHERE nombre = ?", 
                        (nuevo_nombre,))
                    if self.controller.db.cursor.fetchone()[0] > 0:
                        raise ValueError("Ya existe una empresa con ese nombre")
                
                # Actualizar la empresa
                self.controller.db.cursor.execute(
                    "UPDATE empresas_recarga SET nombre = ?, contacto = ?, telefono = ? WHERE id = ?",
                    (nuevo_nombre, nuevo_contacto or None, nuevo_telefono or None, empresa_id))
                
                self.controller.db.registrar_auditoria(
                    self.controller.current_user['id'], 'MODIFICAR_EMPRESA_RECARGA', 'empresas_recarga', empresa_id,
                    f"Empresa modificada: {nombre_actual} -> {nuevo_nombre}")
                
                self.controller.db.conn.commit()
                self._cargar_empresas()
                dialog.destroy()
                messagebox.showinfo("Éxito", "Empresa modificada correctamente")
            except ValueError as e:
                messagebox.showerror("Error", str(e))
            except Exception as e:
                self.controller.db.conn.rollback()
                messagebox.showerror("Error", f"No se pudo modificar la empresa: {str(e)}")
        
        ttk.Button(button_frame, text="Aceptar", command=on_ok).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cancelar", command=dialog.destroy).pack(side=tk.RIGHT, padx=5)
        
        # Centrar la ventana
        dialog.update_idletasks()
        width = dialog.winfo_width()
        height = dialog.winfo_height()
        x = (dialog.winfo_screenwidth() // 2) - (width // 2)
        y = (dialog.winfo_screenheight() // 2) - (height // 2)
        dialog.geometry(f'{width}x{height}+{x}+{y}')
        
        nombre_entry.focus_set()
    
    def _eliminar_empresa(self):
        """Elimina una empresa existente"""
        seleccion = self.treeview.selection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione una empresa para eliminar")
            return
            
        item = self.treeview.item(seleccion[0])
        empresa_id, nombre, _, _ = item['values']
        
        # Verificar si tiene recargas asociadas
        try:
            self.controller.db.cursor.execute(
                "SELECT COUNT(*) FROM recargas_toner WHERE id_empresa = ?", 
                (empresa_id,))
            if self.controller.db.cursor.fetchone()[0] > 0:
                messagebox.showerror("Error", 
                                   "No se puede eliminar: la empresa tiene recargas asociadas")
                return
                
            if messagebox.askyesno("Confirmar", f"¿Eliminar la empresa '{nombre}'?"):
                self.controller.db.cursor.execute(
                    "DELETE FROM empresas_recarga WHERE id = ?", 
                    (empresa_id,))
                
                self.controller.db.registrar_auditoria(
                    self.controller.current_user['id'], 'ELIMINAR_EMPRESA_RECARGA', 'empresas_recarga', empresa_id,
                    f"Empresa eliminada: {nombre}")
                
                self.controller.db.conn.commit()
                self._cargar_empresas()
                messagebox.showinfo("Éxito", "Empresa eliminada correctamente")
        except Exception as e:
            self.controller.db.conn.rollback()
            messagebox.showerror("Error", f"No se pudo eliminar la empresa: {str(e)}")

class MainController:
    """Controlador principal de la aplicación"""


    def initialize(self, *args, **kwargs):
        # Verificar autenticación primero
        if not self.controller.current_user:
            self.controller.mostrar_vista("LoginView")
            return
    
    def __init__(self, root):
        self.root = root
        self.root.withdraw()
        self.db = Database()
        self.current_user = None
        
        # Configurar ventana principal
        self._configure_window()
        self._configure_styles()
        
        # Contenedor para vistas
        self.container = ttk.Frame(self.root)
        self.container.pack(fill=tk.BOTH, expand=True)
        self.container.grid_rowconfigure(0, weight=1)
        self.container.grid_columnconfigure(0, weight=1)
        
        # Diccionario de vistas
        self.views = {}
        self._register_views()

        # Mostrar ventana principal solo después de que el login esté listo
        self.root.after(100, self.root.deiconify)
        
        
        # Mostrar vista inicial
        self.mostrar_vista("LoginView")

        # Ahora mostrar la ventana principal
        self.root.deiconify()

        # Forzar que el LoginView esté al frente
        self.views["LoginView"].tkraise()



        
    def _configure_window(self):
        """Configura la ventana principal con nuevos estilos"""
        self.root.title("Sistema de Gestión de Stock - Poder Judicial")
        self.root.geometry("1200x750")
        self.root.configure(bg=Config.BACKGROUND_COLOR)
        try:
            self.root.iconbitmap(default='logo_pj.ico')
        except:
            pass
        self.root.protocol("WM_DELETE_WINDOW", self.cerrar_aplicacion)
    
    def _configure_styles(self):
        """Configura los estilos globales de la aplicación"""
        style = ttk.Style()
        style.theme_use(Config.UI_THEME)
        
        # Configuración general
        style.configure('.', 
                      background=Config.BACKGROUND_COLOR,
                      foreground=Config.TEXT_COLOR,
                      font=('Helvetica', 10))
        
        # Frames
        style.configure('TFrame', 
                      background=Config.BACKGROUND_COLOR)
        style.configure('Header.TFrame',
                      background=Config.PRIMARY_COLOR)
        
        # Labels
        style.configure('TLabel',
                      font=('Helvetica', 10),
                      background=Config.BACKGROUND_COLOR,
                      foreground=Config.TEXT_COLOR)
        style.configure('Header.TLabel',
                      font=('Helvetica', 12, 'bold'),
                      foreground='white',
                      background=Config.PRIMARY_COLOR)
        style.configure('Title.TLabel',
                      font=('Helvetica', 14, 'bold'),
                      foreground=Config.PRIMARY_COLOR)
        
        # Buttons
        style.configure('TButton',
                      font=('Helvetica', 10),
                      padding=8,
                      background=Config.SECONDARY_COLOR,
                      foreground='white',
                      borderwidth=1,
                      relief="raised")
        style.map('TButton',
                foreground=[('pressed', 'white'), ('active', 'white')],
                background=[('pressed', Config.PRIMARY_COLOR), ('active', '#2185d0')])
        
        style.configure('Secondary.TButton',
                      background=Config.HIGHLIGHT_COLOR,
                      foreground=Config.TEXT_COLOR)
        style.map('Secondary.TButton',
                foreground=[('pressed', Config.TEXT_COLOR), ('active', Config.TEXT_COLOR)],
                background=[('pressed', Config.BORDER_COLOR), ('active', '#e2e6ea')])
        
        # Entries
        style.configure('TEntry',
                      fieldbackground='white',
                      foreground=Config.TEXT_COLOR,
                      bordercolor=Config.BORDER_COLOR,
                      lightcolor=Config.BORDER_COLOR,
                      darkcolor=Config.BORDER_COLOR,
                      padding=5)
        
        # Comboboxes
        style.configure('TCombobox',
                      fieldbackground='white',
                      foreground=Config.TEXT_COLOR,
                      selectbackground=Config.SECONDARY_COLOR,
                      selectforeground='white')
        
        # Treeview
        style.configure('Treeview',
                      font=('Helvetica', 10),
                      rowheight=25,
                      background='white',
                      fieldbackground='white',
                      foreground=Config.TEXT_COLOR)
        style.configure('Treeview.Heading',
                      font=('Helvetica', 10, 'bold'),
                      background=Config.PRIMARY_COLOR,
                      foreground='white',
                      relief='flat')
        style.map('Treeview.Heading',
                background=[('active', Config.SECONDARY_COLOR)])
        
        # Notebook (pestañas)
        style.configure('TNotebook',
                      background=Config.BACKGROUND_COLOR)
        style.configure('TNotebook.Tab',
                      font=('Helvetica', 10, 'bold'),
                      padding=(10, 5),
                      background=Config.HIGHLIGHT_COLOR,
                      foreground=Config.TEXT_COLOR)
        style.map('TNotebook.Tab',
                background=[('selected', Config.SECONDARY_COLOR), ('active', '#e2e6ea')],
                foreground=[('selected', 'white'), ('active', Config.TEXT_COLOR)])
        
        # Scrollbars
        style.configure('Vertical.TScrollbar',
                      background=Config.HIGHLIGHT_COLOR,
                      troughcolor=Config.BACKGROUND_COLOR,
                      bordercolor=Config.BORDER_COLOR,
                      arrowcolor=Config.TEXT_COLOR)
        
        # LabelFrames
        style.configure('TLabelframe',
                      background=Config.BACKGROUND_COLOR,
                      bordercolor=Config.BORDER_COLOR)
        style.configure('TLabelframe.Label',
                      font=('Helvetica', 10, 'bold'),
                      foreground=Config.PRIMARY_COLOR)
    
    def _register_views(self):
        """Registra todas las vistas disponibles"""
        self.views['LoginView'] = LoginView(self.container, self)
        self.views['MainView'] = MainView(self.container, self)
        self.views['EquipmentView'] = EquipmentView(self.container, self)
        self.views['EditarEquipoView'] = EditarEquipoView(self.container, self)  # Nueva vista
        self.views['ReparacionView'] = ReparacionView(self.container, self)
        self.views['MarcasView'] = MarcasView(self.container, self)
        self.views['ModelosView'] = ModelosView(self.container, self)
        self.views['ReporteEquiposView'] = ReporteEquiposView(self.container, self)
        self.views['ReporteRepuestosView'] = ReporteRepuestosView(self.container, self)

        self.views['ReporteResumenRepuestosView'] = ReporteResumenRepuestosView(self.container, self)

        self.views['UsuariosView'] = UsuariosView(self.container, self)


        # Nuevas vistas para toner
        self.views['TonerView'] = TonerView(self.container, self)
        self.views['MarcasTonerView'] = MarcasTonerView(self.container, self)
        self.views['ModelosTonerView'] = ModelosTonerView(self.container, self)
        self.views['RetiroTonerView'] = RetiroTonerView(self.container, self)
        self.views['RecargaTonerView'] = RecargaTonerView(self.container, self)
        self.views['InformesTonerView'] = InformesTonerView(self.container, self)
        
        # Configurar todas las vistas en el contenedor
        for view in self.views.values():
            view.grid(row=0, column=0, sticky="nsew")
    
    def mostrar_vista(self, view_name, *args, **kwargs):
        """Muestra una vista con verificación de autenticación"""
        
        # Permitir siempre el acceso al LoginView
        if view_name != "LoginView" and not self.current_user:
            view_name = "LoginView"
    
        view = self.views.get(view_name)
        if view:
            # Si la vista necesita inicialización
            if hasattr(view, 'initialize'):
                if not args and not kwargs:
                    view.initialize()  # Llamar sin argumentos para actualizar
                else:
                    view.initialize(*args, **kwargs)
        
            # Mostrar la vista
            view.tkraise()
        
            # Actualizar título de la ventana
            if view_name == "MainView" and self.current_user:
                self.root.title(f"Sistema de Gestión - Usuario: {self.current_user['username']}")
            elif view_name == "LoginView":
                self.root.title("Sistema de Gestión - Inicio de sesión")
        else:
            raise ValueError(f"Vista no encontrada: {view_name}")
    
    def mostrar_vista_principal(self, user_data):
        """Muestra la vista principal después del login"""

        if not user_data:
           messagebox.showerror("Error", "Debe iniciar sesión primero")
           return
        
        self.current_user = {
            'id': user_data[0],
            'rol': user_data[1],
            'username': user_data[2]
        }
        self.mostrar_vista("MainView")


        # Actualizar la barra de estado con el usuario
        if "MainView" in self.views:
            self.views["MainView"].actualizar_status(f"Usuario: {self.current_user['username']} | Rol: {self.current_user['rol']}")


        
    
    def autenticar_usuario(self, username, password):
        """Autentica un usuario con verificación estricta"""
        if not username or not password:
            messagebox.showwarning("Error", "Usuario y contraseña son requeridos")
            return None
    
        try:
            resultado = self.db.autenticar_usuario(username, password)
            if resultado:
                logging.info(f"Usuario autenticado: {username}")
                return resultado
        
            messagebox.showerror("Error", "Credenciales inválidas")
            return None
        except Exception as e:
            logging.error(f"Error en autenticación: {e}")
            messagebox.showerror("Error", f"Error al autenticar: {str(e)}")
            return None


    
    def cerrar_aplicacion(self):
        """Cierra la aplicación con confirmación"""
        if messagebox.askokcancel("Salir", "¿Está seguro que desea salir del sistema?"):
            # Registrar en logs
            if self.current_user:
                logging.info(f"Usuario {self.current_user['username']} cerró la aplicación")
            else:
                logging.info("Aplicación cerrada sin sesión activa")
        
            # Cerrar conexión a BD
            if hasattr(self, 'db') and hasattr(self.db, 'conn'):
                self.db.conn.close()
        
            # Destruir ventana
            self.root.destroy()
        
class ExportManager:
    """Manejador de exportación a diferentes formatos"""

    def _sanitize_sheet_name(name):
        """
        Limpia el nombre de hoja para cumplir con las restricciones de Excel:
        - Máximo 31 caracteres
        - "Caracteres inválidos: \\ / * ? [ ]"
        - No puede empezar o terminar con apóstrofe (')
        """
        invalid_chars = r":\/*?[]"
        replace_char = '-'
        
        # Reemplazar caracteres inválidos
        sanitized = ''.join([c if c not in invalid_chars else replace_char for c in name])
        
        # Eliminar espacios al inicio/final
        sanitized = sanitized.strip()
        
        # Limitar longitud
        return sanitized[:31]

    @staticmethod
    def export_to_pdf(data, filename, title="Reporte"):
        """Exporta datos a un archivo PDF"""
        try:
            doc = SimpleDocTemplate(filename, pagesize=Config.PDF_OPTIONS['pagesize'])
            elements = []
            
            styles = getSampleStyleSheet()
            elements.append(Paragraph(title, styles['Title']))
            
            # Convertir datos a tabla
            if not data:
                elements.append(Paragraph("No hay datos para mostrar", styles['BodyText']))
            else:
                if isinstance(data[0], dict):
                    headers = list(data[0].keys())
                    rows = [list(row.values()) for row in data]
                else:
                    headers = data[0]
                    rows = data[1:]
                
                table_data = [headers] + rows
                t = Table(table_data)
                
                # Estilo de tabla
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.grey),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                    ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0,0), (-1,0), 12),
                    ('BOTTOMPADDING', (0,0), (-1,0), 12),
                    ('BACKGROUND', (0,1), (-1,-1), colors.beige),
                    ('GRID', (0,0), (-1,-1), 1, colors.black)
                ]))
                
                elements.append(t)
            
            doc.build(elements)
            return True
        except Exception as e:
            logging.error(f"Error exportando a PDF: {e}")
            return False
    
    @staticmethod

    def export_to_word(data, filename, title="Reporte", subtitle=""):
        """Exporta datos a Word con formato profesional para el Poder Judicial"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN
            import os

            # Crear documento
            doc = Document()

            # --- Configuración de página ---
            section = doc.sections[0]
            section.top_margin = Inches(0.5)      # 2.54 cm
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)   # ~2 cm
            section.right_margin = Inches(0.5)

            # --- Encabezado institucional ---
            # Logo (si existe)
            logo_paths = [
                "logo_pj.png",
                os.path.join(os.path.dirname(__file__), "logo_pj.png"),
                os.path.join(os.path.dirname(__file__), "media", "logo_pj.png")
            ]
        
            logo_path = None
            for path in logo_paths:
                if os.path.exists(path):
                    logo_path = path
                    break

            if logo_path:
                header = doc.add_paragraph()
                header.alignment = WD_ALIGN.CENTER
                run = header.add_run()
                run.add_picture(logo_path, width=Inches(0.7))  # Ajusta el tamaño según necesidad
                #doc.add_paragraph()  # Espacio después del logo

            # Títulos institucionales
            titles = doc.add_paragraph()
            titles.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Línea 1 - PODER JUDICIAL
            line1 = titles.add_run("PODER JUDICIAL DE LA PROVINCIA DE JUJUY\n")
            line1.font.size = Pt(14)
            line1.font.bold = True
            
            # Línea 2 - DEPARTAMENTO
            line2 = titles.add_run("DEPARTAMENTO DE SISTEMAS Y TECNOLOGÍA DE LA INFORMACIÓN\n")
            line2.font.size = Pt(12)
            line2.font.bold = True
            
            # Línea 3 - Dirección
            line3 = titles.add_run("Argañaraz esq. Independencia -- San Salvador de Jujuy\n\n")
            line3.font.size = Pt(10)
            line3.font.bold = False


                

            
            # Títulos institucionales
            #doc.add_paragraph("PODER JUDICIAL DE LA PROVINCIA DE JUJUY", style='Heading 1').alignment = WD_ALIGN.CENTER
            #doc.add_paragraph("Departamento de Sistemas y Tecnología de la Información", style='Heading 2').alignment = WD_ALIGN.CENTER
            #doc.add_paragraph("Sistema de Gestión de Stock Informático", style='Heading 3').alignment = WD_ALIGN.CENTER

            # --- Título del reporte ---
            doc.add_paragraph().add_run(title).bold = True
            doc.add_paragraph(subtitle).alignment = WD_ALIGN.CENTER

            # --- Tabla de datos ---
            if data and len(data) > 1:
                table = doc.add_table(rows=1, cols=len(data[0]))
                table.style = 'Light Shading Accent 1'  # Estilo profesional
            
                # Encabezados
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(data[0]):
                    hdr_cells[i].text = str(header)
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN.CENTER

                # Datos
                for row in data[1:]:
                    row_cells = table.add_row().cells
                    for i, cell_value in enumerate(row):
                        row_cells[i].text = str(cell_value)
                        row_cells[i].paragraphs[0].alignment = WD_ALIGN.CENTER

            # --- Pie de página ---
            doc.add_page_break()  # Opcional: nueva página para el pie
            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"Documento generado automáticamente el {datetime.now().strftime('%d/%m/%Y a las %H:%M')}"
            footer_para.alignment = WD_ALIGN.RIGHT

            # Guardar
            doc.save(filename)
            return True

        except Exception as e:
            logging.error(f"Error al exportar a Word: {str(e)}", exc_info=True)
            return False

   
        
    @staticmethod
    def export_to_excel(data, filename, title="Reporte", sheet_name="Datos"):
        """Exporta datos a Excel con manejo robusto de errores"""
        try:
            # Validaciones iniciales
            if not data:
                logging.error("Intento de exportar datos vacíos a Excel")
                return False

            # Sanitizar nombre de hoja
            sheet_name = ExportManager._sanitize_sheet_name(sheet_name)
            if not sheet_name.strip():
                sheet_name = "Datos"
                logging.warning("Usando nombre de hoja por defecto")

            # Convertir datos a DataFrame
            if isinstance(data[0], dict):
                df = pd.DataFrame(data)
            else:
                if len(data) < 2:
                    logging.error("Solo encabezados sin datos")
                    return False
            
                headers = data[0]
                rows = data[1:]
                df = pd.DataFrame(rows, columns=headers)

            # Crear escritor Excel
            writer = pd.ExcelWriter(
                filename,
                engine='xlsxwriter',
                datetime_format='dd/mm/yyyy',
                date_format='dd/mm/yyyy'
            )

            # Escribir datos
            df.to_excel(
                writer,
                sheet_name=sheet_name,
                index=False,
                header=True,
                startrow=1
            )

            # Configurar formatos
            workbook = writer.book
            worksheet = writer.sheets[sheet_name]

            # Formato para título
            title_format = workbook.add_format({
                'bold': True,
                'size': 14,
                'align': 'center',
                'valign': 'vcenter'
            })

            # Escribir título
            last_col = chr(65 + len(df.columns) - 1)
            worksheet.merge_range(
                f'A1:{last_col}1',
                title,
                title_format
            )

            # Ajustar anchos de columna
            for i, col in enumerate(df.columns):
                max_len = max((
                    df[col].astype(str).map(len).max(),
                    len(str(col))
                )) + 2
                worksheet.set_column(i, i, min(max_len, 50))

            # Guardar
            writer.close()
            return True

        except Exception as e:
            logging.error(f"Error exportando a Excel: {e}")
            return False


    @staticmethod
    def export_informe_tecnico(data, filename):
        """Exporta un informe técnico detallado a Word"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from docx.enum.table import WD_TABLE_ALIGNMENT
            
            # Crear documento
            doc = Document()
            
            # Configurar márgenes
            section = doc.sections[0]
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            
            # Encabezado con logo institucional
            header = doc.add_paragraph()
            header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Intentar agregar logo (si existe)
            try:
                logo_path = None
                possible_paths = ['logo_pj.png', 'media/logo_pj.png']
                for path in possible_paths:
                    if os.path.exists(path):
                        logo_path = path
                        break
                
                if logo_path:
                    header.add_run().add_picture(logo_path, width=Inches(0.7))
                    header.add_run().add_break()
            except Exception as e:
                logging.warning(f"No se pudo agregar logo: {e}")
            
            # Títulos institucionales
            titles = doc.add_paragraph()
            titles.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Línea 1 - PODER JUDICIAL
            line1 = titles.add_run("PODER JUDICIAL DE LA PROVINCIA DE JUJUY\n")
            line1.font.size = Pt(14)
            line1.font.bold = True
            
            # Línea 2 - DEPARTAMENTO
            line2 = titles.add_run("DEPARTAMENTO DE SISTEMAS Y TECNOLOGÍA DE LA INFORMACIÓN\n")
            line2.font.size = Pt(12)
            line2.font.bold = True
            
            # Línea 3 - Dirección
            line3 = titles.add_run("Argañaraz esq. Independencia -- San Salvador de Jujuy\n\n")
            line3.font.size = Pt(10)
            line3.font.bold = False
            
            # Título del informe
            title = doc.add_paragraph(data['titulo'])
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title.style = "Heading 1"
            #title.style = doc.styles['Heading 1']  # ✅ Usando doc.styles
            
            # Fecha de generación
            fecha = doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            fecha.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            fecha.style = "Intense Quote"
            #fecha.style = doc.styles['Intense Quote']  # ✅ Usando doc.styles
            
            # Sección 1: Información del equipo
            doc.add_paragraph("1. Información del Equipo", style='Heading 2')
            
            # Tabla de información del equipo
            table_equipo = doc.add_table(rows=1, cols=2)
            table_equipo.style = 'Light Shading Accent 1'
            table_equipo.autofit = True
            
            for row in data['info_equipo']:
                row_cells = table_equipo.add_row().cells
                row_cells[0].text = str(row[0])
                row_cells[1].text = str(row[1])
                
                # Estilo para la primera columna (negrita)
                row_cells[0].paragraphs[0].runs[0].font.bold = True
            
            # Sección 2: Repuestos utilizados
            doc.add_paragraph("\n2. Repuestos Utilizados", style='Heading 2')
            
            if len(data['repuestos']) > 1:  # Si hay repuestos además del encabezado
                table_repuestos = doc.add_table(rows=1, cols=4, style='Light Shading Accent 1')
                table_repuestos.autofit = True
                
                # Encabezados
                hdr_cells = table_repuestos.rows[0].cells
                for i, header in enumerate(data['repuestos'][0]):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                    hdr_cells[i].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Datos de repuestos
                for row in data['repuestos'][1:]:
                    row_cells = table_repuestos.add_row().cells
                    for i, cell in enumerate(row):
                        row_cells[i].text = str(cell)
                        row_cells[i].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            else:
                doc.add_paragraph("No se utilizaron repuestos en esta reparación.", style='BodyText')
            
            # Sección 3: Observaciones y estado final
            doc.add_paragraph("\n3. Observaciones y Estado Final", style='Heading 2')
            doc.add_paragraph(f"Estado final: {data['estado']}", style='BodyText')
            doc.add_paragraph("Observaciones:", style='BodyText')
            doc.add_paragraph(data['observaciones'], style='List Bullet')
            
            # Pie de página
            doc.add_paragraph("\n")
            footer = doc.add_paragraph("Documento generado automáticamente por el Sistema de Gestión de Stock")
            footer.runs[0].italic = True
            footer.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            # Guardar documento
            doc.save(filename)
            return True
            
        except Exception as e:
            logging.error(f"Error al generar informe técnico: {str(e)}", exc_info=True)
            return False



@staticmethod
def export_informe_toner(data, filename):
    """Exporta un informe de toner a Word con formato profesional"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        # Crear documento
        doc = Document()
        
        # Configurar márgenes
        section = doc.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        # Encabezado con logo institucional
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN.CENTER
        
        # Intentar agregar logo (si existe)
        try:
            logo_path = None
            possible_paths = ['logo_pj.png', 'media/logo_pj.png']
            for path in possible_paths:
                if os.path.exists(path):
                    logo_path = path
                    break
            
            if logo_path:
                header.add_run().add_picture(logo_path, width=Inches(0.7))
                header.add_run().add_break()
        except Exception as e:
            logging.warning(f"No se pudo agregar logo: {e}")
        
        # Títulos institucionales
        titles = doc.add_paragraph()
        titles.alignment = WD_ALIGN.CENTER
        
        # Línea 1 - PODER JUDICIAL
        line1 = titles.add_run("PODER JUDICIAL DE LA PROVINCIA DE JUJUY\n")
        line1.font.size = Pt(14)
        line1.font.bold = True
        
        # Línea 2 - DEPARTAMENTO
        line2 = titles.add_run("DEPARTAMENTO DE SISTEMAS Y TECNOLOGÍA DE LA INFORMACIÓN\n")
        line2.font.size = Pt(12)
        line2.font.bold = True
        
        # Línea 3 - Dirección
        line3 = titles.add_run("Argañaraz esq. Independencia -- San Salvador de Jujuy\n\n")
        line3.font.size = Pt(10)
        line3.font.bold = False
        
        # Título del informe
        title = doc.add_paragraph(data['titulo'])
        title.alignment = WD_ALIGN.CENTER
        title.runs[0].font.size = Pt(14)
        title.runs[0].font.bold = True
        
        # Subtítulo
        if data.get('subtitulo'):
            subtitle = doc.add_paragraph(data['subtitulo'])
            subtitle.alignment = WD_ALIGN.CENTER
            subtitle.runs[0].italic = True
            doc.add_paragraph()
        
        # Fecha de generación
        fecha_gen = doc.add_paragraph()
        fecha_gen.alignment = WD_ALIGN.RIGHT
        fecha_gen.add_run(f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        
        # Tabla de datos principales
        if data['datos']:
            doc.add_paragraph("\nDatos Detallados:", style='Heading 2')
            
            table = doc.add_table(rows=1, cols=len(data['encabezados']), style='Light Shading Accent 1')
            table.autofit = True
            
            # Encabezados
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(data['encabezados']):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Datos
            for row in data['datos']:
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
                    row_cells[i].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Resumen
        if data.get('resumen'):
            doc.add_paragraph("\nResumen:", style='Heading 2')
            
            if isinstance(data['resumen'][0], tuple) and len(data['resumen'][0]) == 3:
                # Resumen de consumos (marca, modelo, cantidad)
                resumen_table = doc.add_table(rows=1, cols=3, style='Light Shading Accent 1')
                resumen_table.autofit = True
                
                # Encabezados
                hdr_cells = resumen_table.rows[0].cells
                hdr_cells[0].text = "Marca"
                hdr_cells[1].text = "Modelo"
                hdr_cells[2].text = "Cantidad Total"
                
                for cell in hdr_cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Datos
                for item in data['resumen']:
                    row_cells = resumen_table.add_row().cells
                    row_cells[0].text = item[0]
                    row_cells[1].text = item[1]
                    row_cells[2].text = str(item[2])
                    
                    for cell in row_cells:
                        cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            
            elif isinstance(data['resumen'][0], tuple) and len(data['resumen'][0]) == 3:
                # Resumen de recargas (empresa, estado, cantidad)
                resumen_table = doc.add_table(rows=1, cols=3, style='Light Shading Accent 1')
                resumen_table.autofit = True
                
                # Encabezados
                hdr_cells = resumen_table.rows[0].cells
                hdr_cells[0].text = "Empresa"
                hdr_cells[1].text = "Estado"
                hdr_cells[2].text = "Cantidad Total"
                
                for cell in hdr_cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Datos
                for item in data['resumen']:
                    row_cells = resumen_table.add_row().cells
                    row_cells[0].text = item[0]
                    row_cells[1].text = item[1]
                    row_cells[2].text = str(item[2])
                    
                    for cell in row_cells:
                        cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Pie de página
        doc.add_page_break()
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Documento generado automáticamente el {datetime.now().strftime('%d/%m/%Y a las %H:%M')}"
        footer_para.alignment = WD_ALIGN.RIGHT
        
        # Guardar
        doc.save(filename)
        return True
    except Exception as e:
        logging.error(f"Error al exportar informe de toner: {str(e)}", exc_info=True)
        return False


@staticmethod
def export_informe_toner(data, filename):
    """Exporta un informe de toner a Word con formato profesional"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        # Crear documento
        doc = Document()
        
        # Configurar márgenes
        section = doc.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        # Encabezado con logo institucional
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN.CENTER
        
        # Intentar agregar logo (si existe)
        try:
            logo_path = None
            possible_paths = ['logo_pj.png', 'media/logo_pj.png']
            for path in possible_paths:
                if os.path.exists(path):
                    logo_path = path
                    break
            
            if logo_path:
                header.add_run().add_picture(logo_path, width=Inches(0.7))
                header.add_run().add_break()
        except Exception as e:
            logging.warning(f"No se pudo agregar logo: {e}")
        
        # Títulos institucionales
        titles = doc.add_paragraph()
        titles.alignment = WD_ALIGN.CENTER
        
        # Línea 1 - PODER JUDICIAL
        line1 = titles.add_run("PODER JUDICIAL DE LA PROVINCIA DE JUJUY\n")
        line1.font.size = Pt(14)
        line1.font.bold = True
        
        # Línea 2 - DEPARTAMENTO
        line2 = titles.add_run("DEPARTAMENTO DE SISTEMAS Y TECNOLOGÍA DE LA INFORMACIÓN\n")
        line2.font.size = Pt(12)
        line2.font.bold = True
        
        # Línea 3 - Dirección
        line3 = titles.add_run("Argañaraz esq. Independencia -- San Salvador de Jujuy\n\n")
        line3.font.size = Pt(10)
        line3.font.bold = False
        
        # Título del informe
        title = doc.add_paragraph(data['titulo'])
        title.alignment = WD_ALIGN.CENTER
        title.runs[0].font.size = Pt(14)
        title.runs[0].font.bold = True
        
        # Subtítulo
        if data.get('subtitulo'):
            subtitle = doc.add_paragraph(data['subtitulo'])
            subtitle.alignment = WD_ALIGN.CENTER
            subtitle.runs[0].italic = True
            doc.add_paragraph()
        
        # Fecha de generación
        fecha_gen = doc.add_paragraph()
        fecha_gen.alignment = WD_ALIGN.RIGHT
        fecha_gen.add_run(f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        
        # Tabla de datos principales
        if data['datos']:
            doc.add_paragraph("\nDatos Detallados:", style='Heading 2')
            
            table = doc.add_table(rows=1, cols=len(data['encabezados']), style='Light Shading Accent 1')
            table.autofit = True
            
            # Encabezados
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(data['encabezados']):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Datos
            for row in data['datos']:
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
                    row_cells[i].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Resumen
        if data.get('resumen'):
            doc.add_paragraph("\nResumen:", style='Heading 2')
            
            if isinstance(data['resumen'][0], tuple) and len(data['resumen'][0]) == 3:
                # Resumen de consumos (marca, modelo, cantidad)
                resumen_table = doc.add_table(rows=1, cols=3, style='Light Shading Accent 1')
                resumen_table.autofit = True
                
                # Encabezados
                hdr_cells = resumen_table.rows[0].cells
                hdr_cells[0].text = "Marca"
                hdr_cells[1].text = "Modelo"
                hdr_cells[2].text = "Cantidad Total"
                
                for cell in hdr_cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Datos
                for item in data['resumen']:
                    row_cells = resumen_table.add_row().cells
                    row_cells[0].text = item[0]
                    row_cells[1].text = item[1]
                    row_cells[2].text = str(item[2])
                    
                    for cell in row_cells:
                        cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            
            elif isinstance(data['resumen'][0], tuple) and len(data['resumen'][0]) == 3:
                # Resumen de recargas (empresa, estado, cantidad)
                resumen_table = doc.add_table(rows=1, cols=3, style='Light Shading Accent 1')
                resumen_table.autofit = True
                
                # Encabezados
                hdr_cells = resumen_table.rows[0].cells
                hdr_cells[0].text = "Empresa"
                hdr_cells[1].text = "Estado"
                hdr_cells[2].text = "Cantidad Total"
                
                for cell in hdr_cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Datos
                for item in data['resumen']:
                    row_cells = resumen_table.add_row().cells
                    row_cells[0].text = item[0]
                    row_cells[1].text = item[1]
                    row_cells[2].text = str(item[2])
                    
                    for cell in row_cells:
                        cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Pie de página
        doc.add_page_break()
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Documento generado automáticamente el {datetime.now().strftime('%d/%m/%Y a las %H:%M')}"
        footer_para.alignment = WD_ALIGN.RIGHT
        
        # Guardar
        doc.save(filename)
        return True
    except Exception as e:
        logging.error(f"Error al exportar informe de toner: {str(e)}", exc_info=True)
        return False   


        

    

if __name__ == "__main__":
    try:
        root = tk.Tk()
        app = MainController(root)
        root.mainloop()
    except Exception as e:
        logging.critical(f"Error crítico: {e}", exc_info=True)
        messagebox.showerror("Error", f"Error crítico: {str(e)}")
  


    
